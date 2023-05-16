# importing libraries
from __future__ import print_function
import glob
import sys
import io
import subprocess
import warnings
from sys import exit
import docx
import msoffcrypto
import numpy as np
import pandas as pd
import streamlit as st
from PIL import Image
from io import BytesIO
from mailmerge import MailMerge
from datetime import date as to_date
import zipfile


warnings.filterwarnings("ignore")


def format_sheet(file_path: str, template_path: str, sheet_num_template=0, run=0):
    """
    file_path: the file you want to design, must bx xlsx
    template_path: the teplate of the design, must bx xlsx
    """
    # format sheet
    xlsx = pd.ExcelWriter(file_path, engine="openpyxl", mode="a")
    for i in range(len(xlsx.book.worksheets)):
        new_sheet = xlsx.book.worksheets[run]
        template = pd.ExcelWriter(template_path, engine="openpyxl", mode="a")
        default_sheet = template.book.worksheets[sheet_num_template]
        from copy import copy

        new_sheet.sheet_view.rightToLeft = True
        for row in default_sheet.rows:
            for cell in row:
                new_cell = new_sheet.cell(row=cell.row, column=cell.col_idx)
                if cell.has_style and new_cell.value != None:
                    new_cell.font = copy(cell.font)
                    new_cell.border = copy(cell.border)
                    new_cell.fill = copy(cell.fill)
                    new_cell.number_format = copy(cell.number_format)
                    new_cell.protection = copy(cell.protection)
                    new_cell.alignment = copy(cell.alignment)
        from openpyxl.utils import get_column_letter

        for i in range(default_sheet.max_column):
            new_sheet.column_dimensions[
                get_column_letter(i + 1)
            ].width = default_sheet.column_dimensions[get_column_letter(i + 1)].width

    xlsx.save()


def change_names_and_order(file_path: str, df: pd.DataFrame):
    """
    name of columns in names must be "old_name", "new_name"
    """
    match_names = pd.read_excel(file_path)
    match_names.index = match_names.old_name
    dict_names = match_names.to_dict()["new_name"]
    ls_names = match_names["new_name"].tolist()
    df.rename(columns=dict_names, inplace=True)
    df = df[ls_names]
    return df


def report_num_gen(business_id, branch_num):
    today = str(to_date.today())
    date_list = today.split("-")
    start_date = date_list[2] + date_list[1] + date_list[0]
    end_date = date_list[0] + date_list[1] + date_list[2]
    time = str("120000")
    report_name = (
        str(start_date)
        + "-"
        + str(time)
        + "-"
        + str(end_date)
        + "-E-"
        + str(business_id)
        + "-"
        + str(branch_num)
        + "-UAR-UST"
    )
    return report_name


def clients2columns(df, clients_df):
    # Sort the dataframe by 'Number'
    df = df.sort_values(by="cumsum")
    # Create an empty list
    data = []
    # Iterate over the unique numbers in the 'Number' column
    for number, group in df.groupby("cumsum"):
        # Append the first name of each group to the data list
        data.append({"cumsum": number, "clients name_0": group.iloc[0]["clients name"]})
        # Append the remaining names of each group as new columns
        for i, name in enumerate(group["clients name"][1:]):
            data[-1][f"clients name_{i + 1}"] = name
    # Create the final dataframe
    result = pd.DataFrame.from_records(data)
    # Iterate over columns that starts with 'Client_'
    for col in result.columns:
        if col.startswith("clients name_"):
            result[col] = result[col].astype(
                str
            )  # todo was added for shery and alex if troubles please check
            result = pd.merge(
                result,
                clients_df[["clients name", "id number", "city", "address"]],
                left_on=col,
                right_on="clients name",
                how="left",
                suffixes=("", col),
            )
            # Drop the 'Name' column from the result dataframe
            result = result.drop(columns="clients name")
            result = result.drop_duplicates(subset="cumsum")
    return result


def clients2columns_gmt(df):
    # Sort the dataframe by 'Number'
    df = df.sort_values(by="owner name")
    # Create an empty list
    data = []
    # Iterate over the unique numbers in the 'Number' column
    for number, group in df.groupby("owner name"):
        # Append the first name of each group to the data list
        data.append(
            {"owner name": number, "clients name_0": group.iloc[0]["clients name"]}
        )
        # Append the remaining names of each group as new columns
        for i, name in enumerate(group["clients name"][1:]):
            data[-1][f"clients name_{i + 1}"] = name
    # Create the final dataframe
    result = pd.DataFrame.from_records(data)
    # Iterate over columns that starts with 'Client_'
    for col in result.columns:
        if col.startswith("clients name_"):
            result = pd.merge(
                result,
                df[["clients name", "id number"]],
                left_on=col,
                right_on="clients name",
                how="left",
                suffixes=("", col),
            )
            # Drop the 'Name' column from the result dataframe
            result = result.drop(columns="clients name")

    return result


def clients2columns_changemat(df, clients_df):
    # Sort the dataframe by 'Number'
    df = df.sort_values(by="cumsum_ow")
    # Create an empty list
    data = []
    # Iterate over the unique numbers in the 'Number' column
    for number, group in df.groupby("cumsum_ow"):
        # Append the first name of each group to the data list
        data.append(
            {"cumsum_ow": number, "clients name_0": group.iloc[0]["clients name"]}
        )
        # Append the remaining names of each group as new columns
        for i, name in enumerate(group["clients name"][1:]):
            data[-1][f"clients name_{i + 1}"] = name
    # Create the final dataframe
    result = pd.DataFrame.from_records(data)
    # Iterate over columns that starts with 'Client_'
    for col in result.columns:
        if col.startswith("clients name_"):
            result = pd.merge(
                result,
                clients_df[["clients name", "id number", "city", "address"]],
                left_on=col,
                right_on="clients name",
                how="left",
                suffixes=("", col),
            )
            # Drop the 'Name' column from the result dataframe
            result = result.drop(columns="clients name")
            result = result.drop_duplicates(subset="cumsum_ow")
    return result


def run_yeshut(check, client, business_info, reported):
    directories = pd.read_excel("./directories.xlsx")
    directories = pd.DataFrame(directories)
    st.dataframe(directories)
    directories = directories.set_index("File type")
    rc_directory = directories.loc["risk countries"].values[0]
    fsp_directory = directories.loc["financial service providers"].values[0]

    template_directory = directories.loc["word template"].values[0]

    yeshut_directory = directories.loc["yeshut content and change names"].values[0]

    # reading the risk country file
    filepath_check = fsp_directory + "\*נותני*"
    txt_check = glob.glob(filepath_check)
    for textfile in txt_check:
        fsp_list = pd.read_excel(textfile)

    # reading the risk country file
    filepath_check = rc_directory + "\*מדינות*"
    txt_check = glob.glob(filepath_check)
    for textfile in txt_check:
        rc = pd.read_excel(textfile)

    st.dataframe(rc)
    # reading the report content file
    filepath_check = yeshut_directory + "\*content*"
    txt_check = glob.glob(filepath_check)
    for textfile in txt_check:
        content = pd.read_excel(textfile)

    # renaming the columns and reorganizing the bank account numbers
    try:
        check["account number"] = check["ב-ס-ח"].str.split("-", 2, expand=True)[2]
        check["branch number"] = check["ב-ס-ח"].str.split("-", 2, expand=True)[1]
        check["bank number"] = check["ב-ס-ח"].str.split("-", 2, expand=True)[0]
    except:
        check["account number"] = check["account number"]
        check["branch number"] = check["branch number"]
        check["bank number"] = check["bank number"]
    check = check.rename(
        columns={
            "עסקה": "deal number",
            "תאריך עיסקה": "date",
            "תאריך עסקה": "date",
            "סוג עיסקה": "type",
            "שם לקוח": "clients name",
            "מספר שיק": "check number",
            "שם מושך": "owner name",
            "ת.פירעון": "delay deposit date",
            "סטטוס": "status",
            "סכום": "nis sum",
        }
    )
    check = check[check["status"] != "הושב ללקוח"]
    # check = check[['deal number','date','type',
    #                 'clients name','check number','owner name',
    #                 'delay deposit date','status','nis sum']]
    check["N"] = check.reset_index().index
    # renaming the columns of the clients file
    client = client.rename(
        columns={
            "קוד לקוח": "clients symbol",
            "שם לקוח": "clients name",
            "מ.מזהה": "id number",
            "תאריך לידה": "dob",
            "מין": "sex",
            "רחוב": "address",
            "עיר": "city",
            "תושבות": "citizenship",
        }
    )

    # drop na
    check = check.dropna(subset="nis sum")
    check = check.dropna(subset="deal number")
    # check nis sum is int
    check["nis sum"] = check["nis sum"].astype(int)
    check["date"] = pd.to_datetime(check["date"], dayfirst=True)
    # stripping client names from client and template df
    client["clients name"] = client["clients name"].str.strip()
    check["clients name"] = check["clients name"].str.strip()
    rc["עברית"] = rc["עברית"].str.strip()

    # reported strip
    reported["שם"] = reported["שם"].str.strip()
    # sort the data in check by deal date and by name
    check["deal number"] = check["deal number"].astype(int)

    check = check.sort_values(["clients name", "date"])

    # adding the sum for every deal number and marking every deal under 50k as True
    deal_sum = check.groupby("deal number").agg({"nis sum": "sum"})
    deal_sum_ow = check.groupby("deal number").agg({"nis sum": "sum"})
    check = check.merge(deal_sum, on="deal number", how="left", suffixes=("", "_deal"))
    check = check.merge(
        deal_sum_ow, on="deal number", how="left", suffixes=("", "_deal_ow")
    )
    # Adding a column stating if the transaction is under or over 50K(False = over 50K)
    check["U50"] = np.where(check["nis sum_deal"] >= 50000, False, True)
    check["U50_ow"] = np.where(check["nis sum_deal_ow"] >= 50000, False, True)
    # Adding a column stating if the transaction is under or over 5K(False = over 5K)
    check["U5"] = np.where(check["nis sum_deal"] >= 5000, False, True)
    check["U5_ow"] = np.where(check["nis sum_deal_ow"] >= 5000, False, True)

    # creating a column with the dates difference by day and creating a new column with the value as a number
    check["date"] = pd.to_datetime(check["date"])
    check["date_diff"] = check["date"].diff()
    check["date_diff_fl"] = check["date_diff"] / pd.to_timedelta(1, unit="D")

    # testing for matching clients in match column
    check["match"] = check["clients name"].eq(check["clients name"].shift())

    # adding a conditional colum based on date diff values
    check["diff_check"] = np.where(
        (check["date_diff_fl"] > 3)
        | (check["date_diff_fl"] < 0)
        | (check["match"] == False),
        1,
        0,
    )

    # cumulative sum as a way to count if the values in the diff check
    check["cumsum"] = np.cumsum(check["diff_check"])
    check = check.drop_duplicates("N")

    # striping blank spaces from citizenship  country
    try:
        client["citizenship"] = client["citizenship"].str.strip()
    except AttributeError:
        client["citizenship"] = client["citizenship"].fillna("ישראל")
        client["citizenship"] = client["citizenship"].str.strip()

    try:
        # id number as type int
        client["id number"] = np.where(
            client["id number"].isna(), 123456789, client["id number"]
        )
        # locating palestinian authority risk clients
        client["citizenship_ps"] = np.where(
            (client["id number"] > 800000000) & (client["id number"] < 999999999),
            "הרשות הפלשטינית",
            client["citizenship"],
        )
        client["citizenship"] = client["citizenship_ps"]
    except TypeError:
        error1 = st.error("חלק ממספרי תעודות הזהות מכילים אותיות")

    client["id number"] = client["id number"].apply(
        lambda x: round(x) if isinstance(x, float) else x
    )

    # building the monitoring df with sum per client, deal distinct count and under 50k deal count
    mont = check.groupby(["clients name", "cumsum"]).agg(
        {"nis sum": "sum", "deal number": "nunique", "U50": "sum", "U5": "sum"}
    )
    mont = mont.reset_index()

    # sorting the values by sum
    mont = mont.sort_values("nis sum", ascending=False)
    mont["U50"] = mont["U50"].astype(int)
    mont["U5"] = mont["U5"].astype(int)

    # preform a reported in the last two month test
    mont = mont.merge(reported["שם"], left_on="clients name", right_on="שם", how="left")

    # adding a column named status based on unique count of deal number and at least on tr under 50k and not reported in the past 2m
    mont["status"] = np.where(
        (mont["deal number"] == 1) | (mont["U50"] == 0) | (mont["שם"].notna()),
        "Regular Report",
        "Check",
    )
    mont = mont.drop_duplicates("cumsum")

    # preform a financial service providers test - results stored in a column named fsp_check
    # adding the id number to the monitoring sheet
    mont = mont.merge(
        client[["id number", "clients name"]], on="clients name", how="left"
    )

    # adding a column מספר מזהה to the monitoring sheet values indicates fsp(financial service provider)
    try:
        mont = mont.merge(
            fsp_list["מספר מזהה"], left_on="id number", right_on="מספר מזהה", how="left"
        )
    except ValueError:
        mont = mont.merge(
            fsp_list["מספר מזהה"].astype(str),
            left_on="id number",
            right_on="מספר מזהה",
            how="left",
        )
    mont["fsp_check"] = np.where(
        mont["מספר מזהה"].isna(), "check", "financial service provider"
    )

    mont = mont.drop_duplicates("cumsum")

    # preform a risk country test
    # merging client detail with monitoring list
    mont_client = mont.merge(client, on="clients name", how="left")

    # a summarized data frame of clients name and country
    mont_client_country = mont_client[["clients name", "citizenship"]]

    # merging the list of risk country
    mont_client_country_risk = mont_client_country.merge(
        rc, left_on="citizenship", right_on="עברית", how="left"
    )

    # merging the risk status for each client to the monitoring df and dropping the cumsum duplicates
    mont = mont.merge(
        mont_client_country_risk[["clients name", "סיכון/ללא דיווח רגיל"]],
        on="clients name",
        how="left",
    )
    mont = mont.drop_duplicates("cumsum")

    # reported test
    mont["reported"] = np.where(mont["שם"].notna(), "reported", "check")

    # marking dirdos transaction as d in a new column named dirdos
    mont["dirdos"] = np.where(
        (mont["nis sum"] >= 47000)
        & (mont["nis sum"] < 50000)
        & (mont["שם"].isna())
        & (mont["fsp_check"] == "check")
        & (mont["U50"] > 0),
        "d",
        "not d",
    )
    mont["dirdos_risk"] = np.where(
        (mont["nis sum"] >= 4700)
        & (mont["nis sum"] < 5000)
        & (mont["שם"].isna())
        & (mont["fsp_check"] == "check")
        & (mont["U5"] > 0),
        "dr",
        "not dr",
    )

    # creating a new data frame containing only clients that need to be reported
    mont2r = mont[
        ((mont["status"] == "Check") | (mont["dirdos"] == "d"))
        & (mont["fsp_check"] == "check")
        & (mont["nis sum"] >= 47000)
    ]
    mont2r = mont2r.drop_duplicates("cumsum")

    # adding the type of the report p-pitzul d-dirdos
    mont2r["type"] = np.where(mont2r["nis sum"] >= 50000, "p", "d")

    # figure out the problem of a client with multiple types of reports
    type_count = mont2r[["clients name", "type"]].groupby("clients name").nunique()

    # adding situations in which there are two types of reports for a client
    mont2r = mont2r.merge(
        type_count, on="clients name", how="left", suffixes=("", "_count")
    )
    mont2r["final type"] = np.where(mont2r["type_count"] > 1, "pd", mont2r["type"])

    # creating a new data frame with duplicate cumsum for report details df= mont2report_w_cumsum
    mont2report_w_cumsum = mont2r
    mont2report = mont2r.drop_duplicates("clients name")

    # creating a new data frame containing only clients that need to be reported for risk clients
    mont2rr = mont[
        ((mont["status"] == "Check") | (mont["dirdos_risk"] == "dr"))
        & (mont["U5"] != 0)
        & (mont["fsp_check"] == "check")
        & (mont["nis sum"] >= 4700)
        & (mont["סיכון/ללא דיווח רגיל"].notna())
    ]
    mont2rr = mont2rr.drop_duplicates("cumsum")

    # adding the type of the report pr-pitzul dr-dirdos (risk)
    mont2rr["final type"] = np.where((mont2rr["nis sum"] >= 5000), "pr", "dr")

    # adding the risk transctions report to the df with cumsum
    mont2report_w_cumsum = mont2report_w_cumsum.append(mont2rr)

    # sort the data in check by deal date and by owner name
    check["deal number"] = check["deal number"].astype(int)
    check_ow = check.sort_values(["owner name", "date"])

    # delete the fsp client from check_ow df
    check_ow = check_ow.merge(
        client[["clients name", "id number"]], on="clients name", how="left"
    )

    # deleting FSP from check ow:
    try:
        check_ow = check_ow.merge(
            fsp_list["מספר מזהה"], left_on="id number", right_on="מספר מזהה", how="left"
        )
    except ValueError:
        check_ow = check_ow.merge(
            fsp_list["מספר מזהה"].astype(str),
            left_on="id number",
            right_on="מספר מזהה",
            how="left",
        )
    check_ow = check_ow[check_ow["מספר מזהה"].isna()]

    # creating a column with the dates difference by day and creating a new column with the value as a number
    check_ow["date_diff"] = check_ow["date"].diff()
    check_ow["date_diff_fl"] = check_ow["date_diff"] / pd.to_timedelta(1, unit="D")

    # creating a column with a test for matching owner in the past transaction
    check_ow["match_ow"] = check_ow["owner name"].eq(check_ow["owner name"].shift())

    # adding a conditional colum based on date diff values
    check_ow["diff_check"] = np.where(
        (check_ow["date_diff_fl"] > 3)
        | (check_ow["date_diff_fl"] < 0)
        | (check_ow["match_ow"] == False),
        1,
        0,
    )

    # cumulative sum as a way to count if the values in the diff check
    check_ow["cumsum"] = np.cumsum(check_ow["diff_check"])
    check_ow = check_ow.drop_duplicates("N")

    # building the monitoring df with sum per client and deal distinct count
    mont_ow_risk = check_ow.groupby(["owner name", "cumsum"]).agg(
        {
            "nis sum": "sum",
            "deal number": "nunique",
            "clients name": "nunique",
            "U5_ow": "sum",
        }
    )

    # sorting the values by sum
    mont_ow_risk = mont_ow_risk.sort_values("nis sum", ascending=False)
    mont_ow_risk["status"] = np.where(
        (mont_ow_risk["deal number"] == 1)
        | (mont_ow_risk["clients name"] == 1)
        | (mont_ow_risk["nis sum"] <= 5000)
        | (mont_ow_risk["U5_ow"] == 0),
        "Regular Report",
        "Check",
    )
    mont_ow_risk = mont_ow_risk[mont_ow_risk["status"] == "Check"]

    # shared owner risk transactions
    mont_ow_risk_merge = mont_ow_risk.merge(
        check_ow[["clients name", "cumsum"]], on="cumsum", how="left"
    )
    mont_ow_risk_merge = mont_ow_risk_merge.merge(
        client[["clients name", "citizenship"]],
        left_on="clients name_y",
        right_on="clients name",
        how="left",
    )
    mont_ow_risk_merge = mont_ow_risk_merge.merge(
        rc[["עברית", "סיכון/ללא דיווח רגיל"]],
        left_on="citizenship",
        right_on="עברית",
        how="left",
    )
    mont_ow_risk_merge = mont_ow_risk_merge[
        mont_ow_risk_merge["סיכון/ללא דיווח רגיל"].notna()
    ]
    mont_ow_clients_w_cumsum = mont_ow_risk_merge
    mont_ow_risk_merge = mont_ow_risk_merge.drop_duplicates("clients name")
    mont_ow_risk_merge = mont_ow_risk_merge.drop_duplicates("cumsum")
    mont_ow_risk_merge["final type_ow"] = "mr"

    # building the monitoring df with sum per client and deal distinct count
    mont_ow = check_ow.groupby(["owner name", "cumsum"]).agg(
        {
            "nis sum": "sum",
            "deal number": "nunique",
            "clients name": "nunique",
            "U50_ow": "sum",
        }
    )

    # sorting the values by sum
    mont_ow = mont_ow.sort_values("nis sum", ascending=False)

    # adding a column named status based on unique count of deal number
    mont_ow["status"] = np.where(
        (mont_ow["deal number"] == 1)
        | (mont_ow["clients name"] == 1)
        | (mont_ow["nis sum"] < 50000)
        | (mont_ow["U50_ow"] == 0),
        "Regular Report",
        "Check",
    )
    mont_ow = mont_ow.reset_index()

    # leaving only suspicious clients
    mont_ow_test = mont_ow
    mont_ow = mont_ow[mont_ow["status"] == "Check"]

    # creating a list of the shared owner report clients name
    mont_ow_merge = mont_ow.merge(
        check_ow[["clients name", "cumsum"]], on="cumsum", how="left"
    )

    # subtracting past 3 month reported clients
    mont_ow_clients = mont_ow_merge.merge(
        reported["שם"], how="left", left_on="clients name_y", right_on="שם"
    )
    mont_ow_clients = mont_ow_clients[mont_ow_clients["שם"].isna()]
    mont_ow_clients = mont_ow_clients[
        -mont_ow_clients["cumsum"].isin(mont_ow_risk_merge["cumsum"])
    ]
    mont_ow_clients = mont_ow_clients.rename(columns={"clients name_y": "clients name"})
    mont_ow_clients_w_cumsum = mont_ow_clients_w_cumsum.append(mont_ow_clients)
    mont_ow_clients = mont_ow_clients.drop_duplicates("cumsum")
    mont_ow_clients = mont_ow_clients.drop_duplicates("clients name")
    mont_ow_clients = mont_ow_clients.reset_index()
    mont_ow_clients["final type"] = "m"

    # normal risk reports
    mont_normal = mont[mont["סיכון/ללא דיווח רגיל"] == "סיכון ללא דיווח רגיל"]
    mont_normal2report = mont_normal[
        ((mont_normal["deal number"] == 1) & (mont_normal["nis sum"] >= 5000))
    ]
    mont_normal2report = mont_normal2report[
        -mont_normal2report["clients name"].isin(mont_ow_clients["clients name"])
    ]
    mont_normal2report = mont_normal2report[
        -mont_normal2report["clients name"].isin(mont_ow_risk_merge["clients name"])
    ]
    mont_normal2report["final type normal"] = "nr"
    mont2report_w_cumsum = mont2report_w_cumsum.append(mont_normal2report)
    # a new data frame containing all kinds of reports
    all_clients2r = mont_ow_clients[
        ["clients name", "final type", "cumsum", "nis sum"]
    ].merge(
        mont2report[["clients name", "final type", "cumsum", "nis sum"]],
        on="clients name",
        how="outer",
        suffixes=("_ow", "_p"),
    )
    all_clients2r = all_clients2r.merge(
        mont2rr[["final type", "clients name", "cumsum", "nis sum"]],
        how="outer",
        on="clients name",
        suffixes=("", "_risk"),
    )
    all_clients2r = all_clients2r.merge(
        mont_ow_risk_merge[["final type_ow", "clients name", "cumsum", "nis sum"]],
        how="outer",
        on="clients name",
        suffixes=("", "_riskow"),
    )
    all_clients2r = all_clients2r.merge(
        mont_normal2report[["final type normal", "clients name", "cumsum", "nis sum"]],
        how="outer",
        on="clients name",
        suffixes=("", "_normal"),
    )

    all_clients2r[
        [
            "final type_p",
            "final type_ow",
            "final type",
            "final type_ow_riskow",
            "final type normal",
        ]
    ] = all_clients2r[
        [
            "final type_p",
            "final type_ow",
            "final type",
            "final type_ow_riskow",
            "final type normal",
        ]
    ].fillna(
        ""
    )
    all_clients2r["finaltypedown"] = (
        all_clients2r["final type_p"]
        + all_clients2r["final type_ow"]
        + all_clients2r["final type"]
        + all_clients2r["final type_ow_riskow"]
        + all_clients2r["final type normal"]
    )

    merged = all_clients2r.drop_duplicates("clients name").reset_index(drop=True)
    merged["cumsum"] = merged["cumsum"].fillna(merged["cumsum_p"])
    merged["cumsum"] = merged["cumsum"].fillna(merged["cumsum_ow"])
    merged["cumsum"] = merged["cumsum"].fillna(merged["cumsum_riskow"])
    merged["nis sum"] = merged["nis sum"].fillna(merged["nis sum_p"])
    merged["nis sum"] = merged["nis sum"].fillna(merged["nis sum_ow"])
    merged["nis sum"] = merged["nis sum"].fillna(merged["nis sum_riskow"])

    # creating the client details for reports data frame
    info2r = (
        merged[["clients name", "finaltypedown", "nis sum", "cumsum"]]
        .merge(client, on="clients name", how="left")
        .dropna(axis=1, how="all")
    )
    try:
        info2r["clients name"] = info2r["clients name"].str.replace("\d+", "")
        info2r["country"] = "ישראל"
    except KeyError:
        error2 = st.error("אין דיווחים בלתי רגילים")

    if "sex" not in info2r:
        try:
            info2r["id number"] = info2r["id number"].fillna(0)
            info2r["id number"] = info2r["id number"].astype(int)
            info2r["sex"] = np.where(
                (
                    (info2r["id number"] >= 500000000)
                    & (info2r["id number"] <= 600000000)
                ),
                "תאגיד",
                "",
            )
        except:
            error3 = st.error("בחלק ממספרי הזהות יש אותיות")

    else:
        info2r["sex"] = info2r["sex"]

    # fixing the business info
    try:
        business_info = business_info.fillna("")
        business_type = business_info.loc["סוג הגוף המדווח"].values[0]
        business_name = business_info.loc["שם מלא של הגורם המדווח"].values[0]
        business_id = business_info.loc["מספר מזהה של הגורם המדווח"].values[0]
        branch_number = business_info.loc["מספר סניף מדווח"].values[0]
        business_address = business_info.loc["מען הסניף"].values[0]
        date = to_date.today()
        business_phone = business_info.loc["טלפון ופקסימילה של הגורם המדווח"].values[0]
        worker_name = business_info.loc["שם פרטי ושם משפחה של עובד עורך הדיווח"].values[
            0
        ]
        worker_id = business_info.loc['ת"ז של עורך הדיווח'].values[0]
        worker_position = business_info.loc["תפקיד עורך הדיווח"].values[0]
        worker_phone = business_info.loc["טלפון עורך הדיווח"].values[0]
        worker_email = business_info.loc["דואר אלקטרוני"].values[0]
    except:
        error4 = st.error("ישנה בעיה בקובץ ה-business information")

    if info2r.empty:
        check4excel = check
        check4excel = check4excel[
            [
                "deal number",
                "date",
                "clients name",
                "check number",
                "bank number",
                "branch number",
                "account number",
                "owner name",
                "nis sum",
            ]
        ]

        check4excel = check4excel.rename(
            columns={
                "deal number": "מזהה פעולה",
                "date": "תאריך פעולה",
                "clients name": "שם הלקוח",
                "check number": "מספר צק",
                "bank number": "מספר בנק",
                "branch number": "מספר סניף",
                "account number": "מספר חשבון",
                "owner name": "שם המושך",
                "nis sum": "סכום בשקלים",
            }
        )
        check4excel[["סטטוס", "סטטוס מושך"]] = "תקין"

        xlsx_data = BytesIO()
        with pd.ExcelWriter(xlsx_data, engine="openpyxl") as writer:
            check4excel.to_excel(writer, sheet_name="ניהול התראות")
        xlsx_data.seek(0)
        st.download_button(
            label="📥 Download Final Report",
            data=xlsx_data,
            file_name="Final Report.xlsx",
        )
        st.error("אין דיווחים בלתי רגילים")
        exit()

    # dropping duplicates and resetting index
    info2r = info2r.drop_duplicates("clients name")
    info2r = info2r.reset_index(drop=True)
    # importing the report number needed
    info2r["report_id_aid"] = reported["מס דיווח"].max() + 1
    info2r["report_id"] = (info2r["report_id_aid"] + range(len(info2r.index))).astype(
        int
    )
    info2r["report_name"] = report_num_gen(business_id, branch_number)

    info2r = info2r.merge(content, left_on="finaltypedown", right_on="type", how="left")

    for i in info2r.index:
        info2r["Title"].iloc[i] = str(info2r["Title"].iloc[i]).replace(
            "מדינה", str(info2r["citizenship"].iloc[i]) + ", מדינה"
        )
        info2r["Content"].iloc[i] = str(info2r["Content"].iloc[i]).replace(
            "מדינה", str(info2r["citizenship"].iloc[i]) + ", מדינה"
        )
    if "dob" not in info2r.columns:
        info2r["dob"] = ""
    if "clients symbol" not in info2r.columns:
        info2r["clients symbol"] = ""

    info2r_excel = info2r[
        [
            "clients name",
            "finaltypedown",
            "clients symbol",
            "id number",
            "cumsum",
            "nis sum",
            "dob",
            "sex",
            "citizenship",
            "report_id",
            "Title",
            "Content",
        ]
    ]
    info2r_excel = info2r_excel.rename(
        columns={
            "clients name": "שם הלקוח",
            "finaltypedown": "מהות הפעילות",
            "clients symbol": "מזהה לקוח",
            "id number": "מספר תעודה",
            "dob": "תאריך לידה\התאגדות",
            "sex": "מין",
            "citizenship": "מדינת תעודה\התאגדות",
            "report_id": "מספר דיווח",
            "Title": "תמצית",
            "Content": "תוכן",
            "cumsum": "מזהה קבוצה",
            "nis sum": "סכום פעולות לדיווח",
        }
    )
    info2r_excel["לדווח/לא לדווח"] = ""
    info2r_excel["הערות"] = ""
    info2r_excel["מהות הפעילות"] = info2r_excel["מהות הפעילות"].replace(
        {
            "pd": "פיצול ודירדוס",
            "pr": "פיצול בסיכון",
            "p": "פיצול",
            "dr": "דירדוס בסיכון",
            "d": "דירדוס",
            "m": "מושך משותף",
            "mr": "מושך משותף בסיכון",
            "pm": "פיצול ומושך משותף",
            "dm": "דירדוס ומושך משותף",
            "prmr": "פיצול ומושך משותף בסיכון",
        }
    )
    xlsx_data = BytesIO()
    with pd.ExcelWriter(xlsx_data, engine="openpyxl") as writer:
        info2r_excel.to_excel(writer, sheet_name="ניהול התראות")
    xlsx_data.seek(0)
    # info2r_excel.to_excel("monitoring.xlsx", sheet_name='ניהול התראות')

    # distinguish between company and private name
    info2r["company name"] = np.where(
        (info2r["sex"] == "תאגיד"), info2r["clients name"], ""
    )
    info2r["person name"] = np.where(
        (info2r["sex"] != "תאגיד"), info2r["clients name"], ""
    )

    # distinguish between company and private dob
    try:
        info2r["dob"] = (info2r["dob"]).astype(str)
        info2r["company birth"] = np.where(
            (info2r["sex"] == "תאגיד"), info2r["dob"], ""
        )
        info2r["person birth"] = np.where((info2r["sex"] != "תאגיד"), info2r["dob"], "")
    except KeyError:
        "dob"
        error5 = st.error(
            'עלולה להיות בעיה בקיטלוג התאריכים בין חברות בע"מ ללקוחות פרטיים'
        )

    # distinguish between company and private citizenship
    info2r["company citizenship"] = np.where(
        (info2r["sex"] == "תאגיד"), info2r["citizenship"], ""
    )
    info2r["person citizenship"] = np.where(
        (info2r["sex"] != "תאגיד"), info2r["citizenship"], ""
    )

    # distinguish between company and private citizenship
    info2r["company id"] = np.where((info2r["sex"] == "תאגיד"), info2r["id number"], "")
    info2r["person id"] = np.where((info2r["sex"] != "תאגיד"), info2r["id number"], "")

    try:
        # distinguish between company and private citizenship
        info2r["company id_type"] = np.where(
            (info2r["sex"] == "תאגיד"), info2r["סוג תעודה ראשונה"], ""
        )
        info2r["person id_type"] = np.where(
            (info2r["sex"] != "תאגיד"), info2r["סוג תעודה ראשונה"], ""
        )
    except KeyError:
        info2r["סוג תעודה ראשונה"] = ""
        info2r["company id_type"] = np.where(
            (info2r["sex"] == "תאגיד"), info2r["סוג תעודה ראשונה"], ""
        )
        info2r["person id_type"] = np.where(
            (info2r["sex"] != "תאגיד"), info2r["סוג תעודה ראשונה"], ""
        )

    mont2report_w_cumsum["clients name"] = mont2report_w_cumsum[
        "clients name"
    ].str.replace("\d+", "")

    # info2r = info2r.merge(mont2report_w_cumsum[['clients name', 'cumsum']], how="left", on='clients name')

    check[
        [
            "deal number",
            "check number",
            "bank number",
            "account number",
            "branch number",
        ]
    ] = check[
        [
            "deal number",
            "check number",
            "bank number",
            "account number",
            "branch number",
        ]
    ].astype(
        str
    )
    check_ow[
        [
            "deal number",
            "check number",
            "bank number",
            "account number",
            "branch number",
        ]
    ] = check_ow[
        [
            "deal number",
            "check number",
            "bank number",
            "account number",
            "branch number",
        ]
    ].astype(
        str
    )

    report_tr = mont2report_w_cumsum[["clients name", "cumsum"]]

    check_min = check[
        [
            "deal number",
            "date",
            "clients name",
            "check number",
            "bank number",
            "account number",
            "branch number",
            "owner name",
            "nis sum",
            "cumsum",
        ]
    ]
    report_tr_min = check_min[
        check_min["cumsum"].isin(report_tr["cumsum"])
    ].drop_duplicates("check number")

    i = 0
    for_word_table = []
    # xlsx = pd.ExcelWriter("monitoring.xlsx", engine='openpyxl', mode='a',
    #                       if_sheet_exists='overlay')
    xlsx_test = pd.ExcelWriter(
        xlsx_data, engine="openpyxl", mode="a", if_sheet_exists="overlay"
    )
    # round the sum in ils
    report_tr_min["nis sum"] = report_tr_min["nis sum"].round()

    # creating a unique list of the sender names
    ls_cumsum = report_tr_min["cumsum"].unique().tolist()

    # loop each sender transactions in a unique table
    for name in ls_cumsum:
        df = report_tr_min[report_tr_min["cumsum"] == name]
        data = [["סהכ", df["nis sum"].sum()]]
        last_row = pd.DataFrame(data, columns=["clients name", "nis sum"])
        orderd = pd.concat([df, last_row])
        orderd = orderd[
            [
                "deal number",
                "date",
                "clients name",
                "check number",
                "bank number",
                "account number",
                "branch number",
                "owner name",
                "nis sum",
            ]
        ].fillna(" ")
        orderd = change_names_and_order(
            yeshut_directory + r"/change_names.xlsx", orderd
        )
        # orderd.to_excel(xlsx, sheet_name='פירוט התראות', startrow=i, index=False)
        with xlsx_test as writer:
            orderd.to_excel(writer, sheet_name="פירוט התראות", startrow=i, index=False)
        xlsx_data.seek(0)
        for_word_table.append(orderd)
        i = i + len(orderd) + 2

        # xlsx.save()

    check_min_ow = check_ow[
        [
            "deal number",
            "date",
            "clients name",
            "check number",
            "bank number",
            "account number",
            "branch number",
            "owner name",
            "nis sum",
            "cumsum",
        ]
    ]
    report_tr_min_ow = check_min_ow[
        check_min_ow["cumsum"].isin(mont_ow_clients_w_cumsum["cumsum"])
    ].drop_duplicates("check number")

    j = 0
    for_word_table_ow = []
    # xlsx = pd.ExcelWriter("monitoring.xlsx", engine='openpyxl', mode='a',
    #                       if_sheet_exists='overlay')
    # round the sum in ils
    report_tr_min_ow["nis sum"] = report_tr_min_ow["nis sum"].round()

    # creating a unique list of the sender names
    ls_customer = report_tr_min["clients name"].unique().tolist()
    ls_cumsum_ow = mont_ow_clients_w_cumsum["cumsum"].unique().tolist()

    # loop each sender transactions in a unique table
    for name in ls_cumsum_ow:
        df = report_tr_min_ow[report_tr_min_ow["cumsum"] == name]
        data = [["סהכ", df["nis sum"].sum()]]
        last_row = pd.DataFrame(data, columns=["clients name", "nis sum"])
        orderd = pd.concat([df, last_row])
        orderd = orderd[
            [
                "deal number",
                "date",
                "clients name",
                "check number",
                "bank number",
                "account number",
                "branch number",
                "owner name",
                "nis sum",
            ]
        ].fillna(" ")
        orderd = change_names_and_order(
            yeshut_directory + r"/change_names.xlsx", orderd
        )
        # orderd.to_excel(xlsx, sheet_name='פירוט התראות מושך', startrow=j, index=False)
        with xlsx_test as writer:
            orderd.to_excel(
                writer, sheet_name="פירוט התראות מושך", startrow=j, index=False
            )
        for_word_table_ow.append(orderd)
        j = j + len(orderd) + 2

        # xlsx.save()

    # creating a new data frame fro info2r with a column that contain the cumsum and client name
    info2r_for_noreported = info2r
    info2r_for_noreported["cumsum"] = info2r_for_noreported["cumsum"].astype(int)
    info2r_for_noreported["cumsum_client"] = info2r_for_noreported[
        "clients name"
    ].astype(str) + info2r_for_noreported["cumsum"].astype(str)
    # creating a DF with the shared owners reports but without the client to report
    report_tr_min_ow_noreported = report_tr_min_ow
    report_tr_min_ow_noreported["cumsum_client"] = report_tr_min_ow_noreported[
        "clients name"
    ].astype(str) + report_tr_min_ow_noreported["cumsum"].astype(str)
    report_tr_min_ow_noreported = report_tr_min_ow_noreported[
        -report_tr_min_ow_noreported["cumsum_client"].isin(
            info2r_for_noreported["cumsum_client"]
        )
    ]
    report_tr_min_ow_noreported = report_tr_min_ow_noreported.drop_duplicates(
        subset="cumsum_client"
    )
    # creating a data frame with each client of each shared owner report details
    shared_owners = clients2columns(report_tr_min_ow_noreported, client)
    for col in shared_owners:
        if col.startswith("clients name"):
            shared_owners[col] = shared_owners[col].str.replace("\d+", "")

    for col in shared_owners.columns:
        if col.startswith("id number"):
            shared_owners[col] = shared_owners[col].fillna("")
    # defining the id numbers astype int
    for col in shared_owners.columns:
        if col.startswith("id number"):
            shared_owners[col] = shared_owners[col].apply(
                lambda x: round(x) if isinstance(x, float) else x
            )

    shared_owners = shared_owners.replace({123456789: ""}).fillna("")

    info2r_status = info2r[
        info2r["finaltypedown"].str.contains("p" or "d", regex=False)
    ]
    check4excel = check.merge(
        info2r_status[["cumsum", "finaltypedown"]], on="cumsum", how="left"
    )

    info2r_status_ow = info2r[info2r["finaltypedown"].str.contains("m", regex=False)]
    check4excel_ow = check_ow.merge(
        info2r_status_ow[["cumsum", "finaltypedown"]], on="cumsum", how="left"
    )

    check4excel = check4excel.merge(
        check4excel_ow[["N", "finaltypedown"]], on="N", how="left", suffixes=("", "_ow")
    )

    check4excel = check4excel[
        [
            "N",
            "deal number",
            "date",
            "clients name",
            "check number",
            "bank number",
            "branch number",
            "account number",
            "owner name",
            "nis sum",
            "finaltypedown",
            "finaltypedown_ow",
        ]
    ]

    check4excel = check4excel.merge(
        client[["clients name", "id number"]], how="left", on="clients name"
    )

    check4excel = check4excel.rename(
        columns={
            "deal number": "מזהה פעולה",
            "date": "תאריך פעולה",
            "clients name": "שם הלקוח",
            "check number": "מספר צק",
            "bank number": "מספר בנק",
            "branch number": "מספר סניף",
            "account number": "מספר חשבון",
            "owner name": "שם המושך",
            "nis sum": "סכום בשקלים",
            "finaltypedown": "סטטוס",
            "finaltypedown_ow": "סטטוס מושך",
        }
    )

    check4excel[["סטטוס", "סטטוס מושך"]] = (
        check4excel[["סטטוס", "סטטוס מושך"]]
        .replace(
            {
                "pd": "פיצול ודירדוס",
                "pr": "פיצול בסיכון",
                "p": "פיצול",
                "dr": "דירדוס בסיכון",
                "d": "דירדוס",
                "m": "מושך משותף",
                "mr": "מושך משותף בסיכון",
                "pm": "פיצול ומושך משותף",
                "nr": "דיווח רגיל בסיכון",
                "prmr": "פיצול ומושך משותף בסיכון",
            }
        )
        .fillna("תקין")
    )

    check4excel["last3month"] = np.where(
        check4excel["שם הלקוח"].isin(reported["שם"]),
        "דווח בשלושת החודשים האחרונים",
        "תקין",
    )
    check4excel["fsp"] = np.where(
        check4excel["id number"].isin(fsp_list["מספר מזהה"].astype(str)), 'נש"מ', "תקין"
    )

    for i in check4excel.index:
        if check4excel["fsp"][i] != "תקין":
            for col in check4excel.columns:
                if col.startswith("סטטוס"):
                    check4excel[col][i] = 'נש"מ'

    for i in check4excel.index:
        if check4excel["last3month"][i] != "תקין":
            for col in check4excel.columns:
                if col.startswith("סטטוס"):
                    check4excel[col][i] = "דווח בשלושת החודשים האחרונים"

    check4excel = check4excel.drop(columns=["last3month", "fsp"])
    check4excel = check4excel.drop_duplicates("N")
    # check4excel.to_excel(xlsx, sheet_name='סטטוס', index=False)
    with xlsx_test as writer:
        check4excel.to_excel(writer, sheet_name="סטטוס", index=False)

    try:
        info2r["last name"] = info2r["person name"].str.split(" ", 1, expand=True)[1]
        info2r["first name"] = info2r["person name"].str.split(" ", 1, expand=True)[0]
        info2r = info2r.fillna("")
    except KeyError:
        info2r["last name"] = ""
        info2r["first name"] = ""

    filepath_check = template_directory + "\*template*"
    txt_check = glob.glob(filepath_check)
    for textfile in txt_check:
        template = textfile
    document = docx.Document(template)
    image_paras = [
        i for i, p in enumerate(document.paragraphs) if "[ChartImage1]" in p.text
    ]
    p = document.paragraphs[image_paras[0]]
    p.text = ""
    r = p.add_run()
    r.add_text("חתימה: ").bold = True
    # try:
    #     # r.add_picture(clients_folder_path + "\signature.png")
    #     document.save(template_directory + "\my_doc.docx")
    # except:
    #     document.save(template_directory + "\my_doc.docx")

    filepath_check = template_directory + "\*my_doc*"
    txt_check = glob.glob(filepath_check)
    for textfile in txt_check:
        template = textfile
    document = MailMerge(template)
    print(document.get_merge_fields())

    # getting ready for the merge
    if "person birth" not in info2r:
        info2r["person birth"] = ""
    else:
        info2r["person birth"] = info2r["person birth"]

    if "company birth" not in info2r:
        info2r["company birth"] = ""
    else:
        info2r["company birth"] = info2r["company birth"]

    if "address" not in info2r:
        info2r["address"] = ""
    else:
        info2r["address"] = info2r["address"]

    if "city" not in info2r:
        info2r["city"] = ""
    else:
        info2r["city"] = info2r["city"]

    if report_tr_min_ow.empty:
        error6 = st.error("אין דיווחי מושך משותף")

    else:
        info2r_m = info2r[info2r["finaltypedown"].str.contains("m")].merge(
            report_tr_min_ow[["cumsum", "owner name"]], on="cumsum", how="left"
        )
        info2r_m = info2r_m.merge(
            shared_owners, on="cumsum", how="left", suffixes=("", "clients name_0")
        )
        info2r_m = info2r_m.drop_duplicates("clients name")
        info2r = info2r_m.append(info2r[-info2r["finaltypedown"].str.contains("m")])

    columns_list = [
        "clients name_",
        "id numberclients name_",
        "cityclients name_",
        "addressclients name_",
    ]

    # Create an empty list to store the new values
    new_columns_list = []

    # Iterate over the values in the original list
    for column in columns_list:
        # Iterate over the numbers from 0 to 8
        for i in range(9):
            # Append the original value with the current number
            new_columns_list.append(column + str(i))

    if "owner name" not in info2r:
        info2r["owner name"] = ""
    else:
        info2r["owner name"] = info2r["owner name"]

    for i in new_columns_list:
        if i not in info2r:
            info2r[i] = ""
        else:
            info2r[i] = info2r[i]

    info2r = info2r.fillna("")
    info2r = info2r.reset_index()
    doc_df = pd.DataFrame(columns=["file_name", "bytes"])
    for i in info2r.index:
        # template1 = 'test1.docx'
        document = MailMerge(template)
        document.merge(
            first_name=str(info2r["first name"][i]),
            last_name=str(info2r["last name"][i]),
            company_name=str(info2r["company name"][i]),
            person_birth=str(info2r["person birth"][i]),
            company_birth=str(info2r["company birth"][i]),
            company_id=str(info2r["company id"][i]),
            Title=str(info2r["Title"][i]),
            person_citizenship=str(info2r["person citizenship"][i]),
            Content=str(info2r["Content"][i]),
            country=str(info2r["country"][i]),
            person_id=str(info2r["person id"][i]),
            report_id=str(info2r["report_id"][i]),
            city=str(info2r["city"][i]),
            address=str(info2r["address"][i]),
            business_name=str(business_name),
            business_id=str(business_id),
            branch_number=str(branch_number),
            business_adress=str(business_address),
            business_phone=str(business_phone),
            worker_name=str(worker_name),
            date=str(date),
            worker_id=str(worker_id),
            workers_phone=str(worker_phone),
            workers_email=str(worker_email),
            format=str("של ניכיון צ'קים"),
            worker_position=str(worker_position),
            business_type=str(business_type),
            clients_name_0=str(info2r["clients name_0"][i]),
            clients_name_1=str(info2r["clients name_1"][i]),
            clients_name_2=str(info2r["clients name_2"][i]),
            clients_name_3=str(info2r["clients name_3"][i]),
            clients_name_4=str(info2r["clients name_4"][i]),
            clients_name_5=str(info2r["clients name_5"][i]),
            clients_name_6=str(info2r["clients name_6"][i]),
            clients_name_7=str(info2r["clients name_7"][i]),
            clients_name_8=str(info2r["clients name_8"][i]),
            owner_name=str(info2r["owner name"][i]),
            id_numberclients_name_0=str(info2r["id numberclients name_0"][i]),
            id_numberclients_name_1=str(info2r["id numberclients name_1"][i]),
            id_numberclients_name_2=str(info2r["id numberclients name_2"][i]),
            id_numberclients_name_3=str(info2r["id numberclients name_3"][i]),
            id_numberclients_name_4=str(info2r["id numberclients name_4"][i]),
            id_numberclients_name_5=str(info2r["id numberclients name_5"][i]),
            id_numberclients_name_6=str(info2r["id numberclients name_6"][i]),
            id_numberclients_name_7=str(info2r["id numberclients name_7"][i]),
            id_numberclients_name_8=str(info2r["id numberclients name_8"][i]),
            cityclients_name_0=str(info2r["cityclients name_0"][i]),
            cityclients_name_1=str(info2r["cityclients name_1"][i]),
            cityclients_name_2=str(info2r["cityclients name_2"][i]),
            cityclients_name_3=str(info2r["cityclients name_3"][i]),
            cityclients_name_4=str(info2r["cityclients name_4"][i]),
            cityclients_name_5=str(info2r["cityclients name_5"][i]),
            cityclients_name_6=str(info2r["cityclients name_6"][i]),
            cityclients_name_7=str(info2r["cityclients name_7"][i]),
            cityclients_name_8=str(info2r["cityclients name_8"][i]),
            addressclients_name_0=str(info2r["addressclients name_0"][i]),
            addressclients_name_1=str(info2r["addressclients name_1"][i]),
            addressclients_name_2=str(info2r["addressclients name_2"][i]),
            addressclients_name_3=str(info2r["addressclients name_3"][i]),
            addressclients_name_4=str(info2r["addressclients name_4"][i]),
            addressclients_name_5=str(info2r["addressclients name_5"][i]),
            addressclients_name_6=str(info2r["addressclients name_6"][i]),
            addressclients_name_7=str(info2r["addressclients name_7"][i]),
            addressclients_name_8=str(info2r["addressclients name_8"][i]),
        )

        output = (
            str(info2r["report_name"][i]) + "-" + str(info2r["report_id"][i]) + ".docx"
        )

        # document.write(output)

        # this part is new and tested todo
        # Set file metadata and parent folder ID
        # file_metadata = {'name': output, 'parents': ['1AFo4UxjY0YQQrk7OstHAZqI2jad6d04u']}
        doc_byte = BytesIO()
        document.write(doc_byte)
        doc4df = pd.DataFrame(
            {
                "file_name": [
                    str(info2r["report_name"][i])
                    + "-"
                    + str(info2r["report_id"][i])
                    + ".docx"
                ],
                "bytes": [doc_byte],
            }
        )
        doc_df = doc_df.append(doc4df)
        # media = MediaIoBaseUpload(doc_byte,
        #                           mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        # # # # Upload the file to Google Drive
        # drive_service = get_gdrive_service()
        # # #
        # drive_service.files().create(body=file_metadata, media_body=media, fields='id').execute()

    info2r = info2r.set_index("cumsum")
    info2r_no_m = info2r[
        (info2r["finaltypedown"].str.contains("p", regex=False))
        | (info2r["finaltypedown"].str.contains("d", regex=False))
    ]

    info2r_nona = info2r_no_m[info2r_no_m.index.notnull()]

    final_reports4word = mont2report_w_cumsum.merge(
        info2r_nona[["clients name", "report_name", "report_id"]],
        on="clients name",
        how="left",
    )
    final_reports4word = final_reports4word.drop_duplicates("cumsum")
    final_reports4word = final_reports4word.set_index("cumsum")
    doc_df = doc_df.set_index("file_name")
    for title in ls_cumsum:
        data = check[check["cumsum"] == title].round()
        table = [["סהכ", data["nis sum"].sum().round(0)]]
        last_row = pd.DataFrame(table, columns=["owner name", "nis sum"])
        orderd = pd.concat([data, last_row])
        orderd = orderd[
            [
                "deal number",
                "date",
                "clients name",
                "check number",
                "bank number",
                "account number",
                "branch number",
                "owner name",
                "nis sum",
            ]
        ].fillna(" ")

        orderd = change_names_and_order(
            yeshut_directory + r"/change_names.xlsx", orderd
        )

        orderd["תאריך פעולה"] = orderd["תאריך פעולה"].astype(str)
        orderd["סכום"] = orderd["סכום"].astype(str)

        # doc = docx.Document(
        #     str(final_reports4word['report_name'].loc[title]) + '-' + str(
        #         final_reports4word['report_id'].loc[title]) + '.docx')

        file_name = (
            str(final_reports4word["report_name"].loc[title])
            + "-"
            + str(final_reports4word["report_id"].loc[title])
            + ".docx"
        )
        doc = docx.Document(doc_df.loc[file_name].values[0])

        doc.add_page_break()
        t = doc.add_table(orderd.shape[0] + 1, orderd.shape[1])

        for j in range(orderd.shape[-1]):
            t.cell(0, j).text = orderd.columns[j]

            # add the rest of the data frame
        for i in range(orderd.shape[0]):
            for j in range(orderd.shape[-1]):
                t.cell(i + 1, j).text = str(orderd.values[i, j])

        t.style = "Grid Table 4 Accent 5"

        # modified_file_content = io.BytesIO()
        # doc.save(modified_file_content)
        doc.save(doc_df.loc[file_name].values[0])

        # Upload the modified docx file to Google Drive
        # media = MediaIoBaseUpload(modified_file_content,
        #                   mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        # body = {'name': file_name}
        # drive_service.files().update(fileId=file_id, body=body, media_body=media, fields='id').execute()

        # doc.save(str(final_reports4word['report_name'].loc[title]) + "-" + str(
        #     final_reports4word['report_id'].loc[title]) + ".docx")

    # todo merge in two stages one by cumsum one by clients name
    info2r = info2r.reset_index()

    mont_ow_clients_w_cumsum["clients name"] = mont_ow_clients_w_cumsum[
        "clients name"
    ].str.replace("\d+", "")

    info2r_only_m = info2r[info2r["finaltypedown"].str.contains("m", regex=False)]
    info2r_ow = info2r_only_m[
        ["report_name", "report_id", "clients name", "owner name"]
    ].merge(
        mont_ow_clients_w_cumsum[["clients name", "cumsum"]],
        on="clients name",
        how="left",
    )
    info2r_ow = info2r_ow.drop_duplicates("cumsum")
    info2r_ow = info2r_ow.set_index("cumsum")

    # final_reports4word_ow = mont_ow_clients_w_cumsum.merge(info2r_ow[['owner name','report_name','report_id']],
    #                                           on = 'owner name',how = 'left')
    # final_reports4word_ow = final_reports4word.set_index('cumsum')

    for title in ls_cumsum_ow:
        data = report_tr_min_ow[report_tr_min_ow["cumsum"] == title]
        table = [["סהכ", data["nis sum"].sum().round(0)]]
        last_row = pd.DataFrame(table, columns=["owner name", "nis sum"])
        orderd = pd.concat([data, last_row])
        orderd = orderd[
            [
                "deal number",
                "date",
                "clients name",
                "check number",
                "bank number",
                "account number",
                "branch number",
                "owner name",
                "nis sum",
            ]
        ].fillna(" ")
        orderd = change_names_and_order(
            yeshut_directory + r"/change_names.xlsx", orderd
        )

        orderd["תאריך פעולה"] = orderd["תאריך פעולה"].astype(str)
        orderd["סכום"] = orderd["סכום"].astype(str)

        file_name = (
            str(info2r_ow["report_name"].loc[title])
            + "-"
            + str(info2r_ow["report_id"].loc[title])
            + ".docx"
        )

        doc = docx.Document(doc_df.loc[file_name].values[0])
        # doc = docx.Document(str(info2r_ow['report_name'].loc[title]) + "-" + str(
        #     info2r_ow['report_id'].loc[title]) + ".docx")

        doc.add_page_break()
        t = doc.add_table(orderd.shape[0] + 1, orderd.shape[1])

        for j in range(orderd.shape[-1]):
            t.cell(0, j).text = orderd.columns[j]

            # add the rest of the data frame
        for i in range(orderd.shape[0]):
            for j in range(orderd.shape[-1]):
                t.cell(i + 1, j).text = str(orderd.values[i, j])

        t.style = "Grid Table 4 Accent 5"

        # doc.save(str(info2r_ow['report_name'].loc[title]) + "-" + str(
        #     info2r_ow['report_id'].loc[title]) + ".docx")
        doc.save(doc_df.loc[file_name].values[0])

    # # Set file metadata and parent folder ID
    # file_metadata = {'name': 'data.xlsx', 'parents': ['1AFo4UxjY0YQQrk7OstHAZqI2jad6d04u']}
    # media = MediaIoBaseUpload(xlsx_data,
    #                           mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    # # # # Upload the file to Google Drive
    # drive_service = get_gdrive_service()
    # # #
    # drive_service.files().create(body=file_metadata, media_body=media, fields='id').execute()
    #
    # for i in doc_df.index:
    #     file_metadata = {'name': i, 'parents': ['1AFo4UxjY0YQQrk7OstHAZqI2jad6d04u']}
    #     media = MediaIoBaseUpload(doc_df['bytes'][i],
    #                               mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    #     # #
    #     drive_service.files().create(body=file_metadata, media_body=media, fields='id').execute()
    # Create an in-memory zip file
    doc_df = doc_df.append(
        pd.DataFrame(
            {"file_name": ["monitoring.xlsx"], "bytes": [xlsx_data]}
        ).set_index("file_name")
    )
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
        for file_name, file_content in doc_df.iterrows():
            # Add each file to the zip file
            file_content = file_content[0]  # Access the BytesIO object
            file_content.seek(0)  # Reset the file pointer to the beginning
            zf.writestr(file_name, file_content.read())

    # Reset the buffer's file pointer to the beginning
    zip_buffer.seek(0)

    download = st.download_button(
        label="📥 Download Reports Zip", data=zip_buffer, file_name="Reports.zip"
    )
    return st.error("הניטור הושלם בהצלחה")


def run_gmt(check, business_info, reported):
    directories = pd.read_excel("directories.xlsx")
    directories = pd.DataFrame(directories)
    directories = directories.set_index("File type")
    rc_directory = directories.loc["risk countries"].values[0]
    rates_directory = directories.loc["rate changer"].values[0]

    template_directory = directories.loc["word template"].values[0]

    gmt_directory = directories.loc["gmt content and change names"].values[0]

    # reading the risk country file
    filepath_check = rc_directory + "\*מדינות*"
    txt_check = glob.glob(filepath_check)
    for textfile in txt_check:
        rc = pd.read_excel(textfile)

    # reading the rates file
    filepath_check = rates_directory + "\*שערים*"
    txt_check = glob.glob(filepath_check)
    for textfile in txt_check:
        rates = pd.read_excel(
            textfile,
            sheet_name="CUR CONV",
            usecols="F:G",
            skiprows=1,
            engine="openpyxl",
        )

    # reading the report content file
    filepath_check = gmt_directory + "\*content*"
    txt_check = glob.glob(filepath_check)
    for textfile in txt_check:
        content = pd.read_excel(textfile)

    # changing the columns name to standard
    gmt = check.rename(
        columns={
            "מדינת יעד/מדינת מקור": "dest country",
            "Country": "dest country",
            "שמ המוטב/שם השולח": "clients name",
            "Beneficiary Name/Sender Name": "clients name",
            "Second side of the transaction": "owner name",
            "הצד השני של העסקה": "owner name",
            "מטבע": "currency",
            "Currency": "currency",
            "Transaction number": "deal number",
            "Amount": "amount",
            "סכום": "amount",
            "תאריך": "date",
            "Date": "date",
            "Date": "date",
            "Transaction type": "סוג עסקה",
            "מספר עסקה": "deal number",
            "מספר זיהוי של לקוח/ה": "id number",
            "Customer ID": "id number",
        }
    )

    try:
        gmt = gmt[gmt["מבוטלת"] != "Canceled"]
    except KeyError:
        gmt = gmt[gmt["Canceled"] != "Canceled"]

    gmt["clients name"] = gmt["clients name"].str.strip()
    gmt["owner name"] = gmt["owner name"].str.strip()
    #
    st.error("הקבצים נקראו בהצלחה")

    gmt = gmt.dropna(subset="amount")
    gmt = gmt.dropna(subset="currency")

    # naming the currency rates
    rates.columns = ["curr", "rate"]

    # calculating the sum in ils
    gmt = gmt.merge(rates, left_on="currency", right_on="curr", how="left")
    gmt["nis sum"] = gmt["rate"] * gmt["amount"]

    # rounding the amount and nis sum:
    gmt = gmt.dropna(subset="amount")
    gmt = gmt.dropna(subset="nis sum")
    gmt["nis sum"] = gmt["nis sum"].astype(int)
    gmt["amount"] = gmt["amount"].astype(int)

    risk_countries = rc[-rc["סיכון/ללא דיווח רגיל"].isna()]
    # find risk transactions gmt
    gmt["dest country"] = gmt["dest country"].str.strip()
    gmt["eng_name"] = gmt["dest country"].str.split(",", 1, expand=True)[0]
    risk_type1_gmt = gmt.merge(
        rc[["English short name (using title case)", "סיכון/ללא דיווח רגיל"]],
        left_on="eng_name",
        right_on="English short name (using title case)",
        how="left",
    )
    risk_type1_gmt = risk_type1_gmt[risk_type1_gmt["סיכון/ללא דיווח רגיל"].notna()]
    risk_type1_gmt = risk_type1_gmt.dropna(subset="dest country")
    # todo check if it works
    # gmt['deal number'] = gmt['deal number'].fillna(gmt['clients name'].str.split(', ', 1, expand = True)[1])
    try:
        gmt["PH"] = (
            gmt["clients name"]
            .str.split(", ", 1, expand=True)[1]
            .str.split("-", 1, expand=True)[0]
            .fillna("")
        )

        risk_type2_gmt = gmt[gmt["PH"].isin(risk_countries["Alpha-2 code"])]
    except KeyError:
        risk_type2_gmt = risk_type1_gmt

    gmt["PH_id"] = gmt["id number"].str.split("-", 1, expand=True)[0].fillna("")
    risk_type3_gmt = gmt[gmt["PH_id"].isin(risk_countries["Alpha-2 code"])]

    all_risk_gmt = pd.concat(
        [risk_type1_gmt, risk_type2_gmt, risk_type3_gmt]
    ).drop_duplicates("deal number")

    # building three kinds of monitoring data frames
    monitoring_name = gmt.groupby("clients name").agg(
        {"nis sum": "sum", "deal number": "nunique"}
    )
    monitoring_ow = gmt.groupby("owner name").agg(
        {"nis sum": "sum", "clients name": "nunique"}
    )

    # filtering need to report transactions:
    monitoring_name_2report = monitoring_name[
        ((monitoring_name["nis sum"] >= 50000) & (monitoring_name["deal number"] > 1))
        | ((monitoring_name["nis sum"] >= 47000) & (monitoring_name["nis sum"] < 50000))
    ]

    # adding the report type:# todo managing the content writen for the different errors
    monitoring_name_2report["report type"] = np.where(
        (monitoring_name_2report["nis sum"] >= 50000), "p", "d"
    )

    try:
        monitoring_name_2report = monitoring_name_2report[
            -monitoring_name_2report["clients name"].isin(reported["שם"])
        ]
    except KeyError:
        st.error("אין דיווחי פיצול או דירדוס לא בסיכון")

    # building three kinds of monitoring data frames
    monitoring_name_risk = all_risk_gmt.groupby("clients name").agg(
        {"nis sum": "sum", "deal number": "nunique"}
    )
    monitoring_ow_risk = all_risk_gmt.groupby("owner name").agg(
        {"nis sum": "sum", "clients name": "nunique"}
    )

    # filtering need to report transctions:
    monitoring_name_risk_2report = monitoring_name_risk[
        (
            (monitoring_name_risk["nis sum"] >= 5000)
            & (monitoring_name_risk["deal number"] > 1)
        )
        | (
            (monitoring_name_risk["nis sum"] >= 4700)
            & (monitoring_name_risk["nis sum"] < 5000)
        )
    ]

    # adding the report type:
    monitoring_name_risk_2report["report type_risk"] = np.where(
        (monitoring_name_risk_2report["nis sum"] >= 5000), "pr", "dr"
    )
    # adding normal reports as 'nr' to monitoring_name_risk_2report df
    monitoring_risk_normal = monitoring_name_risk[
        (
            (monitoring_name_risk["nis sum"] >= 5000)
            & (monitoring_name_risk["deal number"] == 1)
        )
    ]
    monitoring_risk_normal["report type_risk"] = "nr"
    monitoring_name_risk_2report = monitoring_name_risk_2report.append(
        monitoring_risk_normal
    )

    # concatenating the client to report df's
    all_reports = monitoring_name_2report.append(monitoring_name_risk_2report)
    all_reports = all_reports.reset_index()
    all_reports = all_reports[-all_reports["clients name"].isin(reported["שם"])]
    all_reports = all_reports.reset_index()

    # shared destination monitoring
    monitoring_ow_2report = monitoring_ow[
        (monitoring_ow["nis sum"] >= 50000) & (monitoring_ow["clients name"] > 1)
    ]
    monitoring_ow_2report = monitoring_ow_2report.rename(
        columns={"clients name": "clients count"}
    )

    # checking if a client is already reported in previous test
    monitoring_ow_2report_clients = monitoring_ow_2report.merge(
        gmt[["clients name", "owner name"]], on="owner name", how="left"
    )
    monitoring_ow_2report_clients = monitoring_ow_2report_clients.merge(
        all_reports["clients name"], on="clients name", how="left", suffixes=["", "_pd"]
    )
    try:
        monitoring_ow_2report_clients = monitoring_ow_2report_clients[
            monitoring_ow_2report_clients["clients name_pd"].isna()
        ]
    except KeyError:
        st.error("אין לקוחות שביצעו גם עבירת מושך משותף וגם עבירת פיצול")
    try:
        monitoring_ow_2report_clients = monitoring_ow_2report_clients[
            -monitoring_ow_2report_clients["clients name"].isin(reported["שם"])
        ]
    except:
        st.error("אין לקוחות שביצעו עברת מושך משותף ודווחו בשלושת החודשים האחרונים")

    # appending the shared owner clients to the all reports DF
    monitoring_ow_2report_clients["report type ow"] = "m"
    monitoring_ow_2report_clients = monitoring_ow_2report_clients.drop_duplicates(
        "owner name"
    )
    all_reports = all_reports.append(monitoring_ow_2report_clients)

    # shared destination monitoring
    monitoring_ow_2report_risk = monitoring_ow_risk[
        (monitoring_ow_risk["nis sum"] >= 5000)
        & (monitoring_ow_risk["clients name"] > 1)
    ]
    monitoring_ow_2report_risk = monitoring_ow_2report_risk.rename(
        columns={"clients name": "clients count"}
    )

    # checking if a client is already reported in previous tests
    monitoring_ow_2report_risk = monitoring_ow_2report_risk.merge(
        all_risk_gmt[["clients name", "owner name"]], on="owner name", how="left"
    )
    monitoring_ow_2report_risk_clients = monitoring_ow_2report_risk.merge(
        all_reports["clients name"], on="clients name", how="left", suffixes=("", "_pd")
    )
    try:
        monitoring_ow_2report_risk_clients = monitoring_ow_2report_risk_clients[
            monitoring_ow_2report_risk_clients["clients name_pd"].isna()
        ]
    except KeyError:
        st.error("אין לקוחות שביצעו דיווח מושך משותף בסיכון ופיצול בסיכון")
    try:
        monitoring_ow_2report_risk_clients = monitoring_ow_2report_risk_clients[
            -monitoring_ow_2report_risk_clients["clients name"].isin(reported["שם"])
        ]
    except:
        st.error(
            "אין לקוחות בסיכון שביצעו עברת מושך משותף ודווחו בשלושת החודשים האחרונים"
        )

    # appending the risk hared owner clients to the all reports DF
    monitoring_ow_2report_risk_clients["report type ow risk"] = "mr"
    monitoring_ow_2report_risk_clients = (
        monitoring_ow_2report_risk_clients.drop_duplicates("owner name")
    )
    all_reports = all_reports.append(monitoring_ow_2report_risk_clients)

    # joining all the report types into one list
    all_reports = all_reports.fillna("")
    all_reports["final type"] = (
        all_reports["report type"]
        + all_reports["report type_risk"]
        + all_reports["report type ow"]
        + all_reports["report type ow risk"]
    )

    # summarizing the data frame for mailing
    all_reports4word = all_reports[["clients name", "final type", "owner name"]]

    # adding the country of dest details
    all_reports4word = all_reports4word.merge(
        gmt[["clients name", "id number", "dest country"]],
        on="clients name",
        how="left",
    )
    all_reports4word = all_reports4word.drop_duplicates("clients name")

    if all_reports4word.empty == False:
        all_reports4word["name4report"] = all_reports4word["clients name"].str.split(
            ",", 1, expand=True
        )[0]
        try:
            all_reports4word["dest country heb"] = all_reports4word[
                "dest country"
            ].str.split(", ", 1, expand=True)[1]
        except KeyError:
            all_reports4word["dest country heb"] = all_reports4word["dest country"]
        try:
            all_reports4word["id number new"] = all_reports4word[
                "clients name"
            ].str.split(",", 1, expand=True)[1]
            all_reports4word["final id number"] = all_reports4word[
                "id number new"
            ].fillna(all_reports4word["id number"])
        except KeyError:
            all_reports4word["final id number"] = all_reports4word["id number"]

        all_reports4word["country_code"] = (
            all_reports4word["final id number"]
            .str.split("-", 1, expand=True)[0]
            .str.lstrip()
        )
        all_reports4word = all_reports4word.merge(
            rc[["Alpha-2 code", "עברית"]],
            left_on="country_code",
            right_on="Alpha-2 code",
            how="left",
        )
        all_reports4word = all_reports4word.drop_duplicates("clients name")
        try:
            all_reports4word["final risk country"] = all_reports4word[
                "dest country heb"
            ].fillna(all_reports4word["עברית"])
        except:
            all_reports4word["final risk country"] = all_reports4word[
                "dest country heb"
            ]

        all_reports4word_sum = all_reports4word[
            [
                "clients name",
                "final type",
                "final id number",
                "name4report",
                "עברית",
                "dest country heb",
                "final risk country",
            ]
        ]
    else:
        gmt["final type"] = "תקין"
        gmt["final type_ow"] = "תקין"
        gmt4excel = gmt[
            [
                "deal number",
                "date",
                "clients name",
                "סוג עסקה",
                "amount",
                "currency",
                "dest country",
                "id number",
                "owner name",
                "nis sum",
                "final type",
                "final type_ow",
            ]
        ]

        gmt4excel = gmt4excel.rename(
            columns={
                "deal number": "מזהה פעולה",
                "date": "תאריך פעולה",
                "clients name": "שם הלקוח",
                "amount": "סכום במטבע",
                "currency": "שם המטבע",
                "id number": "מספר תעודה מזהה לקוח",
                "dest country": "יעד",
                "owner name": "שם המושך",
                "nis sum": "סכום בשקלים",
                "final type": "סטטוס",
                "final type_ow": "סטטוס מושך",
            }
        )

        gmt4excel["סטטוס סיכון"] = np.where(
            gmt4excel["מזהה פעולה"].isin(all_risk_gmt["deal number"]), "סיכון", "תקין"
        )

        # gmt4excel.to_excel(transaction_folder_path + "\monitoring_GMT.xlsx") todo fix the no report part
        xlsx_data = BytesIO()
        with pd.ExcelWriter(xlsx_data, engine="openpyxl") as writer:
            gmt4excel.to_excel(writer, sheet_name="ניהול התראות")
        xlsx_data.seek(0)
        st.download_button(
            label="📥 Download Final Report",
            data=xlsx_data,
            file_name=str(business_info.loc["שם מלא של הגורם המדווח"].values[0])
            + ".xlsx",
        )
        st.error("אין דיווחים בלתי רגילים")
        st.error("הניטור הושלם בהצלחה")
        exit()

    # merging the report content
    all_reports4word_sum = all_reports4word_sum.merge(
        content, left_on="final type", right_on="type", how="left"
    )

    try:
        all_reports4word_sum["risk_full_content"] = np.where(
            all_reports4word_sum["final type"].str.contains("r", regex=False),
            "הכספים נשלחו ל"
            + all_reports4word_sum["final risk country"]
            + " ,מדינה המוגדרת בסיכון גבוה.",
            "",
        )
        all_reports4word_sum["content2report"] = (
            all_reports4word_sum["Content"]
            + " "
            + all_reports4word_sum["risk_full_content"]
        )
    except KeyError:
        st.error(message="אין דיווחים בלתי רגילים בסיכון")

    # fixing the business info
    business_info = business_info.fillna("")
    business_name = business_info.loc["שם מלא של הגורם המדווח"].values[0]
    business_id = business_info.loc["מספר מזהה של הגורם המדווח"].values[0]
    branch_number = business_info.loc["מספר סניף מדווח"].values[0]
    business_address = business_info.loc["מען הסניף"].values[0]
    date = to_date.today()
    business_phone = business_info.loc["טלפון ופקסימילה של הגורם המדווח"].values[0]
    worker_name = business_info.loc["שם פרטי ושם משפחה של עובד עורך הדיווח"].values[0]
    worker_id = business_info.loc['ת"ז של עורך הדיווח'].values[0]
    worker_position = business_info.loc["תפקיד עורך הדיווח"].values[0]
    worker_phone = business_info.loc["טלפון עורך הדיווח"].values[0]
    worker_email = business_info.loc["דואר אלקטרוני"].values[0]

    # importing the report number needed
    all_reports4word_sum["report_id_aid"] = reported["מס דיווח"].max() + 1
    all_reports4word_sum["report_id"] = (
        all_reports4word_sum["report_id_aid"] + range(len(all_reports4word_sum.index))
    ).astype(int)
    all_reports4word_sum["report_name"] = report_num_gen(business_id, branch_number)
    all_reports4word_sum["first name"] = all_reports4word_sum["name4report"].str.split(
        " ", 1, expand=True
    )[0]
    all_reports4word_sum["last name"] = all_reports4word_sum["name4report"].str.split(
        " ", 1, expand=True
    )[1]

    all_reports4word_sum = all_reports4word_sum.fillna("")
    all_reports4word_sum = all_reports4word_sum.reset_index()
    all_reports4word_excel = all_reports4word_sum[
        [
            "clients name",
            "final type",
            "final id number",
            "עברית",
            "report_id",
            "Title",
            "content2report",
        ]
    ]
    all_reports4word_excel = all_reports4word_excel.rename(
        columns={
            "clients name": "שם הלקוח",
            "final type": "מהות הפעילות",
            "final id number": "מספר תעודה מזהה",
            "עברית": "מוצא",
            "report_id": "מספר דיווח",
            "title": "תמצית",
            "content2report": "תוכן",
        }
    )

    all_reports4word_excel["לדווח/לא לדווח"] = ""
    all_reports4word_excel["הערות"] = ""
    all_reports4word_excel["מהות הפעילות"] = all_reports4word_excel[
        "מהות הפעילות"
    ].replace(
        {
            "pd": "פיצול ודירדוס",
            "pr": "פיצול בסיכון",
            "p": "פיצול",
            "dr": "דירדוס בסיכון",
            "d": "דירדוס",
            "m": "מושך משותף",
            "mr": "מושך משותף בסיכון",
            "pm": "פיצול ומושך משותף",
            "dm": "דירדוס ומושך משותף",
            "nr": "דיווח רגיל בסיכון",
        }
    )

    # all_reports4word_excel.to_excel(transaction_folder_path + "\monitoring_GMT.xlsx", sheet_name='ניהול התראות')
    xlsx_data = BytesIO()
    with pd.ExcelWriter(xlsx_data, engine="openpyxl") as writer:
        all_reports4word_excel.to_excel(writer, sheet_name="ניהול התראות")
    xlsx_data.seek(0)

    # xlsx = pd.ExcelWriter(transaction_folder_path + "\monitoring_GMT.xlsx", engine='openpyxl', mode='a',
    #                       if_sheet_exists='overlay')
    xlsx_test = pd.ExcelWriter(
        xlsx_data, engine="openpyxl", mode="a", if_sheet_exists="overlay"
    )

    # filtering outbound reports
    sent_report = all_reports4word_sum[
        (all_reports4word_sum["final type"] == "p")
        | (all_reports4word_sum["final type"] == "d")
        | (all_reports4word_sum["final type"] == "dr")
        | (all_reports4word_sum["final type"] == "pr")
        | (all_reports4word_sum["final type"] == "nr")
        | (all_reports4word_sum["final type"] == "ppr")
    ]

    # creating a data frames with only need to report transctions
    sent_report_tr = gmt[gmt["clients name"].isin(sent_report["clients name"])]

    # leaving only the important columns
    sent_report_tr = sent_report_tr[
        [
            "deal number",
            "סוג עסקה",
            "date",
            "dest country",
            "clients name",
            "owner name",
            "id number",
            "amount",
            "currency",
            "nis sum",
        ]
    ]

    # creating an Excel sheet with the outbound reports
    i = 0
    for_word_table = []

    # round the sum in ils
    sent_report_tr["nis sum"] = sent_report_tr["nis sum"].round()

    # creating a unique list of the sender names
    ls_customer = sent_report_tr["clients name"].unique().tolist()

    # loop each sender transactions in a unique table
    for name in ls_customer:
        df = sent_report_tr[sent_report_tr["clients name"] == name]
        data = [["סהכ", df["amount"].sum(), df["nis sum"].sum()]]
        last_row = pd.DataFrame(data, columns=["clients name", "amount", "nis sum"])
        orderd = pd.concat([df, last_row])
        with xlsx_test as writer:
            orderd.to_excel(writer, sheet_name="פירוט התראות", startrow=i, index=False)

        for_word_table.append(orderd)
        i = i + len(orderd) + 2

        # xlsx.save()

    # filtering only shared owner reports
    sent_report_ow = all_reports4word_sum[
        (all_reports4word_sum["final type"] == "m")
        | (all_reports4word_sum["final type"] == "mr")
    ]
    sent_report_ow = sent_report_ow.merge(
        gmt[["clients name", "owner name"]], on="clients name", how="left"
    )

    if sent_report_ow.empty == False:
        # concatenating the two kinds of outbound shared owner transaction dataframes
        sent_report_tr_ow = gmt[gmt["owner name"].isin(sent_report_ow["owner name"])]

        # leaving only the important columns
        sent_report_tr_ow = sent_report_tr_ow[
            [
                "deal number",
                "סוג עסקה",
                "date",
                "dest country",
                "clients name",
                "owner name",
                "id number",
                "amount",
                "currency",
                "nis sum",
            ]
        ]

        # creating a new list
        j = 0
        for_word_table_ow = []

        # round the sum in ils
        sent_report_tr_ow["nis sum"] = sent_report_tr_ow["nis sum"].round()

        # creating a unique list of the sender names
        ls_customer_ow = sent_report_tr_ow["owner name"].unique().tolist()

        # loop each sender transactions in a unique table
        for name in ls_customer_ow:
            df = sent_report_tr_ow[sent_report_tr_ow["owner name"] == name]
            data = [["סהכ", df["nis sum"].sum(), df["amount"].sum()]]
            last_row = pd.DataFrame(data, columns=["clients name", "nis sum", "amount"])
            orderd = pd.concat([df, last_row])
            with xlsx_test as writer:
                orderd.to_excel(
                    writer, sheet_name="פירוט התראות יעד משותף", startrow=j, index=False
                )

            for_word_table_ow.append(orderd)
            j = j + len(orderd) + 2

            # xlsx.save()

        # creating a new data frame from info2r with a column that contain the cumsum and client name
        info2r_for_noreported = sent_report_ow
        # creating a DF with the shared owners reports but without the client to report
        report_tr_min_ow_noreported = sent_report_tr_ow
        info2r_for_noreported = info2r_for_noreported[
            ["report_name", "report_id", "clients name"]
        ].merge(
            sent_report_tr_ow[["clients name", "owner name"]],
            on="clients name",
            how="left",
        )
        info2r_for_noreported = info2r_for_noreported.drop_duplicates("owner name")
        report_tr_min_ow_noreported = report_tr_min_ow_noreported[
            -report_tr_min_ow_noreported["clients name"].isin(
                info2r_for_noreported["clients name"]
            )
        ]
        report_tr_min_ow_noreported = report_tr_min_ow_noreported.drop_duplicates(
            subset="clients name"
        )
        report_tr_min_ow_noreported = report_tr_min_ow_noreported.drop_duplicates(
            subset="clients name"
        )
        report_tr_min_ow_noreported = report_tr_min_ow_noreported.drop_duplicates(
            subset="deal number"
        )
        # creating a data frame with each client of each shared owner report details
        shared_owners = clients2columns_gmt(report_tr_min_ow_noreported)
        for col in shared_owners:
            if col.startswith("clients name"):
                shared_owners[col] = shared_owners[col].str.replace("\d+", "")

        # defining the id numbers astype int todo this part is probably useless for gmt
        #    for col in shared_owners.columns:
        #        if col.startswith('id number'):
        #            shared_owners[col] = shared_owners[col].apply(lambda x: round(x) if isinstance(x, float) else x)

        shared_owners = shared_owners.replace({123456789: ""}).fillna("")

        info2r_m = all_reports4word_sum[
            all_reports4word_sum["final type"].str.contains("m")
        ].merge(
            sent_report_tr_ow[["owner name", "clients name"]],
            on="clients name",
            how="left",
        )
        info2r_m = info2r_m.merge(
            shared_owners, on="owner name", how="left", suffixes=("", "clients name_0")
        )
        info2r_m = info2r_m.drop_duplicates("clients name")
        all_reports4word_sum = info2r_m.append(
            all_reports4word_sum[-all_reports4word_sum["final type"].str.contains("m")]
        )

    else:
        st.error("אין דיווחי מושך משותף")

    columns_list = ["clients name_", "id numberclients name_"]

    # Create an empty list to store the new values
    new_columns_list = []

    # Iterate over the values in the original list
    for column in columns_list:
        # Iterate over the numbers from 0 to 8
        for i in range(9):
            # Append the original value with the current number
            new_columns_list.append(column + str(i))

    for i in new_columns_list:
        if i not in all_reports4word_sum:
            all_reports4word_sum[i] = ""
        else:
            all_reports4word_sum[i] = all_reports4word_sum[i]

    if "owner name" not in all_reports4word_sum:
        all_reports4word_sum["owner name"] = ""
    else:
        all_reports4word_sum["owner name"] = all_reports4word_sum["owner name"]

    info2r_status = all_reports4word_sum[
        all_reports4word_sum["final type"].str.contains("p" or "d" or "n", regex=False)
    ]
    check4excel = gmt.merge(
        info2r_status[["clients name", "final type"]], on="clients name", how="left"
    )

    info2r_status_ow = all_reports4word_sum[
        all_reports4word_sum["final type"].str.contains("m", regex=False)
    ]
    check4excel_ow = gmt.merge(
        info2r_status_ow[["clients name", "final type"]], on="clients name", how="left"
    )

    check4excel = check4excel.merge(
        check4excel_ow[["deal number", "final type"]],
        on="deal number",
        how="left",
        suffixes=("", "_ow"),
    )
    check4excel = check4excel[
        [
            "deal number",
            "date",
            "clients name",
            "סוג עסקה",
            "amount",
            "currency",
            "dest country",
            "id number",
            "owner name",
            "nis sum",
            "final type",
            "final type_ow",
        ]
    ]

    check4excel["סטטוס סיכון"] = np.where(
        check4excel["deal number"].isin(all_risk_gmt["deal number"]), "סיכון", "תקין"
    )

    check4excel = check4excel.rename(
        columns={
            "deal number": "מזהה פעולה",
            "date": "תאריך פעולה",
            "clients name": "שם הלקוח",
            "amount": "סכום במטבע",
            "currency": "שם המטבע",
            "id number": "מספר תעודה מזהה לקוח",
            "dest country": "יעד",
            "owner name": "שם המושך",
            "nis sum": "סכום בשקלים",
            "final type": "סטטוס",
            "final type_ow": "סטטוס מושך",
        }
    )

    check4excel[["סטטוס", "סטטוס מושך"]] = (
        check4excel[["סטטוס", "סטטוס מושך"]]
        .replace(
            {
                "pd": "פיצול ודירדוס",
                "pr": "פיצול בסיכון",
                "p": "פיצול",
                "dr": "דירדוס בסיכון",
                "d": "דירדוס",
                "m": "מושך משותף",
                "mr": "מושך משותף בסיכון",
                "pm": "פיצול ומושך משותף",
                "nr": "דיווח רגיל בסיכון",
            }
        )
        .fillna("תקין")
    )

    with xlsx_test as writer:
        all_reports4word_excel.to_excel(writer, sheet_name="סטטוס", index=False)

    all_reports4word_sum = all_reports4word_sum.fillna("")
    all_reports4word_sum = all_reports4word_sum.reset_index()

    filepath_check = template_directory + "\*template*"
    txt_check = glob.glob(filepath_check)
    for textfile in txt_check:
        template = textfile
    document = docx.Document(template)
    image_paras = [
        i for i, p in enumerate(document.paragraphs) if "[ChartImage1]" in p.text
    ]
    p = document.paragraphs[image_paras[0]]
    p.text = ""
    r = p.add_run()
    r.add_text("חתימה: ").bold = True
    try:
        # r.add_picture(clients_folder_path + "\signature.png")
        document.save(template_directory + "\my_doc.docx")
    except:
        document.save(template_directory + "\my_doc.docx")

    filepath_check = template_directory + "\*my_doc*"
    txt_check = glob.glob(filepath_check)
    for textfile in txt_check:
        template = textfile
    document = MailMerge(template)

    doc_df = pd.DataFrame(columns=["file_name", "bytes"])

    for i in all_reports4word_sum.index:
        # template1 = 'test1.docx'
        document = MailMerge(template)
        document.merge(
            first_name=str(all_reports4word_sum["first name"][i]),
            last_name=str(all_reports4word_sum["last name"][i]),
            Title=str(all_reports4word_sum["Title"][i]),
            person_citizenship=str(all_reports4word_sum["עברית"][i]),
            Content=str(all_reports4word_sum["content2report"][i]),
            person_id=str(all_reports4word_sum["final id number"][i]),
            report_id=str(all_reports4word_sum["report_id"][i]),
            business_name=str(business_name),
            business_id=str(business_id),
            branch_number=str(branch_number),
            business_adress=str(business_address),
            business_phone=str(business_phone),
            worker_name=str(worker_name),
            date=str(date),
            worker_id=str(worker_id),
            workers_phone=str(worker_phone),
            worker_position=str(worker_position),
            format=str('של העברת כספים לחו"ל באמצעות רשת GMT'),
            clients_name_0=str(all_reports4word_sum["clients name_0"][i]),
            clients_name_1=str(all_reports4word_sum["clients name_1"][i]),
            clients_name_2=str(all_reports4word_sum["clients name_2"][i]),
            clients_name_3=str(all_reports4word_sum["clients name_3"][i]),
            clients_name_4=str(all_reports4word_sum["clients name_4"][i]),
            clients_name_5=str(all_reports4word_sum["clients name_5"][i]),
            clients_name_6=str(all_reports4word_sum["clients name_6"][i]),
            clients_name_7=str(all_reports4word_sum["clients name_7"][i]),
            clients_name_8=str(all_reports4word_sum["clients name_8"][i]),
            owner_name=str(all_reports4word_sum["owner name"][i]),
            id_numberclients_name_0=str(
                all_reports4word_sum["id numberclients name_0"][i]
            ),
            id_numberclients_name_1=str(
                all_reports4word_sum["id numberclients name_1"][i]
            ),
            id_numberclients_name_2=str(
                all_reports4word_sum["id numberclients name_2"][i]
            ),
            id_numberclients_name_3=str(
                all_reports4word_sum["id numberclients name_3"][i]
            ),
            id_numberclients_name_4=str(
                all_reports4word_sum["id numberclients name_4"][i]
            ),
            id_numberclients_name_5=str(
                all_reports4word_sum["id numberclients name_5"][i]
            ),
            id_numberclients_name_6=str(
                all_reports4word_sum["id numberclients name_6"][i]
            ),
            id_numberclients_name_7=str(
                all_reports4word_sum["id numberclients name_7"][i]
            ),
            id_numberclients_name_8=str(
                all_reports4word_sum["id numberclients name_8"][i]
            ),
        )

        # output = str(all_reports4word_sum['report_name'][i]) + "-" + str(
        #     all_reports4word_sum['report_id'][i]) + '.docx'
        # # document.write(output)
        doc_byte = BytesIO()
        document.write(doc_byte)
        doc4df = pd.DataFrame(
            {
                "file_name": [
                    str(all_reports4word_sum["report_name"][i])
                    + "-"
                    + str(all_reports4word_sum["report_id"][i])
                    + ".docx"
                ],
                "bytes": [doc_byte],
            }
        )
        doc_df = doc_df.append(doc4df)

    # indexing the outbound_report data frame
    info2r = sent_report.set_index("clients name")
    info2r_nona = info2r[info2r.index.notnull()]

    doc_df = doc_df.set_index("file_name")

    for title in ls_customer:
        data = sent_report_tr[sent_report_tr["clients name"] == title].round()
        table = [["סהכ", data["nis sum"].sum().round(), data["amount"].sum()]]
        last_row = pd.DataFrame(table, columns=["clients name", "nis sum", "amount"])
        orderd = pd.concat([data, last_row])
        orderd = orderd.fillna(" ")

        orderd = change_names_and_order(gmt_directory + r"\change_names.xlsx", orderd)

        orderd["תאריך פעולה"] = orderd["תאריך פעולה"].astype(str)

        # doc = docx.Document(transaction_folder_path + '/' + str(info2r_nona['report_name'].loc[title]) + "-" + str(
        #     info2r_nona['report_id'].loc[title]) + ".docx")
        file_name = (
            str(info2r_nona["report_name"].loc[title])
            + "-"
            + str(info2r_nona["report_id"].loc[title])
            + ".docx"
        )
        doc = docx.Document(doc_df.loc[file_name].values[0])
        doc.add_page_break()
        t = doc.add_table(orderd.shape[0] + 1, orderd.shape[1])

        for j in range(orderd.shape[-1]):
            t.cell(0, j).text = orderd.columns[j]

            # add the rest of the data frame
        for i in range(orderd.shape[0]):
            for j in range(orderd.shape[-1]):
                t.cell(i + 1, j).text = str(orderd.values[i, j])

        t.style = "Grid Table 4 Accent 5"

        doc.save(doc_df.loc[file_name].values[0])
        # doc.save(transaction_folder_path + '/' + str(info2r_nona['report_name'].loc[title]) + "-" + str(
        #     info2r_nona['report_id'].loc[title]) + ".docx")

    # creating a list containing the client name and owner name
    if sent_report_ow.empty == False:
        info2r_ow = sent_report_ow[["report_name", "report_id", "clients name"]].merge(
            sent_report_tr_ow[["clients name", "owner name"]],
            on="clients name",
            how="left",
        )
        info2r_ow = info2r_ow.drop_duplicates("owner name")
        info2r_ow = info2r_ow.set_index("owner name")

        for title in ls_customer_ow:
            data = sent_report_tr_ow[sent_report_tr_ow["owner name"] == title]
            table = [["סהכ", data["nis sum"].sum().round(), data["amount"].sum()]]
            last_row = pd.DataFrame(table, columns=["owner name", "nis sum", "amount"])
            orderd = pd.concat([data, last_row])
            orderd = orderd.fillna(" ")
            orderd = change_names_and_order(
                gmt_directory + r"\change_names.xlsx", orderd
            )

            # doc = docx.Document(transaction_folder_path + '/' + str(info2r_ow['report_name'].loc[title]) + "-" + str(
            #     info2r_ow['report_id'].loc[title]) + ".docx")
            file_name = (
                str(info2r_ow["report_name"].loc[title])
                + "-"
                + str(info2r_ow["report_id"].loc[title])
                + ".docx"
            )

            doc = docx.Document(doc_df.loc[file_name].values[0])
            doc.add_page_break()
            t = doc.add_table(orderd.shape[0] + 1, orderd.shape[1])

            for j in range(orderd.shape[-1]):
                t.cell(0, j).text = orderd.columns[j]

                # add the rest of the data frame
            for i in range(orderd.shape[0]):
                for j in range(orderd.shape[-1]):
                    t.cell(i + 1, j).text = str(orderd.values[i, j])

            t.style = "Grid Table 4 Accent 5"

            # doc.save(transaction_folder_path + '/' + str(info2r_ow['report_name'].loc[title]) + "-" + str(
            #     info2r_ow['report_id'].loc[title]) + ".docx")
            doc.save(doc_df.loc[file_name].values[0])
    else:
        st.error("אין דיווחי מושך משותף")

    doc_df = doc_df.append(
        pd.DataFrame(
            {"file_name": ["monitoring.xlsx"], "bytes": [xlsx_data]}
        ).set_index("file_name")
    )
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
        for file_name, file_content in doc_df.iterrows():
            # Add each file to the zip file
            file_content = file_content[0]  # Access the BytesIO object
            file_content.seek(0)  # Reset the file pointer to the beginning
            zf.writestr(file_name, file_content.read())

    # Reset the buffer's file pointer to the beginning
    zip_buffer.seek(0)

    download = st.download_button(
        label="📥 Download Reports Zip", data=zip_buffer, file_name="Reports.zip"
    )
    st.error("הניטור הושלם בהצלחה")


def run_changemat(check, client, business_info, reported):
    # defining the directories of different necessary files by the directories Excel file
    directories = pd.read_excel("directories.xlsx")
    directories = pd.DataFrame(directories)
    directories = directories.set_index("File type")
    rc_directory = directories.loc["risk countries"].values[0]
    fsp_directory = directories.loc["financial service providers"].values[0]
    template_directory = directories.loc["word template"].values[0]
    changemat_directory = directories.loc["changemat content and change names"].values[
        0
    ]

    # reading the financial service providers file into a df fsp_list, file must contain str "נותני"
    filepath_check = fsp_directory + "\*נותני*"
    txt_check = glob.glob(filepath_check)
    for textfile in txt_check:
        fsp_list = pd.read_excel(textfile)

    # reading the risk countries file into a DF named "rc", file must contain str "מדינות"
    filepath_check = rc_directory + "\*מדינות*"
    txt_check = glob.glob(filepath_check)
    for textfile in txt_check:
        rc = pd.read_excel(textfile)

    # reading the content file into a DF named "content", file must contain str "content"
    filepath_check = changemat_directory + "\*content*"
    txt_check = glob.glob(filepath_check)
    for textfile in txt_check:
        content = pd.read_excel(textfile)

    # tkinter message for files were read successfully
    st.error("הקבצים נקראו בהצלחה")

    # striping the tr and client DF columns from whitespaces
    tr = check.rename(columns=lambda x: str(x).strip())
    client = client.rename(columns=lambda x: str(x).strip())

    # renaming the "tr" DF columns
    tr = tr.rename(
        columns={
            "תאריך רישום": "date",
            "פעולה": "transaction id",
            "סוג פעולה": "transaction type",
            "כיוון": "in/out",
            "מטבע": "curr",
            "סכום": "amount",
            "שער": "rate",
            "ארץ": "country",
            "זהות/דרכון": "id number",
            "שם": "clients name",
            "שק - בנק": "bank number",
            "ס בנק": "branch number",
            "חשבון": "account number",
            "מספר שק": "check number",
        }
    )

    # defining the "tr" date column as date in order to remove the extra timestamp
    tr["date"] = pd.to_datetime(tr["date"], errors="coerce").dt.date

    # renaming the client df columns
    client = client.rename(
        columns={
            "ארץ": "country",
            "מספר זיהוי": "id number",
            "שם": "clients name",
            "כתובת": "address",
            "עיר": "city",
            "מין": "sex",
            "תאריך לידה": "dob",
        }
    )

    # try defining the dob column as date in order to remove the extra timestamp,
    # except error if dob column is missing add dob column with blank values
    try:
        client["dob"] = pd.to_datetime(client["dob"], errors="coerce").dt.date
    except KeyError:
        client["dob"] = ""

    # drop transaction with clients name missing
    tr = tr.dropna(subset="clients name")

    # drop transactions with amount missing
    tr = tr.dropna(subset="amount")

    # filter out clients with blank names of whitespace names
    tr = tr[tr["clients name"] != ""]
    tr = tr[tr["clients name"] != " "]

    # define the transaction id number column as type int
    tr["transaction id"] = tr["transaction id"].astype(int)

    # define the amount column as type int
    tr["amount"] = tr["amount"].astype(int)

    # clear the id number column and the clients name column from whitespaces
    tr["id number"] = tr["id number"].astype(str)
    tr["id number"] = tr["id number"].str.strip()
    tr["clients name"] = tr["clients name"].str.strip()

    # create a new column named nis sum by multipling the rate with the amount
    tr["nis sum"] = tr["amount"] * tr["rate"]

    # define the nis sum columns as type int64
    tr["nis sum"] = tr["nis sum"].astype("int64")

    # risk check by merging the tr country with the risk countries df
    tr = tr.merge(
        rc[["Alpha-2 code", "סיכון/ללא דיווח רגיל"]],
        left_on="country",
        right_on="Alpha-2 code",
        how="left",
    )

    # filtering only Nis currency transactions
    tr_only_nis = tr[tr["curr"] == "ILS"]

    # filtering only exchange transactions into a new DF named exchange:
    exchange = tr_only_nis[
        (tr_only_nis["transaction type"] == "המרה")
        | (tr_only_nis["transaction type"] == "המרה בבנק")
    ]

    # sort values by clients name and date
    exchange = exchange.sort_values(["clients name", "date"])

    # adding the sum for every deal number
    deal_sum = exchange.groupby("transaction id").agg({"nis sum": "sum"})
    exchange = exchange.merge(
        deal_sum, on="transaction id", how="left", suffixes=("", "_deal")
    )

    # Adding a column stating if the transaction is under or over 50K(False = over 50K)
    exchange["U50"] = np.where(exchange["nis sum_deal"] >= 50000, False, True)

    # Adding a column stating if the transaction is under or over 5K(False = over 5K)
    exchange["U5"] = np.where(exchange["nis sum_deal"] >= 5000, False, True)

    # creating a column with the dates difference by day and creating a new column with the value as a number
    exchange["date"] = pd.to_datetime(exchange["date"])
    exchange["date_diff"] = exchange["date"].diff()
    exchange["date_diff_fl"] = exchange["date_diff"] / pd.to_timedelta(1, unit="D")

    # testing for matching clients in match column
    exchange["match"] = exchange["clients name"].eq(exchange["clients name"].shift())

    # adding a conditional column based on date diff values
    exchange["diff_check"] = np.where(
        (exchange["date_diff_fl"] > 3)
        | (exchange["date_diff_fl"] < 0)
        | (exchange["match"] == False),
        1,
        0,
    )

    # cumulative sum as a way to count if the values in the diff check
    exchange["cumsum_xc"] = np.cumsum(exchange["diff_check"])

    # risk exchange df from the exchange df by origin of client
    risk_exchange = exchange[exchange["סיכון/ללא דיווח רגיל"].notna()]

    # building the monitoring df with sum per client, deal distinct count and under 50k deal count
    mont_xc = exchange.groupby(["clients name", "cumsum_xc"]).agg(
        {"nis sum": "sum", "transaction id": "nunique", "U50": "sum"}
    )
    mont_xc = mont_xc.reset_index()

    # sorting the values by sum
    mont_xc = mont_xc.sort_values("nis sum", ascending=False)
    mont_xc["U50"] = mont_xc["U50"].astype(int)

    # don't preform a three month reported test if the DF is empty
    if reported["שם"].empty:
        mont_xc["שם"] = ""
    else:
        # preform a reported in the last two-month test
        mont_xc = mont_xc.merge(
            reported["שם"].astype(str),
            left_on="clients name",
            right_on="שם",
            how="left",
        )

    # adding a column named status based on unique count of deal number and at least one transaction
    # under 50k and not reported in the past 2m
    mont_xc["status"] = np.where(
        (mont_xc["transaction id"] == 1)
        | (mont_xc["U50"] == 0)
        | (mont_xc["שם"].notna()),
        "Regular Report",
        "Check",
    )
    mont_xc = mont_xc.drop_duplicates("cumsum_xc")

    # preform a financial service providers test - results stored in a column named fsp_check
    # adding the id number to the monitoring sheet
    mont_xc = mont_xc.merge(
        tr[["id number", "clients name"]], on="clients name", how="left"
    )

    # adding a column מספר מזהה to the monitoring sheet values indicates fsp(financial service provider)
    try:
        mont_xc = mont_xc.merge(
            fsp_list["מספר מזהה"], left_on="id number", right_on="מספר מזהה", how="left"
        )
    except ValueError:
        mont_xc = mont_xc.merge(
            fsp_list["מספר מזהה"].astype(str),
            left_on="id number",
            right_on="מספר מזהה",
            how="left",
        )
    mont_xc["fsp_check"] = np.where(
        mont_xc["מספר מזהה"].isna(), "check", "financial service provider"
    )

    mont_xc = mont_xc.drop_duplicates("cumsum_xc")
    # reported test
    mont_xc["reported"] = np.where(mont_xc["שם"].notna(), "reported", "check")

    # marking dirdos transaction as d in a new column named dirdos
    mont_xc["dirdos"] = np.where(
        (mont_xc["nis sum"] >= 47000)
        & (mont_xc["nis sum"] < 50000)
        & (mont_xc["שם"].isna())
        & (mont_xc["fsp_check"] == "check")
        & (mont_xc["U50"] > 0),
        "xd",
        "not xd",
    )

    # creating a new data frame containing only clients that need to be reported
    mont2r_xc = mont_xc[
        ((mont_xc["status"] == "Check") | (mont_xc["dirdos"] == "xd"))
        & (mont_xc["fsp_check"] == "check")
        & (mont_xc["nis sum"] >= 47000)
    ]
    mont2r_xc = mont2r_xc.drop_duplicates("cumsum_xc")

    # adding the type of the report p-pitzul d-dirdos
    mont2r_xc["type"] = np.where(mont2r_xc["nis sum"] >= 50000, "xp", "xd")

    # figure out the problem of a client with multiple types of reports
    type_count = mont2r_xc[["clients name", "type"]].groupby("clients name").nunique()

    # adding situations in which there are two types of reports for a client
    mont2r_xc = mont2r_xc.merge(
        type_count, on="clients name", how="left", suffixes=("", "_count")
    )
    mont2r_xc["final xc type"] = np.where(
        mont2r_xc["type_count"] > 1, "xpd", mont2r_xc["type"]
    )

    # creating a new data frame with duplicate cumsum for report details df= mont2report_w_cumsum
    mont2r_xc_w_cumsum = mont2r_xc
    mont2r_xc = mont2r_xc.drop_duplicates("clients name")
    info2r = mont2r_xc[
        ["clients name", "id number", "final xc type", "cumsum_xc", "nis sum"]
    ]

    # building the monitoring df with sum per client, deal distinct count and under 50k deal count
    mont_xc_risk = risk_exchange.groupby(["clients name", "cumsum_xc"]).agg(
        {"nis sum": "sum", "transaction id": "nunique", "U5": "sum"}
    )
    mont_xc_risk = mont_xc_risk.reset_index()

    # sorting the values by sum
    mont_xc_risk = mont_xc_risk.sort_values("nis sum", ascending=False)
    mont_xc_risk["U5"] = mont_xc_risk["U5"].astype(int)

    # preform a reported in the last two-month test
    mont_xc_risk = mont_xc_risk.merge(
        reported["שם"], left_on="clients name", right_on="שם", how="left"
    )

    # adding a column named status based on unique count of deal number
    # and at least on transactions under 50k and not reported in the past 2m
    mont_xc_risk["status"] = np.where(
        (mont_xc_risk["transaction id"] == 1)
        | (mont_xc_risk["U5"] == 0)
        | (mont_xc_risk["שם"].notna()),
        "Regular Report",
        "Check",
    )

    mont_xc_risk = mont_xc_risk.drop_duplicates("cumsum_xc")

    # preform a financial service providers test - results stored in a column named fsp_check
    # adding the id number to the monitoring sheet
    mont_xc_risk = mont_xc_risk.merge(
        tr[["id number", "clients name"]], on="clients name", how="left"
    )

    # adding a column מספר מזהה to the monitoring sheet values indicates fsp(financial service provider)
    try:
        mont_xc_risk = mont_xc_risk.merge(
            fsp_list["מספר מזהה"], left_on="id number", right_on="מספר מזהה", how="left"
        )
    except ValueError:
        mont_xc_risk = mont_xc_risk.merge(
            fsp_list["מספר מזהה"].astype(str),
            left_on="id number",
            right_on="מספר מזהה",
            how="left",
        )
    mont_xc_risk["fsp_check"] = np.where(
        mont_xc_risk["מספר מזהה"].isna(), "check", "financial service provider"
    )

    mont_xc_risk = mont_xc_risk.drop_duplicates("cumsum_xc")
    # reported test
    mont_xc_risk["reported"] = np.where(mont_xc_risk["שם"].notna(), "reported", "check")

    # marking dirdos transaction as d in a new column named dirdos
    mont_xc_risk["dirdos_risk"] = np.where(
        (mont_xc_risk["nis sum"] >= 4700)
        & (mont_xc_risk["nis sum"] < 5000)
        & (mont_xc_risk["שם"].isna())
        & (mont_xc_risk["fsp_check"] == "check")
        & (mont_xc_risk["U5"] > 0),
        "xdr",
        "not xdr",
    )

    # creating a new data frame containing only clients that need to be reported
    mont2r_xc_risk = mont_xc_risk[
        ((mont_xc_risk["status"] == "Check") | (mont_xc_risk["dirdos_risk"] == "xdr"))
        & (mont_xc_risk["fsp_check"] == "check")
        & (mont_xc_risk["nis sum"] >= 4700)
    ]

    mont2r_xc_risk = mont2r_xc_risk.drop_duplicates("cumsum_xc")

    # adding the type of the report p-pitzul d-dirdos
    mont2r_xc_risk["type_risk"] = np.where(
        mont2r_xc_risk["nis sum"] >= 5000, "xpr", "xdr"
    )

    # figure out the problem of a client with multiple types of reports
    type_count_risk = (
        mont2r_xc_risk[["clients name", "type_risk"]].groupby("clients name").nunique()
    )

    # adding situations in which there are two types of reports for a client
    mont2r_xc_risk = mont2r_xc_risk.merge(
        type_count_risk, on="clients name", how="left", suffixes=("", "_count")
    )

    mont2r_xc_risk["final type xc risk"] = np.where(
        mont2r_xc_risk["type_risk_count"] > 1, "xpdr", mont2r_xc_risk["type_risk"]
    )

    # creating a new data frame with duplicate cumsum for report details df= mont2report_w_cumsum
    mont2r_xc_w_cumsum = mont2r_xc_w_cumsum.append(mont2r_xc_risk)
    mont2r_xc_w_cumsum = mont2r_xc_w_cumsum.fillna("")
    mont2r_xc_w_cumsum["final type xc"] = (
        mont2r_xc_w_cumsum["final xc type"] + mont2r_xc_w_cumsum["final type xc risk"]
    )
    mont2r_xc_risk = mont2r_xc_risk.drop_duplicates("clients name")
    info2r = info2r.append(
        mont2r_xc_risk[
            ["clients name", "id number", "final type xc risk", "cumsum_xc", "nis sum"]
        ]
    )

    # filtering only check transactions into a new DF named check:
    check = tr_only_nis[
        (tr_only_nis["transaction type"] == "נכיונות")
        | (tr_only_nis["transaction type"] == "מסירה")
    ]
    check = check[check["in/out"] == "כניסה"]

    # sort values by clients name and date
    check = check.sort_values(["clients name", "date"])

    # adding the sum for every deal number and marking every deal under 50k as True
    deal_sum = check.groupby("transaction id").agg({"nis sum": "sum"})
    check = check.merge(
        deal_sum, on="transaction id", how="left", suffixes=("", "_deal")
    )

    # Adding a column stating if the transaction is under or over 50K(False = over 50K)
    check["U50"] = np.where(check["nis sum_deal"] >= 50000, False, True)

    # Adding a column stating if the transaction is under or over 5K(False = over 5K)
    check["U5"] = np.where(check["nis sum_deal"] >= 5000, False, True)

    # creating a column with the dates difference by day and creating a new column with the value as a number
    check["date"] = pd.to_datetime(check["date"])
    check["date_diff"] = check["date"].diff()
    check["date_diff_fl"] = check["date_diff"] / pd.to_timedelta(1, unit="D")

    # testing for matching clients in match column
    check["match"] = check["clients name"].eq(check["clients name"].shift())

    # adding a conditional colum based on date diff values
    check["diff_check"] = np.where(
        (check["date_diff_fl"] > 3)
        | (check["date_diff_fl"] < 0)
        | (check["match"] == False),
        1,
        0,
    )

    # cumulative sum as a way to count if the values in the diff check
    check["cumsum"] = np.cumsum(check["diff_check"])

    # risk check
    risk_check = check[check["סיכון/ללא דיווח רגיל"].notna()]

    # building the monitoring df with sum per client, deal distinct count and under 50k deal count
    mont_check = check.groupby(["clients name", "cumsum"]).agg(
        {"nis sum": "sum", "transaction id": "nunique", "U50": "sum"}
    )
    mont_check = mont_check.reset_index()

    # sorting the values by sum
    mont_check = mont_check.sort_values("nis sum", ascending=False)
    mont_check["U50"] = mont_check["U50"].astype(int)

    # preform a reported in the last two-month test
    mont_check = mont_check.merge(
        reported["שם"], left_on="clients name", right_on="שם", how="left"
    )

    # adding a column named status based on unique count of deal number and at least one transaction
    # under 50k and not reported in the past 2m
    mont_check["status"] = np.where(
        (mont_check["transaction id"] == 1)
        | (mont_check["U50"] == 0)
        | (mont_check["שם"].notna()),
        "Regular Report",
        "Check",
    )
    mont_check = mont_check.drop_duplicates("cumsum")

    # preform a financial service providers test - results stored in a column named fsp_check
    # adding the id number to the monitoring sheet
    mont_check = mont_check.merge(
        tr[["id number", "clients name"]], on="clients name", how="left"
    )

    # adding a column מספר מזהה to the monitoring sheet values indicates fsp(financial service provider)
    try:
        mont_check = mont_check.merge(
            fsp_list["מספר מזהה"], left_on="id number", right_on="מספר מזהה", how="left"
        )
    except ValueError:
        mont_check = mont_check.merge(
            fsp_list["מספר מזהה"].astype(str),
            left_on="id number",
            right_on="מספר מזהה",
            how="left",
        )
    mont_check["fsp_check"] = np.where(
        mont_check["מספר מזהה"].isna(), "check", "financial service provider"
    )

    mont_check = mont_check.drop_duplicates("cumsum")
    # reported test
    mont_check["reported"] = np.where(mont_check["שם"].notna(), "reported", "check")

    # marking dirdos transaction as d in a new column named dirdos
    mont_check["dirdos"] = np.where(
        (mont_check["nis sum"] >= 47000)
        & (mont_check["nis sum"] < 50000)
        & (mont_check["שם"].isna())
        & (mont_check["fsp_check"] == "check")
        & (mont_check["U50"] > 0),
        "d",
        "not d",
    )

    # creating a new data frame containing only clients that need to be reported
    mont2r_check = mont_check[
        ((mont_check["status"] == "Check") | (mont_check["dirdos"] == "d"))
        & (mont_check["fsp_check"] == "check")
        & (mont_check["nis sum"] >= 47000)
    ]
    mont2r_check = mont2r_check.drop_duplicates("cumsum")

    # adding the type of the report p-pitzul d-dirdos
    mont2r_check["type"] = np.where(mont2r_check["nis sum"] >= 50000, "p", "d")

    # figure out the problem of a client with multiple types of reports
    type_count = (
        mont2r_check[["clients name", "type"]].groupby("clients name").nunique()
    )

    # adding situations in which there are two types of reports for a client
    mont2r_check = mont2r_check.merge(
        type_count, on="clients name", how="left", suffixes=("", "_count")
    )
    mont2r_check["final type"] = np.where(
        mont2r_check["type_count"] > 1, "pd", mont2r_check["type"]
    )

    # creating a new data frame with duplicate cumsum for report details df= mont2report_w_cumsum
    mont2r_check_w_cumsum = mont2r_check
    mont2r_check = mont2r_check.drop_duplicates("clients name")

    # building the monitoring df with sum per client, deal distinct count and under 50k deal count
    mont_check_risk = risk_check.groupby(["clients name", "cumsum"]).agg(
        {"nis sum": "sum", "transaction id": "nunique", "U5": "sum"}
    )
    mont_check_risk = mont_check_risk.reset_index()

    # sorting the values by sum
    mont_check_risk = mont_check_risk.sort_values("nis sum", ascending=False)
    mont_check_risk["U5"] = mont_check_risk["U5"].astype(int)

    # preform a reported in the last two-month test
    mont_check_risk = mont_check_risk.merge(
        reported["שם"], left_on="clients name", right_on="שם", how="left"
    )

    # adding a column named status based on unique count of deal number
    # and at least on transactions under 50k and not reported in the past 2m
    mont_check_risk["status"] = np.where(
        (mont_check_risk["transaction id"] == 1)
        | (mont_check_risk["U5"] == 0)
        | (mont_check_risk["שם"].notna()),
        "Regular Report",
        "Check",
    )

    mont_check_risk = mont_check_risk.drop_duplicates("cumsum")

    # preform a financial service providers test - results stored in a column named fsp_check
    # adding the id number to the monitoring sheet
    mont_check_risk = mont_check_risk.merge(
        tr[["id number", "clients name"]], on="clients name", how="left"
    )

    # adding a column מספר מזהה to the monitoring sheet values indicates fsp(financial service provider)
    try:
        mont_check_risk = mont_check_risk.merge(
            fsp_list["מספר מזהה"], left_on="id number", right_on="מספר מזהה", how="left"
        )
    except ValueError:
        mont_check_risk = mont_check_risk.merge(
            fsp_list["מספר מזהה"].astype(str),
            left_on="id number",
            right_on="מספר מזהה",
            how="left",
        )
    mont_check_risk["fsp_check"] = np.where(
        mont_check_risk["מספר מזהה"].isna(), "check", "financial service provider"
    )

    mont_check_risk = mont_check_risk.drop_duplicates("cumsum")
    # reported test
    mont_check_risk["reported"] = np.where(
        mont_check_risk["שם"].notna(), "reported", "check"
    )

    # marking dirdos transaction as d in a new column named dirdos
    mont_check_risk["dirdos_risk"] = np.where(
        (mont_check_risk["nis sum"] >= 4700)
        & (mont_check_risk["nis sum"] < 5000)
        & (mont_check_risk["שם"].isna())
        & (mont_check_risk["fsp_check"] == "check")
        & (mont_check_risk["U5"] > 0),
        "dr",
        "not dr",
    )

    # creating a new data frame containing only clients that need to be reported
    mont2r_check_risk = mont_check_risk[
        (
            (mont_check_risk["status"] == "Check")
            | (mont_check_risk["dirdos_risk"] == "dr")
        )
        & (mont_check_risk["fsp_check"] == "check")
        & (mont_check_risk["nis sum"] >= 4700)
    ]

    mont2r_check_risk = mont2r_check_risk.drop_duplicates("cumsum")

    # adding the type of the report p-pitzul d-dirdos
    mont2r_check_risk["type_risk"] = np.where(
        mont2r_check_risk["nis sum"] >= 5000, "pr", "dr"
    )

    # figure out the problem of a client with multiple types of reports
    type_count_risk = (
        mont2r_check_risk[["clients name", "type_risk"]]
        .groupby("clients name")
        .nunique()
    )

    # adding situations in which there are two types of reports for a client
    mont2r_check_risk = mont2r_check_risk.merge(
        type_count_risk, on="clients name", how="left", suffixes=("", "_count")
    )

    mont2r_check_risk["final type risk"] = np.where(
        mont2r_check_risk["type_risk_count"] > 1, "pdr", mont2r_check_risk["type_risk"]
    )
    mont2r_check_w_cumsum = mont2r_check_w_cumsum.append(mont2r_check_risk)
    mont2r_check_w_cumsum = mont2r_check_w_cumsum.fillna("")
    mont2r_check_w_cumsum["final type check"] = (
        mont2r_check_w_cumsum["final type"] + mont2r_check_w_cumsum["final type risk"]
    )
    mont2r_check_risk = mont2r_check_risk.drop_duplicates("clients name")

    info2r = info2r.append(
        mont2r_check[["clients name", "id number", "final type", "cumsum", "nis sum"]]
    )
    info2r = info2r.append(
        mont2r_check_risk[
            ["clients name", "id number", "final type risk", "cumsum", "nis sum"]
        ]
    )

    check_account = check.dropna(subset="account number")

    # sort values by clients name and date
    check_account = check_account.sort_values(["account number", "date"])

    # adding the sum for every deal number and marking every deal under 50k as True
    deal_sum = check_account.groupby("transaction id").agg({"nis sum": "sum"})
    check_account = check_account.merge(
        deal_sum, on="transaction id", how="left", suffixes=("", "_deal")
    )

    # Adding a column stating if the transaction is under or over 50K(False = over 50K)
    check_account["U50"] = np.where(check_account["nis sum_deal"] >= 50000, False, True)

    # Adding a column stating if the transaction is under or over 5K(False = over 5K)
    check_account["U5"] = np.where(check_account["nis sum_deal"] >= 5000, False, True)

    # creating a column with the dates difference by day and creating a new column with the value as a number
    check_account["date"] = pd.to_datetime(check_account["date"])
    check_account["date_diff"] = check_account["date"].diff()
    check_account["date_diff_fl"] = check_account["date_diff"] / pd.to_timedelta(
        1, unit="D"
    )

    # testing for matching clients in match column
    check_account["match"] = check_account["account number"].eq(
        check_account["account number"].shift()
    )

    # adding a conditional colum based on date diff values
    check_account["diff_check"] = np.where(
        (check_account["date_diff_fl"] > 3)
        | (check_account["date_diff_fl"] < 0)
        | (check_account["match"] == False),
        1,
        0,
    )

    # cumulative sum as a way to count if the values in the diff check
    check_account["cumsum_ow"] = np.cumsum(check_account["diff_check"])

    # risk check
    risk_check_account = check_account[check_account["סיכון/ללא דיווח רגיל"].notna()]

    # building the monitoring df with sum per client, deal distinct count and under 50k deal count
    mont_check_account = check_account.groupby(["account number", "cumsum_ow"]).agg(
        {
            "nis sum": "sum",
            "transaction id": "nunique",
            "U50": "sum",
            "clients name": "nunique",
        }
    )

    mont_check_account = mont_check_account.rename(
        columns={"clients name": "clients name_count"}
    )

    mont_check_account = mont_check_account.reset_index()

    # sorting the values by sum
    mont_check_account = mont_check_account.sort_values("nis sum", ascending=False)
    mont_check_account["U50"] = mont_check_account["U50"].astype(int)

    mont_check_account = mont_check_account.merge(
        check_account[["cumsum_ow", "clients name"]], on="cumsum_ow", how="left"
    )
    # preform a reported in the last two-month test
    mont_check_account = mont_check_account.merge(
        reported["שם"], left_on="clients name", right_on="שם", how="left"
    )

    # adding a column named status based on unique count of deal number and at least one transaction
    # under 50k and not reported in the past 2m
    mont_check_account["status"] = np.where(
        (mont_check_account["clients name_count"] == 1)
        | (mont_check_account["U50"] == 0)
        | (mont_check_account["שם"].notna()),
        "Regular Report",
        "Check",
    )

    # preform a financial service providers test - results stored in a column named fsp_check
    # adding the id number to the monitoring sheet
    mont_check_account = mont_check_account.merge(
        check_account[["id number", "clients name"]], on="clients name", how="left"
    )

    # adding a column מספר מזהה to the monitoring sheet values indicates fsp(financial service provider)
    try:
        mont_check_account = mont_check_account.merge(
            fsp_list["מספר מזהה"], left_on="id number", right_on="מספר מזהה", how="left"
        )
    except ValueError:
        mont_check_account = mont_check_account.merge(
            fsp_list["מספר מזהה"].astype(str),
            left_on="id number",
            right_on="מספר מזהה",
            how="left",
        )
    mont_check_account["fsp_check"] = np.where(
        mont_check_account["מספר מזהה"].isna(), "check", "financial service provider"
    )

    # reported test
    mont_check_account["reported"] = np.where(
        mont_check_account["שם"].notna(), "reported", "check"
    )

    # creating a new data frame containing only clients that need to be reported
    mont2r_check_account = mont_check_account[
        (mont_check_account["status"] == "Check")
        & (mont_check_account["fsp_check"] == "check")
        & (mont_check_account["nis sum"] >= 50000)
    ]

    # adding the type of the report p-pitzul d-dirdos m-mosheh
    mont2r_check_account["type"] = "m"

    # creating a new data frame with duplicate cumsum for report details df= mont2report_w_cumsum
    mont2r_check_account_w_cumsum = mont2r_check_account

    mont2r_check_account = mont2r_check_account[
        mont2r_check_account["reported"] == "check"
    ]
    mont2r_check_account = mont2r_check_account[
        -mont2r_check_account["clients name"].isin(info2r["clients name"])
    ]

    mont2r_check_account = mont2r_check_account.drop_duplicates("cumsum_ow")

    info2r = info2r.append(
        mont2r_check_account[
            ["clients name", "id number", "type", "cumsum_ow", "nis sum"]
        ]
    )

    # building the monitoring df with sum per client, deal distinct count and under 50k deal count
    mont_risk_check_account = risk_check_account.groupby(
        ["account number", "cumsum_ow"]
    ).agg(
        {
            "nis sum": "sum",
            "transaction id": "nunique",
            "U5": "sum",
            "clients name": "nunique",
        }
    )

    mont_risk_check_account = mont_risk_check_account.rename(
        columns={"clients name": "clients name_count"}
    )

    mont_risk_check_account = mont_risk_check_account.reset_index()

    # sorting the values by sum
    mont_risk_check_account = mont_risk_check_account.sort_values(
        "nis sum", ascending=False
    )
    mont_risk_check_account["U50"] = mont_risk_check_account["U5"].astype(int)

    mont_risk_check_account = mont_risk_check_account.merge(
        risk_check_account[["cumsum_ow", "clients name"]], on="cumsum_ow", how="left"
    )
    # preform a reported in the last two-month test
    mont_risk_check_account = mont_risk_check_account.merge(
        reported["שם"], left_on="clients name", right_on="שם", how="left"
    )

    # adding a column named status based on unique count of deal number and at least one transaction
    # under 50k and not reported in the past 2m
    mont_risk_check_account["status"] = np.where(
        (mont_risk_check_account["clients name_count"] == 1)
        | (mont_risk_check_account["U5"] == 0)
        | (mont_risk_check_account["שם"].notna()),
        "Regular Report",
        "Check",
    )

    # preform a financial service providers test - results stored in a column named fsp_check
    # adding the id number to the monitoring sheet
    mont_risk_check_account = mont_risk_check_account.merge(
        risk_check_account[["id number", "clients name"]], on="clients name", how="left"
    )

    # adding a column מספר מזהה to the monitoring sheet values indicates fsp(financial service provider)
    try:
        mont_risk_check_account = mont_risk_check_account.merge(
            fsp_list["מספר מזהה"], left_on="id number", right_on="מספר מזהה", how="left"
        )
    except ValueError:
        mont_risk_check_account = mont_risk_check_account.merge(
            fsp_list["מספר מזהה"].astype(str),
            left_on="id number",
            right_on="מספר מזהה",
            how="left",
        )
    mont_risk_check_account["fsp_check"] = np.where(
        mont_risk_check_account["מספר מזהה"].isna(),
        "check",
        "financial service provider",
    )

    # reported test
    mont_risk_check_account["reported"] = np.where(
        mont_risk_check_account["שם"].notna(), "reported", "check"
    )

    # creating a new data frame containing only clients that need to be reported
    mont2r_risk_check_account = mont_risk_check_account[
        (mont_risk_check_account["status"] == "Check")
        & (mont_risk_check_account["fsp_check"] == "check")
        & (mont_risk_check_account["nis sum"] >= 5000)
    ]

    # adding the type of the report p-pitzul d-dirdos m-mosheh
    mont2r_risk_check_account["type risk"] = "mr"

    # creating a new data frame with duplicate cumsum for report details df= mont2report_w_cumsum
    mont2r_risk_check_account_w_cumsum = mont2r_risk_check_account

    mont2r_risk_check_account = mont2r_risk_check_account[
        mont2r_risk_check_account["reported"] == "check"
    ]
    mont2r_risk_check_account = mont2r_risk_check_account[
        -mont2r_risk_check_account["clients name"].isin(info2r["clients name"])
    ]

    mont2r_risk_check_account = mont2r_risk_check_account.drop_duplicates("cumsum_ow")

    info2r = info2r.append(
        mont2r_risk_check_account[
            ["clients name", "id number", "type risk", "cumsum_ow", "nis sum"]
        ]
    )

    for column in info2r.columns:
        if column.startswith("cumsum"):
            info2r[column] = info2r[column].fillna("")
            info2r[column] = info2r[column].apply(
                lambda x: round(x) if isinstance(x, float) else x
            )

    info2r["nis sum"] = info2r["nis sum"].astype("int64")

    info2r = info2r.fillna("")
    info2r["real final type"] = (
        info2r["final xc type"]
        + info2r["final type xc risk"]
        + info2r["final type"]
        + info2r["final type risk"]
        + info2r["type"]
        + info2r["type risk"]
    )
    # todo this try except might cause problems with the clients details in reports
    try:
        info2r = info2r[
            [
                "clients name",
                "cumsum",
                "cumsum_ow",
                "cumsum_xc",
                "nis sum",
                "real final type",
            ]
        ].merge(
            client[
                [
                    "clients name",
                    "id number",
                    "dob",
                    "sex",
                    "country",
                    "city",
                    "address",
                ]
            ],
            on="clients name",
            how="left",
        )
    except KeyError:
        info2r = info2r[
            [
                "clients name",
                "cumsum",
                "cumsum_ow",
                "cumsum_xc",
                "nis sum",
                "real final type",
            ]
        ].merge(
            client[["clients name", "id number", "country"]],
            on="clients name",
            how="left",
        )

    # creating columns for private customer or company based on the sex column
    info2r["person name"] = np.where(
        info2r["sex"] == "תאגיד", "", info2r["clients name"]
    )
    info2r["company name"] = np.where(
        info2r["sex"] == "תאגיד", info2r["clients name"], ""
    )
    info2r["person dob"] = np.where(info2r["sex"] == "תאגיד", "", info2r["dob"])
    info2r["company dob"] = np.where(info2r["sex"] == "תאגיד", info2r["dob"], "")
    info2r["person id number"] = np.where(
        info2r["sex"] == "תאגיד", "", info2r["id number"]
    )
    info2r["company id number"] = np.where(
        info2r["sex"] == "תאגיד", info2r["id number"], ""
    )

    # merging the content to the info2r DF
    info2r = info2r.merge(
        content, left_on="real final type", right_on="type", how="left"
    )

    # fixing the business info
    business_info = business_info
    business_info = business_info.rename(columns=lambda x: str(x).strip())
    business_info = business_info.fillna("")
    business_type = business_info.loc["סוג הגוף המדווח"].values[0]
    business_name = business_info.loc["שם מלא של הגורם המדווח"].values[0]
    business_id = business_info.loc["מספר מזהה של הגורם המדווח"].values[0]
    branch_number = business_info.loc["מספר סניף מדווח"].values[0]
    business_address = business_info.loc["מען הסניף"].values[0]
    date = to_date.today()
    business_phone = business_info.loc["טלפון ופקסימילה של הגורם המדווח"].values[0]
    worker_name = business_info.loc["שם פרטי ושם משפחה של עובד עורך הדיווח"].values[0]
    worker_id = business_info.loc['ת"ז של עורך הדיווח'].values[0]
    worker_position = business_info.loc["תפקיד עורך הדיווח"].values[0]
    worker_phone = business_info.loc["טלפון עורך הדיווח"].values[0]
    worker_email = business_info.loc["דואר אלקטרוני"].values[0]

    info2r = info2r.merge(
        rc[["Alpha-2 code", "עברית"]],
        left_on="country",
        right_on="Alpha-2 code",
        how="left",
    )

    # adding the country of origin to the report content
    for i in info2r.index:
        info2r["Title"].iloc[i] = str(info2r["Title"].iloc[i]).replace(
            "מדינה", str(info2r["עברית"].iloc[i]) + ", מדינה"
        )
        info2r["Content"].iloc[i] = str(info2r["Content"].iloc[i]).replace(
            "מדינה", str(info2r["עברית"].iloc[i]) + ", מדינה"
        )
    info2r["country"].notna = info2r["עברית"]
    info2r = info2r.drop(columns=["עברית", "Alpha-2 code"])

    # dropping duplicates and resetting index
    info2r = info2r.drop_duplicates("clients name")
    info2r = info2r.reset_index(drop=True)
    # importing the report number needed
    info2r["report_id_aid"] = reported["מס דיווח"].max() + 1
    info2r["report_id"] = (info2r["report_id_aid"] + range(len(info2r.index))).astype(
        int
    )
    info2r["report_name"] = report_num_gen(business_id, branch_number)

    # creating a DF named info2r_excel with the column for the final report
    info2r_excel = info2r[
        [
            "clients name",
            "real final type",
            "id number",
            "cumsum",
            "cumsum_xc",
            "cumsum_ow",
            "nis sum",
            "dob",
            "sex",
            "country",
            "report_id",
            "Title",
            "Content",
        ]
    ]

    # renaming the columns of info2r_excel DF
    info2r_excel = info2r_excel.rename(
        columns={
            "clients name": "שם הלקוח",
            "real final type": "מהות הפעילות",
            "id number": "מספר תעודה",
            "dob": "תאריך לידה\התאגדות",
            "sex": "מין",
            "country": "מדינת תעודה\התאגדות",
            "report_id": "מספר דיווח",
            "Title": "תמצית",
            "Content": "תוכן",
            "cumsum": "מזהה קבוצה ניכיון",
            "cumsum_ow": "מזהה קבוצה מושך",
            "cumsum_xc": "מזהה קבוצה המרה",
            "nis sum": "סכום פעולות לדיווח",
        }
    )

    # creating new empty columns for the customer to write notes
    info2r_excel["לדווח/לא לדווח"] = ""
    info2r_excel["הערות"] = ""

    # replacing the content in מהות הםעילות from english shortcut to heb real name
    info2r_excel["מהות הפעילות"] = info2r_excel["מהות הפעילות"].replace(
        {
            "pd": "פיצול ודירדוס",
            "pr": "פיצול בסיכון",
            "p": "פיצול",
            "dr": "דירדוס בסיכון",
            "d": "דירדוס",
            "m": "מושך משותף",
            "mr": "מושך משותף בסיכון",
            "pm": "פיצול ומושך משותף",
            "dm": "דירדוס ומושך משותף",
            "xp": "פיצול המרה",
            "xd": "דירדוס המרה",
        }
    )

    xlsx_data = BytesIO()
    with pd.ExcelWriter(xlsx_data, engine="openpyxl") as writer:
        info2r_excel.to_excel(writer, sheet_name="ניהול התראות")
    xlsx_data.seek(0)

    # report_tr is a DF containing all the cumsum to report for checks cleaning
    report_tr = mont2r_check_w_cumsum[["clients name", "cumsum"]]

    # check_min is a df with all the check transactions but with only the relevant info for report
    check_min = check[
        [
            "1",
            "transaction id",
            "date",
            "transaction type",
            "אמצעי",
            "in/out",
            "clients name",
            "check number",
            "bank number",
            "account number",
            "branch number",
            "nis sum",
            "cumsum",
        ]
    ]

    # creating a new df report_tr_min with all the transactions to report with only the important columns
    report_tr_min = check_min[check_min["cumsum"].isin(report_tr["cumsum"])]
    report_tr_min = report_tr_min.drop_duplicates("1")
    i = 0
    for_word_table = []
    xlsx = pd.ExcelWriter(
        xlsx_data, engine="openpyxl", mode="a", if_sheet_exists="overlay"
    )

    # round the sum in ils
    report_tr_min["nis sum"] = report_tr_min["nis sum"].round()

    # creating a unique list of the sender names
    ls_cumsum = report_tr_min["cumsum"].unique().tolist()

    # loop each sender transactions in a unique table
    for name in ls_cumsum:
        df = report_tr_min[report_tr_min["cumsum"] == name]
        data = [["סהכ", df["nis sum"].sum()]]
        last_row = pd.DataFrame(data, columns=["clients name", "nis sum"])
        df["transaction id"] = df["transaction id"].astype(int)
        orderd = pd.concat([df, last_row])
        orderd = orderd[
            [
                "transaction id",
                "date",
                "transaction type",
                "אמצעי",
                "in/out",
                "clients name",
                "check number",
                "bank number",
                "account number",
                "branch number",
                "nis sum",
            ]
        ].fillna(" ")
        orderd = change_names_and_order(
            changemat_directory + r"/change_names.xlsx", orderd
        )
        with xlsx as writer:
            orderd.to_excel(
                writer, sheet_name="פירוט התראות ניכיון", startrow=i, index=False
            )
        for_word_table.append(orderd)
        i = i + len(orderd) + 2

    # report_tr is a DF containing all the cumsum to report for checks cleaning
    report_tr_xc = mont2r_xc_w_cumsum[["clients name", "cumsum_xc"]]
    tr_xc = tr[tr["transaction type"] == "המרה"]
    tr_min = tr_xc[
        [
            "1",
            "transaction id",
            "date",
            "transaction type",
            "אמצעי",
            "in/out",
            "clients name",
            "amount",
            "curr",
            "nis sum",
        ]
    ].merge(exchange[["transaction id", "cumsum_xc"]], on="transaction id", how="left")

    # check_min is a df with all the check transactions but with only the relevant info for report
    xc_min = tr_min[
        [
            "1",
            "transaction id",
            "date",
            "transaction type",
            "אמצעי",
            "in/out",
            "clients name",
            "amount",
            "curr",
            "nis sum",
            "cumsum_xc",
        ]
    ]

    # creating a new df report_tr_min with all the transactions to report with only the important columns
    report_tr_min_xc = xc_min[xc_min["cumsum_xc"].isin(report_tr_xc["cumsum_xc"])]
    report_tr_min_xc = report_tr_min_xc.drop_duplicates("1")

    i = 0
    for_word_table = []

    # creating a unique list of the sender names
    ls_cumsum_xc = report_tr_min_xc["cumsum_xc"].unique().tolist()

    # loop each sender transactions in a unique table
    for name in ls_cumsum_xc:
        df = report_tr_min_xc[report_tr_min_xc["cumsum_xc"] == name]
        data = [
            [
                "סהכ",
                df["nis sum"][df["curr"] == "ILS"].sum(),
                df["amount"][df["curr"] != "ILS"].sum(),
            ]
        ]
        last_row = pd.DataFrame(data, columns=["clients name", "nis sum", "amount"])
        for j in df.index:
            if df["curr"][j] == "ILS":
                df["amount"][j] = ""
            else:
                df["nis sum"][j] = ""
        df["transaction id"] = df["transaction id"].astype(int)
        orderd = pd.concat([df, last_row])
        orderd = orderd[
            [
                "transaction id",
                "date",
                "transaction type",
                "אמצעי",
                "in/out",
                "clients name",
                "curr",
                "amount",
                "nis sum",
            ]
        ].fillna(" ")

        orderd = change_names_and_order(
            changemat_directory + r"/change_names_xc.xlsx", orderd
        )
        with xlsx as writer:
            orderd.to_excel(
                writer, sheet_name="פירוט התראות המרה", startrow=i, index=False
            )
        for_word_table.append(orderd)
        i = i + len(orderd) + 2

    # report_tr is a DF containing all the cumsum to report for checks cleaning
    report_tr_ow = mont2r_risk_check_account_w_cumsum[["clients name", "cumsum_ow"]]

    tr_min = tr[
        [
            "transaction id",
            "date",
            "transaction type",
            "אמצעי",
            "in/out",
            "clients name",
            "amount",
            "curr",
            "nis sum",
        ]
    ].merge(exchange[["transaction id", "cumsum_xc"]], on="transaction id", how="left")
    # check_min is a df with all the check transactions but with only the relevant info for report

    # report_tr is a DF containing all the cumsum to report for checks cleaning
    report_tr_ow = mont2r_check_account_w_cumsum[["clients name", "cumsum_ow"]]

    # check_min is a df with all the check transactions but with only the relevant info for report
    check_min_ow = check_account[
        [
            "transaction id",
            "date",
            "transaction type",
            "אמצעי",
            "in/out",
            "clients name",
            "check number",
            "bank number",
            "account number",
            "branch number",
            "nis sum",
            "cumsum_ow",
        ]
    ]

    # creating a new df report_tr_min with all the transactions to report with only the important columns
    report_tr_min_ow = check_min_ow[
        check_min_ow["cumsum_ow"].isin(report_tr_ow["cumsum_ow"])
    ]

    i = 0
    for_word_table = []

    # round the sum in ils
    report_tr_min_ow["nis sum"] = report_tr_min_ow["nis sum"].round()

    # creating a unique list of the sender names
    ls_cumsum_ow = report_tr_min_ow["cumsum_ow"].unique().tolist()

    # loop each sender transactions in a unique table
    for name in ls_cumsum_ow:
        df = report_tr_min_ow[report_tr_min_ow["cumsum_ow"] == name]
        data = [["סהכ", df["nis sum"].sum()]]
        last_row = pd.DataFrame(data, columns=["clients name", "nis sum"])
        df["transaction id"] = df["transaction id"].astype(int)
        orderd = pd.concat([df, last_row])
        orderd = orderd[
            [
                "transaction id",
                "date",
                "transaction type",
                "אמצעי",
                "in/out",
                "clients name",
                "check number",
                "bank number",
                "account number",
                "branch number",
                "nis sum",
            ]
        ].fillna(" ")
        orderd = change_names_and_order(
            changemat_directory + r"/change_names.xlsx", orderd
        )
        with xlsx as writer:
            orderd.to_excel(
                writer, sheet_name="פירוט התראות מושך", startrow=i, index=False
            )
        for_word_table.append(orderd)
        i = i + len(orderd) + 2

    tr_for_excel = tr.merge(
        exchange[["transaction id", "cumsum_xc"]], on="transaction id", how="left"
    ).merge(
        mont2r_xc_w_cumsum[["cumsum_xc", "final type xc"]], on="cumsum_xc", how="left"
    )
    tr_for_excel = tr_for_excel.merge(
        check[["transaction id", "cumsum"]], on="transaction id", how="left"
    ).merge(
        mont2r_check_w_cumsum[["cumsum", "final type check"]], on="cumsum", how="left"
    )
    tr_for_excel = tr_for_excel.merge(
        check_account[["transaction id", "cumsum_ow"]], on="transaction id", how="left"
    ).merge(
        mont2r_check_account_w_cumsum[["cumsum_ow", "type"]], on="cumsum_ow", how="left"
    )

    tr_for_excel[["final type check", "final type xc", "type"]] = tr_for_excel[
        ["final type check", "final type xc", "type"]
    ].fillna("תקין")
    tr_for_excel = tr_for_excel.rename(
        columns={
            "final type xc": "סטטוס המרה",
            "final type check": "סטטוס נכיון",
            "type": "סטטוס מושך",
        }
    )

    tr_for_excel[["סטטוס המרה", "סטטוס נכיון", "סטטוס מושך"]] = tr_for_excel[
        ["סטטוס המרה", "סטטוס נכיון", "סטטוס מושך"]
    ].replace(
        {
            "pd": "פיצול ודירדוס",
            "pr": "פיצול בסיכון",
            "p": "פיצול",
            "dr": "דירדוס בסיכון",
            "d": "דירדוס",
            "m": "מושך משותף",
            "mr": "מושך משותף בסיכון",
            "pm": "פיצול ומושך משותף",
            "dm": "דירדוס ומושך משותף",
            "xp": "פיצול המרה",
            "xd": "דירדוס המרה",
        }
    )

    tr_for_excel["last3month"] = np.where(
        tr_for_excel["clients name"].isin(reported["שם"]),
        "דווח בשלושת החודשים האחרונים",
        "תקין",
    )
    tr_for_excel["fsp"] = np.where(
        tr_for_excel["id number"].isin(fsp_list["מספר מזהה"].astype(str)),
        'נש"מ',
        "תקין",
    )

    tr_for_excel = tr_for_excel.drop_duplicates("1")

    for i in tr_for_excel.index:
        if tr_for_excel["fsp"][i] != "תקין":
            for col in tr_for_excel.columns:
                if col.startswith("סטטוס"):
                    tr_for_excel[col][i] = 'נש"מ'
                else:
                    tr_for_excel[col][i] = tr_for_excel[col][i]

    for i in tr_for_excel.index:
        if tr_for_excel["last3month"][i] != "תקין":
            for col in tr_for_excel.columns:
                if col.startswith("סטטוס"):
                    tr_for_excel[col][i] = "דווח בשלושת החודשים האחרונים"
                else:
                    tr_for_excel[col][i] = tr_for_excel[col][i]

    tr_for_excel = tr_for_excel[
        [
            "date",
            "transaction id",
            "transaction type",
            "in/out",
            "curr",
            "אמצעי",
            "amount",
            "rate",
            "country",
            "id number",
            "clients name",
            "bank number",
            "account number",
            "branch number",
            "check number",
            "nis sum",
            "סטטוס נכיון",
            "סטטוס המרה",
            "סטטוס מושך",
        ]
    ]

    with xlsx as writer:
        tr_for_excel.to_excel(writer, sheet_name="סטטוס", index=False)

    try:
        info2r["last name"] = info2r["person name"].str.split(" ", 1, expand=True)[1]
        info2r["first name"] = info2r["person name"].str.split(" ", 1, expand=True)[0]
        info2r = info2r.fillna("")
    except KeyError:
        info2r["last name"] = ""
        info2r["first name"] = ""

    filepath_check = template_directory + "\*template*"
    txt_check = glob.glob(filepath_check)
    for textfile in txt_check:
        template = textfile
    document = docx.Document(template)
    image_paras = [
        i for i, p in enumerate(document.paragraphs) if "[ChartImage1]" in p.text
    ]
    p = document.paragraphs[image_paras[0]]
    p.text = ""
    r = p.add_run()
    r.add_text("חתימה: ").bold = True
    try:
        document.save(template_directory + "\my_doc.docx")
    except:
        document.save(template_directory + "\my_doc.docx")

    filepath_check = template_directory + "\*my_doc*"
    txt_check = glob.glob(filepath_check)
    for textfile in txt_check:
        template = textfile
    document = MailMerge(template)
    print(document.get_merge_fields())

    # getting ready for the merge
    if "person dob" not in info2r:
        info2r["person dob"] = ""
    else:
        info2r["person dob"] = info2r["person dob"]

    if "company dob" not in info2r:
        info2r["company dob"] = ""
    else:
        info2r["company dob"] = info2r["company dob"]

    if "address" not in info2r:
        info2r["address"] = ""
    else:
        info2r["address"] = info2r["address"]

    if "city" not in info2r:
        info2r["city"] = ""
    else:
        info2r["city"] = info2r["city"]

    if mont2r_check_account.empty:
        st.error("אין דיווחי מושך משותף")
    else:
        info2r_m = info2r[info2r["real final type"].str.contains("m")].merge(
            report_tr_min_ow[["cumsum_ow", "account number"]],
            on="cumsum_ow",
            how="left",
        )
        shared_owners = clients2columns_changemat(
            mont2r_check_account_w_cumsum.drop_duplicates(
                ["cumsum_ow", "clients name"]
            ),
            client,
        )
        info2r_m = info2r_m.merge(
            shared_owners, on="cumsum_ow", how="left", suffixes=("", "clients name_0")
        )
        info2r_m = info2r_m.drop_duplicates("clients name")
        info2r = info2r_m.append(info2r[-info2r["real final type"].str.contains("m")])

    columns_list = [
        "clients name_",
        "id numberclients name_",
        "cityclients name_",
        "addressclients name_",
    ]

    # Create an empty list to store the new values
    new_columns_list = []

    # Iterate over the values in the original list
    for column in columns_list:
        # Iterate over the numbers from 0 to 8
        for i in range(9):
            # Append the original value with the current number
            new_columns_list.append(column + str(i))

    if "account number" not in info2r:
        info2r["owner name"] = ""
    else:
        info2r["account number"] = info2r["account number"]

    for i in new_columns_list:
        if i not in info2r:
            info2r[i] = ""
        else:
            info2r[i] = info2r[i]

    info2r = info2r.fillna("")
    info2r = info2r.reset_index()
    doc_df = pd.DataFrame(columns=["file_name", "bytes"])

    for i in info2r.index:
        # template1 = 'test1.docx'
        document = MailMerge(template)
        document.merge(
            first_name=str(info2r["first name"][i]),
            last_name=str(info2r["last name"][i]),
            company_name=str(info2r["company name"][i]),
            person_birth=str(info2r["person dob"][i]),
            company_birth=str(info2r["company dob"][i]),
            company_id=str(info2r["company id number"][i]),
            Title=str(info2r["Title"][i]),
            person_citizenship=str(info2r["country"][i]),
            Content=str(info2r["Content"][i]),
            country=str("ישראל"),
            person_id=str(info2r["person id number"][i]),
            report_id=str(info2r["report_id"][i]),
            city=str(info2r["city"][i]),
            address=str(info2r["address"][i]),
            business_name=str(business_name),
            business_id=str(business_id),
            branch_number=str(branch_number),
            business_adress=str(business_address),
            business_phone=str(business_phone),
            worker_name=str(worker_name),
            date=str(date),
            worker_id=str(worker_id),
            workers_phone=str(worker_phone),
            workers_email=str(worker_email),
            worker_position=str(worker_position),
            business_type=str(business_type),
            clients_name_0=str(info2r["clients name_0"][i]),
            clients_name_1=str(info2r["clients name_1"][i]),
            clients_name_2=str(info2r["clients name_2"][i]),
            clients_name_3=str(info2r["clients name_3"][i]),
            clients_name_4=str(info2r["clients name_4"][i]),
            clients_name_5=str(info2r["clients name_5"][i]),
            clients_name_6=str(info2r["clients name_6"][i]),
            clients_name_7=str(info2r["clients name_7"][i]),
            clients_name_8=str(info2r["clients name_8"][i]),
            id_numberclients_name_0=str(info2r["id numberclients name_0"][i]),
            id_numberclients_name_1=str(info2r["id numberclients name_1"][i]),
            id_numberclients_name_2=str(info2r["id numberclients name_2"][i]),
            id_numberclients_name_3=str(info2r["id numberclients name_3"][i]),
            id_numberclients_name_4=str(info2r["id numberclients name_4"][i]),
            id_numberclients_name_5=str(info2r["id numberclients name_5"][i]),
            id_numberclients_name_6=str(info2r["id numberclients name_6"][i]),
            id_numberclients_name_7=str(info2r["id numberclients name_7"][i]),
            id_numberclients_name_8=str(info2r["id numberclients name_8"][i]),
            cityclients_name_0=str(info2r["cityclients name_0"][i]),
            cityclients_name_1=str(info2r["cityclients name_1"][i]),
            cityclients_name_2=str(info2r["cityclients name_2"][i]),
            cityclients_name_3=str(info2r["cityclients name_3"][i]),
            cityclients_name_4=str(info2r["cityclients name_4"][i]),
            cityclients_name_5=str(info2r["cityclients name_5"][i]),
            cityclients_name_6=str(info2r["cityclients name_6"][i]),
            cityclients_name_7=str(info2r["cityclients name_7"][i]),
            cityclients_name_8=str(info2r["cityclients name_8"][i]),
            addressclients_name_0=str(info2r["addressclients name_0"][i]),
            addressclients_name_1=str(info2r["addressclients name_1"][i]),
            addressclients_name_2=str(info2r["addressclients name_2"][i]),
            addressclients_name_3=str(info2r["addressclients name_3"][i]),
            addressclients_name_4=str(info2r["addressclients name_4"][i]),
            addressclients_name_5=str(info2r["addressclients name_5"][i]),
            addressclients_name_6=str(info2r["addressclients name_6"][i]),
            addressclients_name_7=str(info2r["addressclients name_7"][i]),
            addressclients_name_8=str(info2r["addressclients name_8"][i]),
        )

        # output =str(info2r['report_name'][i]) + "-" + str(info2r['report_id'][i]) + '.docx'
        # document.write(output)
        doc_byte = BytesIO()
        document.write(doc_byte)
        doc4df = pd.DataFrame(
            {
                "file_name": [
                    str(info2r["report_name"][i])
                    + "-"
                    + str(info2r["report_id"][i])
                    + ".docx"
                ],
                "bytes": [doc_byte],
            }
        )
        doc_df = doc_df.append(doc4df)

    info2r = info2r.set_index("cumsum")
    info2r_no_m = info2r[
        (info2r["real final type"].str.contains("p", regex=False))
        | (info2r["real final type"].str.contains("d", regex=False))
    ]

    info2r_nona = info2r_no_m[info2r_no_m.index.notnull()]

    final_reports4word = mont2r_check_w_cumsum.merge(
        info2r_nona[["clients name", "report_name", "report_id"]],
        on="clients name",
        how="left",
    )
    final_reports4word = final_reports4word.set_index("cumsum")
    doc_df = doc_df.set_index("file_name")
    for title in ls_cumsum:
        data = check[check["cumsum"] == title].round()
        table = [["סהכ", data["nis sum"].sum().round(0)]]
        last_row = pd.DataFrame(table, columns=["clients name", "nis sum"])
        df["transaction id"] = df["transaction id"].astype(int)
        orderd = pd.concat([data, last_row])
        orderd = orderd[
            [
                "transaction id",
                "transaction type",
                "date",
                "clients name",
                "check number",
                "bank number",
                "account number",
                "branch number",
                "in/out",
                "nis sum",
                "אמצעי",
            ]
        ].fillna(" ")

        orderd = change_names_and_order(
            changemat_directory + r"/change_names.xlsx", orderd
        )

        orderd["תאריך פעולה"] = orderd["תאריך פעולה"].astype(str)

        file_name = (
            str(final_reports4word["report_name"].loc[title])
            + "-"
            + str(final_reports4word["report_id"].loc[title])
            + ".docx"
        )
        doc = docx.Document(doc_df.loc[file_name].values[0])

        doc.add_page_break()
        t = doc.add_table(orderd.shape[0] + 1, orderd.shape[1])

        for j in range(orderd.shape[-1]):
            t.cell(0, j).text = orderd.columns[j]

            # add the rest of the data frame
        for i in range(orderd.shape[0]):
            for j in range(orderd.shape[-1]):
                t.cell(i + 1, j).text = str(orderd.values[i, j])

        t.style = "Grid Table 4 Accent 5"

        doc.save(doc_df.loc[file_name].values[0])

    info2r = info2r.set_index("cumsum_xc")
    info2r_no_m = info2r[
        (info2r["real final type"].str.contains("p", regex=False))
        | (info2r["real final type"].str.contains("d", regex=False))
    ]

    info2r_nona = info2r_no_m[info2r_no_m.index.notnull()]

    final_reports4word = mont2r_xc_w_cumsum.merge(
        info2r_nona[["clients name", "report_name", "report_id"]],
        on="clients name",
        how="left",
    )
    final_reports4word = final_reports4word.set_index("cumsum_xc")

    for title in ls_cumsum_xc:
        df = report_tr_min_xc[report_tr_min_xc["cumsum_xc"] == title]
        data = [
            [
                "סהכ",
                df["nis sum"][df["curr"] == "ILS"].sum(),
                df["amount"][df["curr"] != "ILS"].sum(),
            ]
        ]
        last_row = pd.DataFrame(data, columns=["clients name", "nis sum", "amount"])
        for j in df.index:
            if df["curr"][j] == "ILS":
                df["amount"][j] = ""
            else:
                df["nis sum"][j] = ""
        df["transaction id"] = df["transaction id"].astype(int)
        orderd = pd.concat([df, last_row])
        orderd = orderd[
            [
                "transaction id",
                "date",
                "transaction type",
                "אמצעי",
                "in/out",
                "clients name",
                "curr",
                "amount",
                "nis sum",
            ]
        ].fillna("")
        orderd = change_names_and_order(
            changemat_directory + r"/change_names_xc.xlsx", orderd
        )

        file_name = (
            str(final_reports4word["report_name"].loc[title])
            + "-"
            + str(final_reports4word["report_id"].loc[title])
            + ".docx"
        )

        doc = docx.Document(doc_df.loc[file_name].values[0])

        doc.add_page_break()
        t = doc.add_table(orderd.shape[0] + 1, orderd.shape[1])

        for j in range(orderd.shape[-1]):
            t.cell(0, j).text = orderd.columns[j]

            # add the rest of the data frame
        for i in range(orderd.shape[0]):
            for j in range(orderd.shape[-1]):
                t.cell(i + 1, j).text = str(orderd.values[i, j])

        t.style = "Grid Table 4 Accent 5"

        doc.save(doc_df.loc[file_name].values[0])

    info2r_only_m = info2r[(info2r["real final type"].str.contains("m", regex=False))]
    info2r_only_m = info2r_only_m.set_index("cumsum_ow")

    if info2r_only_m.empty:
        info2r_only_m = info2r_only_m
    else:
        info2r_ow = info2r_only_m[
            ["report_name", "report_id", "clients name", "account number"]
        ].merge(
            mont2r_check_account_w_cumsum[["clients name", "cumsum_ow"]],
            on="clients name",
            how="left",
        )
        info2r_ow = info2r_ow.drop_duplicates("cumsum_ow")
        final_reports4word = info2r_ow.set_index("cumsum_ow")

        for name in ls_cumsum_ow:
            df = report_tr_min_ow[report_tr_min_ow["cumsum_ow"] == name]
            data = [["סהכ", df["nis sum"].sum()]]
            last_row = pd.DataFrame(data, columns=["clients name", "nis sum"])
            df["transaction id"] = df["transaction id"].astype(int)
            orderd = pd.concat([df, last_row])
            orderd = orderd[
                [
                    "transaction id",
                    "date",
                    "transaction type",
                    "אמצעי",
                    "in/out",
                    "clients name",
                    "check number",
                    "bank number",
                    "account number",
                    "branch number",
                    "nis sum",
                ]
            ].fillna(" ")

            orderd = change_names_and_order(
                changemat_directory + r"/change_names.xlsx", orderd
            )

            file_name = (
                str(final_reports4word["report_name"].loc[name])
                + "-"
                + str(final_reports4word["report_id"].loc[name])
                + ".docx"
            )

            doc = docx.Document(doc_df.loc[file_name].values[0])

            doc.add_page_break()
            t = doc.add_table(orderd.shape[0] + 1, orderd.shape[1])

            for j in range(orderd.shape[-1]):
                t.cell(0, j).text = orderd.columns[j]

                # add the rest of the data frame
            for i in range(orderd.shape[0]):
                for j in range(orderd.shape[-1]):
                    t.cell(i + 1, j).text = str(orderd.values[i, j])

            t.style = "Grid Table 4 Accent 5"

            doc.save(doc_df.loc[file_name].values[0])

    doc_df = doc_df.append(
        pd.DataFrame(
            {"file_name": ["monitoring.xlsx"], "bytes": [xlsx_data]}
        ).set_index("file_name")
    )
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
        for file_name, file_content in doc_df.iterrows():
            # Add each file to the zip file
            file_content = file_content[0]  # Access the BytesIO object
            file_content.seek(0)  # Reset the file pointer to the beginning
            zf.writestr(file_name, file_content.read())

    # Reset the buffer's file pointer to the beginning
    zip_buffer.seek(0)

    download = st.download_button(
        label="📥 Download Reports Zip", data=zip_buffer, file_name="Reports.zip"
    )
    st.error("הניטור הושלם בהצלחה")


def run_cox(check, client, business_info, reported):
    # defining the directories of different necessary files by the directories Excel file
    directories = pd.read_excel("directories.xlsx")
    directories = pd.DataFrame(directories)
    directories = directories.set_index("File type")
    rc_directory = directories.loc["risk countries"].values[0]
    fsp_directory = directories.loc["financial service providers"].values[0]
    template_directory = directories.loc["word template"].values[0]
    cox_dirrectory = directories.loc["cox content and change names"].values[0]

    if client is None:
        client = pd.DataFrame(
            columns=[
                "N",
                "Name",
                "LastName",
                "Passport",
                "Kind",
                "BDate",
                "Country",
                "City",
                "Street",
                "Home",
            ]
        )

    # reading the financial service providers file into a df fsp_list, file must contain str "נותני"
    filepath_check = fsp_directory + "\*נותני*"
    txt_check = glob.glob(filepath_check)
    for textfile in txt_check:
        fsp_list = pd.read_excel(textfile)

    # reading the risk countries file into a DF named "rc", file must contain str "מדינות"
    filepath_check = rc_directory + "\*מדינות*"
    txt_check = glob.glob(filepath_check)
    for textfile in txt_check:
        rc = pd.read_excel(textfile)

    # reading the content file into a DF named "content", file must contain str "content"
    filepath_check = cox_dirrectory + "\*content*"
    txt_check = glob.glob(filepath_check)
    for textfile in txt_check:
        content = pd.read_excel(textfile)

    # tkinter message for files were read successfully
    st.error("הקבצים נקראו בהצלחה")

    # reported cleaning
    reported["שם"] = reported["שם"].str.strip()
    # striping the tr and client DF columns from whitespaces
    check["1"] = check.reset_index().index
    tr = check.rename(columns=lambda x: str(x).strip())
    if client is None:
        client = pd.DataFrame(
            columns=[
                "N",
                "Name",
                "LastName",
                "Passport",
                "Kind",
                "BDate",
                "Country",
                "City",
                "Street",
                "Home",
            ]
        )
    client = client.rename(columns=lambda x: str(x).strip())

    # renaming the "tr" DF columns
    tr = tr.rename(
        columns={
            "Date": "date",
            "N": "transaction id",
            "Operation": "transaction type",
            "Mode": "in/out",
            "ISO": "curr",
            "ILS": "nis sum",
            "Rate": "rate",
            "Client": "clients name",
            "Bank": "bank number",
            "Branch": "branch number",
            "Account": "account number",
            "Doc.N": "check number",
        }
    )
    tr = tr.dropna(subset=["transaction id", "Status", "transaction type"])
    tr["transaction id"] = tr["transaction id"].astype(str)
    tr["date"] = pd.to_datetime(tr["date"], dayfirst=True)
    # filter canceled transaction
    tr = tr[-tr["Status"].str.contains("בוטל")]

    # defining the "tr" date column as date in order to remove the extra timestamp
    tr["date"] = pd.to_datetime(tr["date"], errors="coerce").dt.date

    # cleaning blank spaces from clients name in clients df
    client["Name"] = client["Name"].str.strip()
    client["LastName"] = client["LastName"].str.strip()
    client["full_name"] = client["Name"] + " " + client["LastName"]
    client["Passport"] = client["Passport"].astype(str)
    client["Passport"] = client["Passport"].str.strip()

    # renaming the client df columns
    client = client.rename(
        columns={
            "Country": "country",
            "Passport": "id number",
            "full_name": "clients name",
            "Street": "address",
            "City": "city",
            "Kind": "sex",
            "BDate": "dob",
        }
    )

    # striping client country
    client["country"] = client["country"].str.strip()
    # replacing the sex column values from numbers to words
    client["sex"] = client["sex"].replace({"1": "זכר", "2": "נקבה", "3": "תאגיד"})

    # try defining the dob column as date in order to remove the extra timestamp,
    # except error if dob column is missing add dob column with blank values
    try:
        client["dob"] = pd.to_datetime(client["dob"], errors="coerce").dt.date
    except KeyError:
        client["dob"] = ""

    # fixing the check transctions client name
    tr["clients name"] = np.where(
        tr["in/out"].str.contains("צ'ק"), tr["account number"], tr["clients name"]
    )
    tr["clients name"] = np.where(
        tr["clients name"] == 0, tr["account number"], tr["clients name"]
    )
    # drop transaction with clients name missing
    tr = tr.dropna(subset="clients name")

    # drop transactions with amount missing
    tr = tr.dropna(subset="nis sum")

    # clear the id number column and the clients name column from whitespaces
    tr["clients name"] = tr["clients name"].astype(str)
    tr["clients name"] = tr["clients name"].str.strip()

    # drop transaction with clients name missing
    tr = tr[tr["clients name"] != "0"]

    # filter out clients with blank names of whitespace names
    tr = tr[tr["clients name"] != ""]
    tr = tr[tr["clients name"] != " "]

    # setting the amount sum:
    tr["amount"] = tr["nis sum"] / tr["rate"]

    # define the amount column as type int
    tr["amount"] = tr["amount"].astype(int)

    # creating a new column named id number with the last word of the clients name
    tr["id number"] = tr["clients name"].str.rsplit(" ", 1, expand=True)[1]

    # define the nis sum columns as type int64
    tr["nis sum"] = tr["nis sum"].astype("int64")
    tr["id number"] = tr["id number"].astype(str)

    # deleting leading 0 from the id number
    tr["id number"] = tr["id number"].str.lstrip("0")

    # merging the client and tr dataframes to discover new info
    tr = tr.merge(client[["country", "id number"]], on="id number", how="left")

    # risk check by merging the tr country with the risk countries df
    tr = tr.merge(
        rc[["Alpha-2 code", "סיכון/ללא דיווח רגיל"]],
        left_on="country",
        right_on="Alpha-2 code",
        how="left",
    )
    tr = tr.drop_duplicates("1")
    # striping the values inside curr and type from extra blank spaces
    tr["curr"] = tr["curr"].str.strip()
    tr["transaction type"] = tr["transaction type"].str.strip()

    # filtering only Nis currency transactions
    tr_only_nis = tr[tr["curr"] == "ILS"]

    # filtering only exchange transactions into a new DF named exchange:
    exchange = tr_only_nis[
        (tr_only_nis["transaction type"] == "פעולה")
        | (tr_only_nis["transaction type"] == "המרה בבנק")
    ]

    # sort values by clients name and date
    exchange = exchange.sort_values(["clients name", "date"])

    # adding the sum for every deal number
    deal_sum = exchange.groupby("transaction id").agg({"nis sum": "sum"})
    exchange = exchange.merge(
        deal_sum, on="transaction id", how="left", suffixes=("", "_deal")
    )
    exchange = exchange.drop_duplicates("1")
    # Adding a column stating if the transaction is under or over 50K(False = over 50K)
    exchange["U50"] = np.where(exchange["nis sum_deal"] >= 50000, False, True)

    # Adding a column stating if the transaction is under or over 5K(False = over 5K)
    exchange["U5"] = np.where(exchange["nis sum_deal"] >= 5000, False, True)

    # creating a column with the dates difference by day and creating a new column with the value as a number
    exchange["date"] = pd.to_datetime(exchange["date"])
    exchange["date_diff"] = exchange["date"].diff()
    exchange["date_diff_fl"] = exchange["date_diff"] / pd.to_timedelta(1, unit="D")

    # testing for matching clients in match column
    exchange["match"] = exchange["clients name"].eq(exchange["clients name"].shift())

    # adding a conditional column based on date diff values
    exchange["diff_check"] = np.where(
        (exchange["date_diff_fl"] > 3)
        | (exchange["date_diff_fl"] < 0)
        | (exchange["match"] == False),
        1,
        0,
    )

    # cumulative sum as a way to count if the values in the diff check
    exchange["cumsum_xc"] = np.cumsum(exchange["diff_check"])

    # risk exchange df from the exchange df by origin of client
    risk_exchange = exchange[exchange["סיכון/ללא דיווח רגיל"].notna()]

    # building the monitoring df with sum per client, deal distinct count and under 50k deal count
    mont_xc = exchange.groupby(["clients name", "cumsum_xc"]).agg(
        {"nis sum": "sum", "transaction id": "nunique", "U50": "sum"}
    )
    mont_xc = mont_xc.reset_index()

    # sorting the values by sum
    mont_xc = mont_xc.sort_values("nis sum", ascending=False)
    mont_xc["U50"] = mont_xc["U50"].astype(int)

    # don't preform a three month reported test if the DF is empty
    if reported["שם"].empty:
        mont_xc["שם"] = ""
    else:
        # preform a reported in the last two-month test
        mont_xc = mont_xc.merge(
            reported["שם"].astype(str).str.strip(),
            left_on="clients name",
            right_on="שם",
            how="left",
        )

    # adding a column named status based on unique count of deal number and at least one transaction
    # under 50k and not reported in the past 2m
    mont_xc["status"] = np.where(
        (mont_xc["transaction id"] == 1)
        | (mont_xc["U50"] == 0)
        | (mont_xc["שם"].notna()),
        "Regular Report",
        "Check",
    )
    mont_xc = mont_xc.drop_duplicates("cumsum_xc")

    # preform a financial service providers test - results stored in a column named fsp_check
    # adding the id number to the monitoring sheet
    mont_xc = mont_xc.merge(
        tr[["id number", "clients name"]], on="clients name", how="left"
    )

    # adding a column מספר מזהה to the monitoring sheet values indicates fsp(financial service provider)
    try:
        mont_xc = mont_xc.merge(
            fsp_list["מספר מזהה"], left_on="id number", right_on="מספר מזהה", how="left"
        )
    except ValueError:
        mont_xc = mont_xc.merge(
            fsp_list["מספר מזהה"].astype(str),
            left_on="id number",
            right_on="מספר מזהה",
            how="left",
        )
    mont_xc["fsp_check"] = np.where(
        mont_xc["מספר מזהה"].isna(), "check", "financial service provider"
    )

    mont_xc = mont_xc.drop_duplicates("cumsum_xc")
    # reported test
    mont_xc["reported"] = np.where(mont_xc["שם"].notna(), "reported", "check")

    # marking dirdos transaction as d in a new column named dirdos
    mont_xc["dirdos"] = np.where(
        (mont_xc["nis sum"] >= 47000)
        & (mont_xc["nis sum"] < 50000)
        & (mont_xc["שם"].isna())
        & (mont_xc["fsp_check"] == "check")
        & (mont_xc["U50"] > 0),
        "xd",
        "not xd",
    )

    # creating a new data frame containing only clients that need to be reported
    mont2r_xc = mont_xc[
        ((mont_xc["status"] == "Check") | (mont_xc["dirdos"] == "xd"))
        & (mont_xc["fsp_check"] == "check")
        & (mont_xc["nis sum"] >= 47000)
    ]
    mont2r_xc = mont2r_xc.drop_duplicates("cumsum_xc")

    # adding the type of the report p-pitzul d-dirdos
    mont2r_xc["type"] = np.where(mont2r_xc["nis sum"] >= 50000, "xp", "xd")

    # figure out the problem of a client with multiple types of reports
    type_count = mont2r_xc[["clients name", "type"]].groupby("clients name").nunique()

    # adding situations in which there are two types of reports for a client
    mont2r_xc = mont2r_xc.merge(
        type_count, on="clients name", how="left", suffixes=("", "_count")
    )
    mont2r_xc["final xc type"] = np.where(
        mont2r_xc["type_count"] > 1, "xpd", mont2r_xc["type"]
    )

    # creating a new data frame with duplicate cumsum for report details df= mont2report_w_cumsum
    mont2r_xc_w_cumsum = mont2r_xc
    mont2r_xc = mont2r_xc.drop_duplicates("clients name")
    info2r = mont2r_xc[
        ["clients name", "id number", "final xc type", "cumsum_xc", "nis sum"]
    ]

    # building the monitoring df with sum per client, deal distinct count and under 50k deal count
    mont_xc_risk = risk_exchange.groupby(["clients name", "cumsum_xc"]).agg(
        {"nis sum": "sum", "transaction id": "nunique", "U5": "sum"}
    )
    mont_xc_risk = mont_xc_risk.reset_index()

    # sorting the values by sum
    mont_xc_risk = mont_xc_risk.sort_values("nis sum", ascending=False)
    mont_xc_risk["U5"] = mont_xc_risk["U5"].astype(int)

    # preform a reported in the last two-month test
    mont_xc_risk = mont_xc_risk.merge(
        reported["שם"], left_on="clients name", right_on="שם", how="left"
    )

    # adding a column named status based on unique count of deal number
    # and at least on transactions under 50k and not reported in the past 2m
    mont_xc_risk["status"] = np.where(
        (mont_xc_risk["transaction id"] == 1)
        | (mont_xc_risk["U5"] == 0)
        | (mont_xc_risk["שם"].notna()),
        "Regular Report",
        "Check",
    )

    mont_xc_risk = mont_xc_risk.drop_duplicates("cumsum_xc")

    # preform a financial service providers test - results stored in a column named fsp_check
    # adding the id number to the monitoring sheet
    mont_xc_risk = mont_xc_risk.merge(
        tr[["id number", "clients name"]], on="clients name", how="left"
    )

    # adding a column מספר מזהה to the monitoring sheet values indicates fsp(financial service provider)
    try:
        mont_xc_risk = mont_xc_risk.merge(
            fsp_list["מספר מזהה"], left_on="id number", right_on="מספר מזהה", how="left"
        )
    except ValueError:
        mont_xc_risk = mont_xc_risk.merge(
            fsp_list["מספר מזהה"].astype(str),
            left_on="id number",
            right_on="מספר מזהה",
            how="left",
        )
    mont_xc_risk["fsp_check"] = np.where(
        mont_xc_risk["מספר מזהה"].isna(), "check", "financial service provider"
    )

    mont_xc_risk = mont_xc_risk.drop_duplicates("cumsum_xc")
    # reported test
    mont_xc_risk["reported"] = np.where(mont_xc_risk["שם"].notna(), "reported", "check")

    # marking dirdos transaction as d in a new column named dirdos
    mont_xc_risk["dirdos_risk"] = np.where(
        (mont_xc_risk["nis sum"] >= 4700)
        & (mont_xc_risk["nis sum"] < 5000)
        & (mont_xc_risk["שם"].isna())
        & (mont_xc_risk["fsp_check"] == "check")
        & (mont_xc_risk["U5"] > 0),
        "xdr",
        "not xdr",
    )

    # creating a new data frame containing only clients that need to be reported
    mont2r_xc_risk = mont_xc_risk[
        ((mont_xc_risk["status"] == "Check") | (mont_xc_risk["dirdos_risk"] == "xdr"))
        & (mont_xc_risk["fsp_check"] == "check")
        & (mont_xc_risk["nis sum"] >= 4700)
    ]

    mont2r_xc_risk = mont2r_xc_risk.drop_duplicates("cumsum_xc")

    # adding the type of the report p-pitzul d-dirdos
    mont2r_xc_risk["type_risk"] = np.where(
        mont2r_xc_risk["nis sum"] >= 5000, "xpr", "xdr"
    )

    # figure out the problem of a client with multiple types of reports
    type_count_risk = (
        mont2r_xc_risk[["clients name", "type_risk"]].groupby("clients name").nunique()
    )

    # adding situations in which there are two types of reports for a client
    mont2r_xc_risk = mont2r_xc_risk.merge(
        type_count_risk, on="clients name", how="left", suffixes=("", "_count")
    )

    mont2r_xc_risk["final type xc risk"] = np.where(
        mont2r_xc_risk["type_risk_count"] > 1, "xpdr", mont2r_xc_risk["type_risk"]
    )

    # creating a new data frame with duplicate cumsum for report details df= mont2report_w_cumsum
    mont2r_xc_w_cumsum = mont2r_xc_w_cumsum.append(mont2r_xc_risk)
    mont2r_xc_w_cumsum = mont2r_xc_w_cumsum.fillna("")
    mont2r_xc_w_cumsum["final type xc"] = (
        mont2r_xc_w_cumsum["final xc type"] + mont2r_xc_w_cumsum["final type xc risk"]
    )
    mont2r_xc_risk = mont2r_xc_risk.drop_duplicates("clients name")
    info2r = info2r.append(
        mont2r_xc_risk[
            ["clients name", "id number", "final type xc risk", "cumsum_xc", "nis sum"]
        ]
    )

    # filtering only check transactions into a new DF named check:
    check = tr_only_nis[
        (tr_only_nis["transaction type"] == "נכיון צ'קים")
        | (tr_only_nis["transaction type"] == "מסירה")
    ]
    check = check[check["in/out"].str.contains("צ'ק")]
    check["clients name"] = check["account number"]
    # sort values by clients name and date
    check = check.sort_values(["clients name", "date"])

    # adding the sum for every deal number and marking every deal under 50k as True
    deal_sum = check.groupby("transaction id").agg({"nis sum": "sum"})
    check = check.merge(
        deal_sum, on="transaction id", how="left", suffixes=("", "_deal")
    )

    # Adding a column stating if the transaction is under or over 50K(False = over 50K)
    check["U50"] = np.where(check["nis sum_deal"] >= 50000, False, True)

    # Adding a column stating if the transaction is under or over 5K(False = over 5K)
    check["U5"] = np.where(check["nis sum_deal"] >= 5000, False, True)

    # creating a column with the dates difference by day and creating a new column with the value as a number
    check["date"] = pd.to_datetime(check["date"])
    check["date_diff"] = check["date"].diff()
    check["date_diff_fl"] = check["date_diff"] / pd.to_timedelta(1, unit="D")

    # testing for matching clients in match column
    check["match"] = check["clients name"].eq(check["clients name"].shift())

    # adding a conditional colum based on date diff values
    check["diff_check"] = np.where(
        (check["date_diff_fl"] > 3)
        | (check["date_diff_fl"] < 0)
        | (check["match"] == False),
        1,
        0,
    )

    # cumulative sum as a way to count if the values in the diff check
    check["cumsum"] = np.cumsum(check["diff_check"])

    # risk check
    risk_check = check[check["סיכון/ללא דיווח רגיל"].notna()]

    # building the monitoring df with sum per client, deal distinct count and under 50k deal count
    mont_check = check.groupby(["clients name", "cumsum"]).agg(
        {"nis sum": "sum", "transaction id": "nunique", "U50": "sum"}
    )
    mont_check = mont_check.reset_index()

    # sorting the values by sum
    mont_check = mont_check.sort_values("nis sum", ascending=False)
    mont_check["U50"] = mont_check["U50"].astype(int)

    # preform a reported in the last two-month test
    mont_check = mont_check.merge(
        reported["שם"], left_on="clients name", right_on="שם", how="left"
    )

    # adding a column named status based on unique count of deal number and at least one transaction
    # under 50k and not reported in the past 2m
    mont_check["status"] = np.where(
        (mont_check["transaction id"] == 1)
        | (mont_check["U50"] == 0)
        | (mont_check["שם"].notna()),
        "Regular Report",
        "Check",
    )
    mont_check = mont_check.drop_duplicates("cumsum")

    # preform a financial service providers test - results stored in a column named fsp_check
    # adding the id number to the monitoring sheet
    mont_check = mont_check.merge(
        tr[["id number", "clients name"]], on="clients name", how="left"
    )

    # adding a column מספר מזהה to the monitoring sheet values indicates fsp(financial service provider)
    try:
        mont_check = mont_check.merge(
            fsp_list["מספר מזהה"], left_on="id number", right_on="מספר מזהה", how="left"
        )
    except ValueError:
        mont_check = mont_check.merge(
            fsp_list["מספר מזהה"].astype(str),
            left_on="id number",
            right_on="מספר מזהה",
            how="left",
        )
    mont_check["fsp_check"] = np.where(
        mont_check["מספר מזהה"].isna(), "check", "financial service provider"
    )

    mont_check = mont_check.drop_duplicates("cumsum")
    # reported test
    mont_check["reported"] = np.where(mont_check["שם"].notna(), "reported", "check")

    # marking dirdos transaction as d in a new column named dirdos
    mont_check["dirdos"] = np.where(
        (mont_check["nis sum"] >= 47000)
        & (mont_check["nis sum"] < 50000)
        & (mont_check["שם"].isna())
        & (mont_check["fsp_check"] == "check")
        & (mont_check["U50"] > 0),
        "d",
        "not d",
    )

    # creating a new data frame containing only clients that need to be reported
    mont2r_check = mont_check[
        ((mont_check["status"] == "Check") | (mont_check["dirdos"] == "d"))
        & (mont_check["fsp_check"] == "check")
        & (mont_check["nis sum"] >= 47000)
    ]
    mont2r_check = mont2r_check.drop_duplicates("cumsum")

    # adding the type of the report p-pitzul d-dirdos
    mont2r_check["type"] = np.where(mont2r_check["nis sum"] >= 50000, "p", "d")

    # figure out the problem of a client with multiple types of reports
    type_count = (
        mont2r_check[["clients name", "type"]].groupby("clients name").nunique()
    )

    # adding situations in which there are two types of reports for a client
    mont2r_check = mont2r_check.merge(
        type_count, on="clients name", how="left", suffixes=("", "_count")
    )
    mont2r_check["final type"] = np.where(
        mont2r_check["type_count"] > 1, "pd", mont2r_check["type"]
    )

    # creating a new data frame with duplicate cumsum for report details df= mont2report_w_cumsum
    mont2r_check_w_cumsum = mont2r_check
    mont2r_check = mont2r_check.drop_duplicates("clients name")

    # building the monitoring df with sum per client, deal distinct count and under 50k deal count
    mont_check_risk = risk_check.groupby(["clients name", "cumsum"]).agg(
        {"nis sum": "sum", "transaction id": "nunique", "U5": "sum"}
    )
    mont_check_risk = mont_check_risk.reset_index()

    # sorting the values by sum
    mont_check_risk = mont_check_risk.sort_values("nis sum", ascending=False)
    mont_check_risk["U5"] = mont_check_risk["U5"].astype(int)

    # preform a reported in the last two-month test
    mont_check_risk = mont_check_risk.merge(
        reported["שם"], left_on="clients name", right_on="שם", how="left"
    )

    # adding a column named status based on unique count of deal number
    # and at least on transactions under 50k and not reported in the past 2m
    mont_check_risk["status"] = np.where(
        (mont_check_risk["transaction id"] == 1)
        | (mont_check_risk["U5"] == 0)
        | (mont_check_risk["שם"].notna()),
        "Regular Report",
        "Check",
    )

    mont_check_risk = mont_check_risk.drop_duplicates("cumsum")

    # preform a financial service providers test - results stored in a column named fsp_check
    # adding the id number to the monitoring sheet
    mont_check_risk = mont_check_risk.merge(
        tr[["id number", "clients name"]], on="clients name", how="left"
    )

    # adding a column מספר מזהה to the monitoring sheet values indicates fsp(financial service provider)
    try:
        mont_check_risk = mont_check_risk.merge(
            fsp_list["מספר מזהה"], left_on="id number", right_on="מספר מזהה", how="left"
        )
    except ValueError:
        mont_check_risk = mont_check_risk.merge(
            fsp_list["מספר מזהה"].astype(str),
            left_on="id number",
            right_on="מספר מזהה",
            how="left",
        )
    mont_check_risk["fsp_check"] = np.where(
        mont_check_risk["מספר מזהה"].isna(), "check", "financial service provider"
    )

    mont_check_risk = mont_check_risk.drop_duplicates("cumsum")
    # reported test
    mont_check_risk["reported"] = np.where(
        mont_check_risk["שם"].notna(), "reported", "check"
    )

    # marking dirdos transaction as d in a new column named dirdos
    mont_check_risk["dirdos_risk"] = np.where(
        (mont_check_risk["nis sum"] >= 4700)
        & (mont_check_risk["nis sum"] < 5000)
        & (mont_check_risk["שם"].isna())
        & (mont_check_risk["fsp_check"] == "check")
        & (mont_check_risk["U5"] > 0),
        "dr",
        "not dr",
    )

    # creating a new data frame containing only clients that need to be reported
    mont2r_check_risk = mont_check_risk[
        (
            (mont_check_risk["status"] == "Check")
            | (mont_check_risk["dirdos_risk"] == "dr")
        )
        & (mont_check_risk["fsp_check"] == "check")
        & (mont_check_risk["nis sum"] >= 4700)
    ]

    mont2r_check_risk = mont2r_check_risk.drop_duplicates("cumsum")

    # adding the type of the report p-pitzul d-dirdos
    mont2r_check_risk["type_risk"] = np.where(
        mont2r_check_risk["nis sum"] >= 5000, "pr", "dr"
    )

    # figure out the problem of a client with multiple types of reports
    type_count_risk = (
        mont2r_check_risk[["clients name", "type_risk"]]
        .groupby("clients name")
        .nunique()
    )

    # adding situations in which there are two types of reports for a client
    mont2r_check_risk = mont2r_check_risk.merge(
        type_count_risk, on="clients name", how="left", suffixes=("", "_count")
    )

    mont2r_check_risk["final type risk"] = np.where(
        mont2r_check_risk["type_risk_count"] > 1, "pdr", mont2r_check_risk["type_risk"]
    )
    mont2r_check_w_cumsum = mont2r_check_w_cumsum.append(mont2r_check_risk)
    mont2r_check_w_cumsum = mont2r_check_w_cumsum.fillna("")
    mont2r_check_w_cumsum["final type check"] = (
        mont2r_check_w_cumsum["final type"] + mont2r_check_w_cumsum["final type risk"]
    )
    mont2r_check_risk = mont2r_check_risk.drop_duplicates("clients name")

    info2r = info2r.append(
        mont2r_check[["clients name", "id number", "final type", "cumsum", "nis sum"]]
    )
    info2r = info2r.append(
        mont2r_check_risk[
            ["clients name", "id number", "final type risk", "cumsum", "nis sum"]
        ]
    )

    check["account number"] = check["branch number"]
    check["branch number"] = check["bank number"]
    check["bank number"] = check["check number"]
    check["check number"] = check["in/out"].replace({"'נכנס צ'ק": ""})
    check_account = check.dropna(subset="account number")

    # sort values by clients name and date
    check_account = check_account.sort_values(["account number", "date"])

    # adding the sum for every deal number and marking every deal under 50k as True
    deal_sum = check_account.groupby("transaction id").agg({"nis sum": "sum"})
    check_account = check_account.merge(
        deal_sum, on="transaction id", how="left", suffixes=("", "_deal")
    )

    # Adding a column stating if the transaction is under or over 50K(False = over 50K)
    check_account["U50"] = np.where(check_account["nis sum_deal"] >= 50000, False, True)

    # Adding a column stating if the transaction is under or over 5K(False = over 5K)
    check_account["U5"] = np.where(check_account["nis sum_deal"] >= 5000, False, True)

    # creating a column with the dates difference by day and creating a new column with the value as a number
    check_account["date"] = pd.to_datetime(check_account["date"])
    check_account["date_diff"] = check_account["date"].diff()
    check_account["date_diff_fl"] = check_account["date_diff"] / pd.to_timedelta(
        1, unit="D"
    )

    # testing for matching clients in match column
    check_account["match"] = check_account["account number"].eq(
        check_account["account number"].shift()
    )

    # adding a conditional colum based on date diff values
    check_account["diff_check"] = np.where(
        (check_account["date_diff_fl"] > 3)
        | (check_account["date_diff_fl"] < 0)
        | (check_account["match"] == False),
        1,
        0,
    )

    # cumulative sum as a way to count if the values in the diff check
    check_account["cumsum_ow"] = np.cumsum(check_account["diff_check"])

    # risk check
    risk_check_account = check_account[check_account["סיכון/ללא דיווח רגיל"].notna()]

    # building the monitoring df with sum per client, deal distinct count and under 50k deal count
    mont_check_account = check_account.groupby(["account number", "cumsum_ow"]).agg(
        {
            "nis sum": "sum",
            "transaction id": "nunique",
            "U50": "sum",
            "clients name": "nunique",
        }
    )

    mont_check_account = mont_check_account.rename(
        columns={"clients name": "clients name_count"}
    )

    mont_check_account = mont_check_account.reset_index()

    # sorting the values by sum
    mont_check_account = mont_check_account.sort_values("nis sum", ascending=False)
    mont_check_account["U50"] = mont_check_account["U50"].astype(int)

    mont_check_account = mont_check_account.merge(
        check_account[["cumsum_ow", "clients name"]], on="cumsum_ow", how="left"
    )
    # preform a reported in the last two-month test
    mont_check_account = mont_check_account.merge(
        reported["שם"], left_on="clients name", right_on="שם", how="left"
    )

    # adding a column named status based on unique count of deal number and at least one transaction
    # under 50k and not reported in the past 2m
    mont_check_account["status"] = np.where(
        (mont_check_account["clients name_count"] == 1)
        | (mont_check_account["U50"] == 0)
        | (mont_check_account["שם"].notna()),
        "Regular Report",
        "Check",
    )

    # preform a financial service providers test - results stored in a column named fsp_check
    # adding the id number to the monitoring sheet
    mont_check_account = mont_check_account.merge(
        check_account[["id number", "clients name"]], on="clients name", how="left"
    )

    # adding a column מספר מזהה to the monitoring sheet values indicates fsp(financial service provider)
    try:
        mont_check_account = mont_check_account.merge(
            fsp_list["מספר מזהה"], left_on="id number", right_on="מספר מזהה", how="left"
        )
    except ValueError:
        mont_check_account = mont_check_account.merge(
            fsp_list["מספר מזהה"].astype(str),
            left_on="id number",
            right_on="מספר מזהה",
            how="left",
        )
    mont_check_account["fsp_check"] = np.where(
        mont_check_account["מספר מזהה"].isna(), "check", "financial service provider"
    )

    # reported test
    mont_check_account["reported"] = np.where(
        mont_check_account["שם"].notna(), "reported", "check"
    )

    # creating a new data frame containing only clients that need to be reported
    mont2r_check_account = mont_check_account[
        (mont_check_account["status"] == "Check")
        & (mont_check_account["fsp_check"] == "check")
        & (mont_check_account["nis sum"] >= 50000)
    ]

    # adding the type of the report p-pitzul d-dirdos m-mosheh
    mont2r_check_account["type"] = "m"

    # creating a new data frame with duplicate cumsum for report details df= mont2report_w_cumsum
    mont2r_check_account_w_cumsum = mont2r_check_account

    mont2r_check_account = mont2r_check_account[
        mont2r_check_account["reported"] == "check"
    ]
    mont2r_check_account = mont2r_check_account[
        -mont2r_check_account["clients name"].isin(info2r["clients name"])
    ]

    mont2r_check_account = mont2r_check_account.drop_duplicates("cumsum_ow")

    info2r = info2r.append(
        mont2r_check_account[
            ["clients name", "id number", "type", "cumsum_ow", "nis sum"]
        ]
    )

    # building the monitoring df with sum per client, deal distinct count and under 50k deal count
    mont_risk_check_account = risk_check_account.groupby(
        ["account number", "cumsum_ow"]
    ).agg(
        {
            "nis sum": "sum",
            "transaction id": "nunique",
            "U5": "sum",
            "clients name": "nunique",
        }
    )

    mont_risk_check_account = mont_risk_check_account.rename(
        columns={"clients name": "clients name_count"}
    )

    mont_risk_check_account = mont_risk_check_account.reset_index()

    # sorting the values by sum
    mont_risk_check_account = mont_risk_check_account.sort_values(
        "nis sum", ascending=False
    )
    mont_risk_check_account["U50"] = mont_risk_check_account["U5"].astype(int)

    mont_risk_check_account = mont_risk_check_account.merge(
        risk_check_account[["cumsum_ow", "clients name"]], on="cumsum_ow", how="left"
    )
    # preform a reported in the last two-month test
    mont_risk_check_account = mont_risk_check_account.merge(
        reported["שם"], left_on="clients name", right_on="שם", how="left"
    )

    # adding a column named status based on unique count of deal number and at least one transaction
    # under 50k and not reported in the past 2m
    mont_risk_check_account["status"] = np.where(
        (mont_risk_check_account["clients name_count"] == 1)
        | (mont_risk_check_account["U5"] == 0)
        | (mont_risk_check_account["שם"].notna()),
        "Regular Report",
        "Check",
    )

    # preform a financial service providers test - results stored in a column named fsp_check
    # adding the id number to the monitoring sheet
    mont_risk_check_account = mont_risk_check_account.merge(
        risk_check_account[["id number", "clients name"]], on="clients name", how="left"
    )

    # adding a column מספר מזהה to the monitoring sheet values indicates fsp(financial service provider)
    try:
        mont_risk_check_account = mont_risk_check_account.merge(
            fsp_list["מספר מזהה"], left_on="id number", right_on="מספר מזהה", how="left"
        )
    except ValueError:
        mont_risk_check_account = mont_risk_check_account.merge(
            fsp_list["מספר מזהה"].astype(str),
            left_on="id number",
            right_on="מספר מזהה",
            how="left",
        )
    mont_risk_check_account["fsp_check"] = np.where(
        mont_risk_check_account["מספר מזהה"].isna(),
        "check",
        "financial service provider",
    )

    # reported test
    mont_risk_check_account["reported"] = np.where(
        mont_risk_check_account["שם"].notna(), "reported", "check"
    )

    # creating a new data frame containing only clients that need to be reported
    mont2r_risk_check_account = mont_risk_check_account[
        (mont_risk_check_account["status"] == "Check")
        & (mont_risk_check_account["fsp_check"] == "check")
        & (mont_risk_check_account["nis sum"] >= 5000)
    ]

    # adding the type of the report p-pitzul d-dirdos m-mosheh
    mont2r_risk_check_account["type risk"] = "mr"

    # creating a new data frame with duplicate cumsum for report details df= mont2report_w_cumsum
    mont2r_risk_check_account_w_cumsum = mont2r_risk_check_account

    mont2r_risk_check_account = mont2r_risk_check_account[
        mont2r_risk_check_account["reported"] == "check"
    ]
    mont2r_risk_check_account = mont2r_risk_check_account[
        -mont2r_risk_check_account["clients name"].isin(info2r["clients name"])
    ]

    mont2r_risk_check_account = mont2r_risk_check_account.drop_duplicates("cumsum_ow")

    info2r = info2r.append(
        mont2r_risk_check_account[
            ["clients name", "id number", "type risk", "cumsum_ow", "nis sum"]
        ]
    )

    for column in info2r.columns:
        if column.startswith("cumsum"):
            info2r[column] = info2r[column].fillna("")
            info2r[column] = info2r[column].apply(
                lambda x: round(x) if isinstance(x, float) else x
            )

    info2r["nis sum"] = info2r["nis sum"].astype("int64")

    info2r = info2r.fillna("")
    info2r["real final type"] = (
        info2r["final xc type"]
        + info2r["final type xc risk"]
        + info2r["final type"]
        + info2r["final type risk"]
        + info2r["type"]
        + info2r["type risk"]
    )
    if info2r.empty:
        tr_for_excel = tr.merge(
            exchange[["transaction id", "cumsum_xc"]], on="transaction id", how="left"
        ).merge(
            mont2r_xc_w_cumsum[["cumsum_xc", "final type xc"]],
            on="cumsum_xc",
            how="left",
        )
        tr_for_excel = tr_for_excel.merge(
            check[["transaction id", "cumsum"]], on="transaction id", how="left"
        ).merge(
            mont2r_check_w_cumsum[["cumsum", "final type check"]],
            on="cumsum",
            how="left",
        )
        tr_for_excel = tr_for_excel.merge(
            check_account[["transaction id", "cumsum_ow"]],
            on="transaction id",
            how="left",
        ).merge(
            mont2r_check_account_w_cumsum[["cumsum_ow", "type"]],
            on="cumsum_ow",
            how="left",
        )
        tr_for_excel[["final type check", "final type xc", "type"]] = tr_for_excel[
            ["final type check", "final type xc", "type"]
        ].fillna("תקין")
        tr_for_excel = tr_for_excel.rename(
            columns={
                "final type xc": "סטטוס המרה",
                "final type check": "סטטוס נכיון",
                "type": "סטטוס מושך",
            }
        )

        tr_for_excel[["סטטוס המרה", "סטטוס נכיון", "סטטוס מושך"]] = tr_for_excel[
            ["סטטוס המרה", "סטטוס נכיון", "סטטוס מושך"]
        ].replace(
            {
                "pd": "פיצול ודירדוס",
                "pr": "פיצול בסיכון",
                "p": "פיצול",
                "dr": "דירדוס בסיכון",
                "d": "דירדוס",
                "m": "מושך משותף",
                "mr": "מושך משותף בסיכון",
                "pm": "פיצול ומושך משותף",
                "dm": "דירדוס ומושך משותף",
                "xp": "פיצול המרה",
                "xd": "דירדוס המרה",
            }
        )

        tr_for_excel["last3month"] = np.where(
            tr_for_excel["clients name"].isin(reported["שם"]),
            "דווח בשלושת החודשים האחרונים",
            "תקין",
        )
        tr_for_excel["fsp"] = np.where(
            tr_for_excel["id number"].isin(fsp_list["מספר מזהה"].astype(str)),
            'נש"מ',
            "תקין",
        )

        for i in tr_for_excel.index:
            if tr_for_excel["fsp"][i] != "תקין":
                for col in tr_for_excel.columns:
                    if col.startswith("סטטוס"):
                        tr_for_excel[col][i] = 'נש"מ'

        for i in tr_for_excel.index:
            if tr_for_excel["last3month"][i] != "תקין":
                for col in tr_for_excel.columns:
                    if col.startswith("סטטוס"):
                        tr_for_excel[col][i] = "דווח בשלושת החודשים האחרונים"

        tr_for_excel = tr_for_excel.drop_duplicates("1")
        tr_for_excel = tr_for_excel[
            [
                "date",
                "transaction id",
                "transaction type",
                "in/out",
                "curr",
                "amount",
                "rate",
                "country",
                "id number",
                "clients name",
                "bank number",
                "account number",
                "branch number",
                "check number",
                "nis sum",
                "סטטוס נכיון",
                "סטטוס המרה",
                "סטטוס מושך",
            ]
        ]
        xlsx_data = BytesIO()
        with pd.ExcelWriter(xlsx_data, engine="openpyxl") as writer:
            tr_for_excel.to_excel(writer, sheet_name="סטטוס", index=False)
        st.download_button(
            label="📥 Download Final Report",
            data=xlsx_data,
            file_name="Final Report.xlsx",
        )
        st.error("אין דיווחים בלתי רגילים")
        exit()
    else:
        st.error("יש דיווחים בלתי רגילים")

    # todo this try except might cause problems with the clients details in reports
    try:
        info2r["id number"] = info2r["clients name"].str.rsplit(" ", 1, expand=True)[1]
        info2r["id number"] = info2r["id number"].str.lstrip("0")
    except KeyError:
        tkinter.messagebox.showinfo(
            message="הייתה בעיה בהפרדה בין שם הלקוח למספר מזהה, אם אין דיווחים בלתי רגילים ניתן להתעלם מהודעה זאת"
        )
    try:
        info2r = info2r[
            [
                "clients name",
                "cumsum",
                "cumsum_ow",
                "cumsum_xc",
                "nis sum",
                "real final type",
                "id number",
            ]
        ].merge(
            client[["id number", "dob", "sex", "country", "city", "address"]],
            on="id number",
            how="left",
        )
    except KeyError:
        info2r = info2r[
            [
                "clients name",
                "cumsum",
                "cumsum_ow",
                "cumsum_xc",
                "nis sum",
                "real final type",
            ]
        ].merge(
            client[["clients name", "id number", "country"]],
            on="clients name",
            how="left",
        )

    info2r["clients name only"] = info2r["clients name"].str.rsplit(
        " ", 1, expand=True
    )[0]

    # creating columns for private customer or company based on the sex column
    info2r["person name"] = np.where(
        info2r["sex"] == "תאגיד", "", info2r["clients name only"]
    )
    info2r["company name"] = np.where(
        info2r["sex"] == "תאגיד", info2r["clients name only"], ""
    )
    info2r["person dob"] = np.where(info2r["sex"] == "תאגיד", "", info2r["dob"])
    info2r["company dob"] = np.where(info2r["sex"] == "תאגיד", info2r["dob"], "")
    info2r["person id number"] = np.where(
        info2r["sex"] == "תאגיד", "", info2r["id number"]
    )
    info2r["company id number"] = np.where(
        info2r["sex"] == "תאגיד", info2r["id number"], ""
    )

    # merging the content to the info2r DF
    info2r = info2r.merge(
        content, left_on="real final type", right_on="type", how="left"
    )

    # fixing the business info
    business_info = business_info
    business_info = business_info.rename(columns=lambda x: str(x).strip())
    business_info = business_info.fillna("")
    business_type = business_info.loc["סוג הגוף המדווח"].values[0]
    business_name = business_info.loc["שם מלא של הגורם המדווח"].values[0]
    business_id = business_info.loc["מספר מזהה של הגורם המדווח"].values[0]
    branch_number = business_info.loc["מספר סניף מדווח"].values[0]
    business_address = business_info.loc["מען הסניף"].values[0]
    date = to_date.today()
    business_phone = business_info.loc["טלפון ופקסימילה של הגורם המדווח"].values[0]
    worker_name = business_info.loc["שם פרטי ושם משפחה של עובד עורך הדיווח"].values[0]
    worker_id = business_info.loc['ת"ז של עורך הדיווח'].values[0]
    worker_position = business_info.loc["תפקיד עורך הדיווח"].values[0]
    worker_phone = business_info.loc["טלפון עורך הדיווח"].values[0]
    worker_email = business_info.loc["דואר אלקטרוני"].values[0]

    # creating a countries dictionary
    info2r = info2r.merge(
        rc[["Alpha-2 code", "עברית"]],
        left_on="country",
        right_on="Alpha-2 code",
        how="left",
    )

    # adding the country of origin to the report content
    for i in info2r.index:
        info2r["Title"].iloc[i] = str(info2r["Title"].iloc[i]).replace(
            "מדינה", str(info2r["עברית"].iloc[i]) + ", מדינה"
        )
        info2r["Content"].iloc[i] = str(info2r["Content"].iloc[i]).replace(
            "מדינה", str(info2r["עברית"].iloc[i]) + ", מדינה"
        )

    info2r["country"].notna = info2r["עברית"]
    info2r = info2r.drop(columns=["עברית", "Alpha-2 code"])

    # dropping duplicates and resetting index
    info2r = info2r.drop_duplicates("clients name")
    info2r = info2r.reset_index(drop=True)
    # importing the report number needed
    info2r["report_id_aid"] = reported["מס דיווח"].max() + 1
    info2r["report_id"] = (info2r["report_id_aid"] + range(len(info2r.index))).astype(
        int
    )
    info2r["report_name"] = report_num_gen(business_id, branch_number)

    # creating a DF named info2r_excel with the column for the final report
    info2r_excel = info2r[
        [
            "clients name",
            "real final type",
            "id number",
            "cumsum",
            "cumsum_xc",
            "cumsum_ow",
            "nis sum",
            "dob",
            "sex",
            "country",
            "report_id",
            "Title",
            "Content",
        ]
    ]

    # renaming the columns of info2r_excel DF
    info2r_excel = info2r_excel.rename(
        columns={
            "clients name": "שם הלקוח",
            "real final type": "מהות הפעילות",
            "id number": "מספר תעודה",
            "dob": "תאריך לידה\התאגדות",
            "sex": "מין",
            "country": "מדינת תעודה\התאגדות",
            "report_id": "מספר דיווח",
            "Title": "תמצית",
            "Content": "תוכן",
            "cumsum": "מזהה קבוצה ניכיון",
            "cumsum_ow": "מזהה קבוצה מושך",
            "cumsum_xc": "מזהה קבוצה המרה",
            "nis sum": "סכום פעולות לדיווח",
        }
    )

    # creating new empty columns for the customer to write notes
    info2r_excel["לדווח/לא לדווח"] = ""
    info2r_excel["הערות"] = ""

    # replacing the content in מהות הםעילות from english shortcut to heb real name
    info2r_excel["מהות הפעילות"] = info2r_excel["מהות הפעילות"].replace(
        {
            "pd": "פיצול ודירדוס",
            "pr": "פיצול בסיכון",
            "p": "פיצול",
            "dr": "דירדוס בסיכון",
            "d": "דירדוס",
            "m": "מושך משותף",
            "mr": "מושך משותף בסיכון",
            "pm": "פיצול ומושך משותף",
            "dm": "דירדוס ומושך משותף",
            "xp": "פיצול המרה",
            "xd": "דירדוס המרה",
            "xpr": "פיצול המרה בסיכון",
            "xdr": "דירדוס המרה בסיכון,",
        }
    )

    info2r_excel["מין"] = info2r_excel["מין"].replace(
        {"1": "זכר", "2": "נקבה", "3": "תאגיד"}
    )

    # saving info2r_excel as an Excel sheet named ניהול התראות in a xlsx file named monitoring at the transactions directory
    xlsx_data = BytesIO()
    with pd.ExcelWriter(xlsx_data, engine="openpyxl") as writer:
        info2r_excel.to_excel(writer, sheet_name="ניהול התראות")

    # report_tr is a DF containing all the cumsum to report for checks cleaning
    report_tr = mont2r_check_w_cumsum[["clients name", "cumsum"]]

    # check_min is a df with all the check transactions but with only the relevant info for report
    check_min = check[
        [
            "1",
            "transaction id",
            "date",
            "transaction type",
            "in/out",
            "clients name",
            "check number",
            "bank number",
            "account number",
            "branch number",
            "nis sum",
            "cumsum",
        ]
    ]

    # creating a new df report_tr_min with all the transactions to report with only the important columns
    report_tr_min = check_min[check_min["cumsum"].isin(report_tr["cumsum"])]
    report_tr_min = report_tr_min.drop_duplicates("1")
    i = 0
    for_word_table = []
    xlsx = pd.ExcelWriter(
        xlsx_data, engine="openpyxl", mode="a", if_sheet_exists="overlay"
    )
    # round the sum in ils
    report_tr_min["nis sum"] = report_tr_min["nis sum"].round()

    # creating a unique list of the sender names
    ls_cumsum = report_tr_min["cumsum"].unique().tolist()

    # loop each sender transactions in a unique table
    for name in ls_cumsum:
        df = report_tr_min[report_tr_min["cumsum"] == name]
        data = [["סהכ", df["nis sum"].sum()]]
        last_row = pd.DataFrame(data, columns=["clients name", "nis sum"])
        orderd = pd.concat([df, last_row])
        orderd = orderd[
            [
                "transaction id",
                "date",
                "transaction type",
                "in/out",
                "clients name",
                "check number",
                "bank number",
                "account number",
                "branch number",
                "nis sum",
            ]
        ].fillna(" ")
        orderd = change_names_and_order(cox_dirrectory + r"/change_names.xlsx", orderd)
        with xlsx as writer:
            orderd.to_excel(
                writer, sheet_name="פירוט התראות ניכיון", startrow=i, index=False
            )
        for_word_table.append(orderd)
        i = i + len(orderd) + 2

    # report_tr is a DF containing all the cumsum to report for checks cleaning
    report_tr_xc = mont2r_xc_w_cumsum[["clients name", "cumsum_xc"]]
    tr_xc = tr[tr["transaction type"] == "פעולה"]
    tr_min = tr_xc[
        [
            "1",
            "transaction id",
            "date",
            "transaction type",
            "in/out",
            "clients name",
            "amount",
            "curr",
            "nis sum",
        ]
    ].merge(exchange[["transaction id", "cumsum_xc"]], on="transaction id", how="left")

    # check_min is a df with all the check transactions but with only the relevant info for report
    xc_min = tr_min[
        [
            "1",
            "transaction id",
            "date",
            "transaction type",
            "in/out",
            "clients name",
            "amount",
            "curr",
            "nis sum",
            "cumsum_xc",
        ]
    ]

    # creating a new df report_tr_min with all the transactions to report with only the important columns
    report_tr_min_xc = xc_min[xc_min["cumsum_xc"].isin(report_tr_xc["cumsum_xc"])]
    report_tr_min_xc = report_tr_min_xc.drop_duplicates("1")

    i = 0
    for_word_table = []

    # creating a unique list of the sender names
    ls_cumsum_xc = report_tr_min_xc["cumsum_xc"].unique().tolist()

    # loop each sender transactions in a unique table
    for name in ls_cumsum_xc:
        df = report_tr_min_xc[report_tr_min_xc["cumsum_xc"] == name]
        data = [
            [
                "סהכ",
                df["nis sum"][df["curr"] == "ILS"].sum(),
                df["amount"][df["curr"] != "ILS"].sum(),
            ]
        ]
        last_row = pd.DataFrame(data, columns=["clients name", "nis sum", "amount"])
        for j in df.index:
            if df["curr"][j] == "ILS":
                df["amount"][j] = ""
            else:
                df["nis sum"][j] = ""
        orderd = pd.concat([df, last_row])
        orderd = orderd[
            [
                "transaction id",
                "date",
                "transaction type",
                "in/out",
                "clients name",
                "curr",
                "amount",
                "nis sum",
            ]
        ].fillna(" ")

        orderd = change_names_and_order(
            cox_dirrectory + r"/change_names_xc.xlsx", orderd
        )
        with xlsx as writer:
            orderd.to_excel(
                writer, sheet_name="פירוט התראות המרה", startrow=i, index=False
            )
        for_word_table.append(orderd)
        i = i + len(orderd) + 2

    # report_tr is a DF containing all the cumsum to report for checks cleaning
    report_tr_ow = mont2r_risk_check_account_w_cumsum[["clients name", "cumsum_ow"]]

    tr_min = tr[
        [
            "transaction id",
            "date",
            "transaction type",
            "in/out",
            "clients name",
            "amount",
            "curr",
            "nis sum",
        ]
    ].merge(exchange[["transaction id", "cumsum_xc"]], on="transaction id", how="left")
    # check_min is a df with all the check transactions but with only the relevant info for report

    # report_tr is a DF containing all the cumsum to report for checks cleaning
    report_tr_ow = mont2r_check_account_w_cumsum[["clients name", "cumsum_ow"]]

    # check_min is a df with all the check transactions but with only the relevant info for report
    check_min_ow = check_account[
        [
            "transaction id",
            "date",
            "transaction type",
            "in/out",
            "clients name",
            "check number",
            "bank number",
            "account number",
            "branch number",
            "nis sum",
            "cumsum_ow",
        ]
    ]

    # creating a new df report_tr_min with all the transactions to report with only the important columns
    report_tr_min_ow = check_min_ow[
        check_min_ow["cumsum_ow"].isin(report_tr_ow["cumsum_ow"])
    ]

    i = 0
    for_word_table = []

    # round the sum in ils
    report_tr_min_ow["nis sum"] = report_tr_min_ow["nis sum"].round()

    # creating a unique list of the sender names
    ls_cumsum_ow = report_tr_min_ow["cumsum_ow"].unique().tolist()

    # loop each sender transactions in a unique table
    for name in ls_cumsum_ow:
        df = report_tr_min_ow[report_tr_min_ow["cumsum_ow"] == name]
        data = [["סהכ", df["nis sum"].sum()]]
        last_row = pd.DataFrame(data, columns=["clients name", "nis sum"])
        orderd = pd.concat([df, last_row])
        orderd = orderd[
            [
                "transaction id",
                "date",
                "transaction type",
                "in/out",
                "clients name",
                "check number",
                "bank number",
                "account number",
                "branch number",
                "nis sum",
            ]
        ].fillna(" ")
        orderd = change_names_and_order(cox_dirrectory + r"/change_names.xlsx", orderd)
        with xlsx as writer:
            orderd.to_excel(
                writer, sheet_name="פירוט התראות מושך", startrow=i, index=False
            )
        for_word_table.append(orderd)
        i = i + len(orderd) + 2

    tr_for_excel = tr.merge(
        exchange[["transaction id", "cumsum_xc"]], on="transaction id", how="left"
    ).merge(
        mont2r_xc_w_cumsum[["cumsum_xc", "final type xc"]], on="cumsum_xc", how="left"
    )
    tr_for_excel = tr_for_excel.merge(
        check[["transaction id", "cumsum"]], on="transaction id", how="left"
    ).merge(
        mont2r_check_w_cumsum[["cumsum", "final type check"]], on="cumsum", how="left"
    )
    tr_for_excel = tr_for_excel.merge(
        check_account[["transaction id", "cumsum_ow"]], on="transaction id", how="left"
    ).merge(
        mont2r_check_account_w_cumsum[["cumsum_ow", "type"]], on="cumsum_ow", how="left"
    )

    tr_for_excel[["final type check", "final type xc", "type"]] = tr_for_excel[
        ["final type check", "final type xc", "type"]
    ].fillna("תקין")
    tr_for_excel = tr_for_excel.rename(
        columns={
            "final type xc": "סטטוס המרה",
            "final type check": "סטטוס נכיון",
            "type": "סטטוס מושך",
        }
    )

    tr_for_excel[["סטטוס המרה", "סטטוס נכיון", "סטטוס מושך"]] = tr_for_excel[
        ["סטטוס המרה", "סטטוס נכיון", "סטטוס מושך"]
    ].replace(
        {
            "pd": "פיצול ודירדוס",
            "pr": "פיצול בסיכון",
            "p": "פיצול",
            "dr": "דירדוס בסיכון",
            "d": "דירדוס",
            "m": "מושך משותף",
            "mr": "מושך משותף בסיכון",
            "pm": "פיצול ומושך משותף",
            "dm": "דירדוס ומושך משותף",
            "xp": "פיצול המרה",
            "xd": "דירדוס המרה",
        }
    )

    tr_for_excel["last3month"] = np.where(
        tr_for_excel["clients name"].isin(reported["שם"]),
        "דווח בשלושת החודשים האחרונים",
        "תקין",
    )
    tr_for_excel["fsp"] = np.where(
        tr_for_excel["id number"].isin(fsp_list["מספר מזהה"].astype(str)),
        'נש"מ',
        "תקין",
    )

    for i in tr_for_excel.index:
        if tr_for_excel["fsp"][i] != "תקין":
            for col in tr_for_excel.columns:
                if col.startswith("סטטוס"):
                    tr_for_excel[col][i] = 'נש"מ'

    for i in tr_for_excel.index:
        if tr_for_excel["last3month"][i] != "תקין":
            for col in tr_for_excel.columns:
                if col.startswith("סטטוס"):
                    tr_for_excel[col][i] = "דווח בשלושת החודשים האחרונים"

    tr_for_excel = tr_for_excel.drop_duplicates("1")
    tr_for_excel = tr_for_excel[
        [
            "date",
            "transaction id",
            "transaction type",
            "in/out",
            "curr",
            "amount",
            "rate",
            "country",
            "id number",
            "clients name",
            "bank number",
            "account number",
            "branch number",
            "check number",
            "nis sum",
            "סטטוס נכיון",
            "סטטוס המרה",
            "סטטוס מושך",
        ]
    ]
    with xlsx as writer:
        tr_for_excel.to_excel(writer, sheet_name="סטטוס", index=False)

    try:
        info2r["last name"] = info2r["person name"].str.split(" ", 1, expand=True)[1]
        info2r["first name"] = info2r["person name"].str.split(" ", 1, expand=True)[0]
        info2r = info2r.fillna("")
    except KeyError:
        info2r["last name"] = ""
        info2r["first name"] = ""

    filepath_check = template_directory + "\*template*"
    txt_check = glob.glob(filepath_check)
    for textfile in txt_check:
        template = textfile
    document = docx.Document(template)
    image_paras = [
        i for i, p in enumerate(document.paragraphs) if "[ChartImage1]" in p.text
    ]
    p = document.paragraphs[image_paras[0]]
    p.text = ""
    r = p.add_run()
    r.add_text("חתימה: ").bold = True
    try:
        document.save(template_directory + "\my_doc.docx")
    except:
        document.save(template_directory + "\my_doc.docx")

    filepath_check = template_directory + "\*my_doc*"
    txt_check = glob.glob(filepath_check)
    for textfile in txt_check:
        template = textfile
    document = MailMerge(template)
    print(document.get_merge_fields())

    # getting ready for the merge
    if "person dob" not in info2r:
        info2r["person dob"] = ""
    else:
        info2r["person dob"] = info2r["person dob"]

    if "company dob" not in info2r:
        info2r["company dob"] = ""
    else:
        info2r["company dob"] = info2r["company dob"]

    if "address" not in info2r:
        info2r["address"] = ""
    else:
        info2r["address"] = info2r["address"]

    if "city" not in info2r:
        info2r["city"] = ""
    else:
        info2r["city"] = info2r["city"]

    if mont2r_check_account.empty:
        st.error("אין דיווחי מושך משותף")
    else:
        info2r_m = info2r[info2r["real final type"].str.contains("m")].merge(
            report_tr_min_ow[["cumsum_ow", "account number"]],
            on="cumsum_ow",
            how="left",
        )
        shared_owners = clients2columns_changemat(
            mont2r_check_account_w_cumsum.drop_duplicates(
                ["cumsum_ow", "clients name"]
            ),
            client,
        )
        info2r_m = info2r_m.merge(
            shared_owners, on="cumsum_ow", how="left", suffixes=("", "clients name_0")
        )
        info2r_m = info2r_m.drop_duplicates("clients name")
        info2r = info2r_m.append(info2r[-info2r["real final type"].str.contains("m")])

    columns_list = [
        "clients name_",
        "id numberclients name_",
        "cityclients name_",
        "addressclients name_",
    ]

    # Create an empty list to store the new values
    new_columns_list = []

    # Iterate over the values in the original list
    for column in columns_list:
        # Iterate over the numbers from 0 to 8
        for i in range(9):
            # Append the original value with the current number
            new_columns_list.append(column + str(i))

    if "account number" not in info2r:
        info2r["owner name"] = ""
    else:
        info2r["account number"] = info2r["account number"]

    for i in new_columns_list:
        if i not in info2r:
            info2r[i] = ""
        else:
            info2r[i] = info2r[i]

    info2r = info2r.fillna("")
    info2r = info2r.reset_index()
    doc_df = pd.DataFrame(columns=["file_name", "bytes"])
    for i in info2r.index:
        # template1 = 'test1.docx'
        document = MailMerge(template)
        document.merge(
            first_name=str(info2r["first name"][i]),
            last_name=str(info2r["last name"][i]),
            company_name=str(info2r["company name"][i]),
            person_birth=str(info2r["person dob"][i]),
            company_birth=str(info2r["company dob"][i]),
            company_id=str(info2r["company id number"][i]),
            Title=str(info2r["Title"][i]),
            person_citizenship=str(info2r["country"][i]),
            Content=str(info2r["Content"][i]),
            country=str("ישראל"),
            person_id=str(info2r["person id number"][i]),
            report_id=str(info2r["report_id"][i]),
            city=str(info2r["city"][i]),
            address=str(info2r["address"][i]),
            business_name=str(business_name),
            business_id=str(business_id),
            branch_number=str(branch_number),
            business_adress=str(business_address),
            business_phone=str(business_phone),
            worker_name=str(worker_name),
            date=str(date),
            worker_id=str(worker_id),
            workers_phone=str(worker_phone),
            workers_email=str(worker_email),
            worker_position=str(worker_position),
            business_type=str(business_type),
            clients_name_0=str(info2r["clients name_0"][i]),
            clients_name_1=str(info2r["clients name_1"][i]),
            clients_name_2=str(info2r["clients name_2"][i]),
            clients_name_3=str(info2r["clients name_3"][i]),
            clients_name_4=str(info2r["clients name_4"][i]),
            clients_name_5=str(info2r["clients name_5"][i]),
            clients_name_6=str(info2r["clients name_6"][i]),
            clients_name_7=str(info2r["clients name_7"][i]),
            clients_name_8=str(info2r["clients name_8"][i]),
            id_numberclients_name_0=str(info2r["id numberclients name_0"][i]),
            id_numberclients_name_1=str(info2r["id numberclients name_1"][i]),
            id_numberclients_name_2=str(info2r["id numberclients name_2"][i]),
            id_numberclients_name_3=str(info2r["id numberclients name_3"][i]),
            id_numberclients_name_4=str(info2r["id numberclients name_4"][i]),
            id_numberclients_name_5=str(info2r["id numberclients name_5"][i]),
            id_numberclients_name_6=str(info2r["id numberclients name_6"][i]),
            id_numberclients_name_7=str(info2r["id numberclients name_7"][i]),
            id_numberclients_name_8=str(info2r["id numberclients name_8"][i]),
            cityclients_name_0=str(info2r["cityclients name_0"][i]),
            cityclients_name_1=str(info2r["cityclients name_1"][i]),
            cityclients_name_2=str(info2r["cityclients name_2"][i]),
            cityclients_name_3=str(info2r["cityclients name_3"][i]),
            cityclients_name_4=str(info2r["cityclients name_4"][i]),
            cityclients_name_5=str(info2r["cityclients name_5"][i]),
            cityclients_name_6=str(info2r["cityclients name_6"][i]),
            cityclients_name_7=str(info2r["cityclients name_7"][i]),
            cityclients_name_8=str(info2r["cityclients name_8"][i]),
            addressclients_name_0=str(info2r["addressclients name_0"][i]),
            addressclients_name_1=str(info2r["addressclients name_1"][i]),
            addressclients_name_2=str(info2r["addressclients name_2"][i]),
            addressclients_name_3=str(info2r["addressclients name_3"][i]),
            addressclients_name_4=str(info2r["addressclients name_4"][i]),
            addressclients_name_5=str(info2r["addressclients name_5"][i]),
            addressclients_name_6=str(info2r["addressclients name_6"][i]),
            addressclients_name_7=str(info2r["addressclients name_7"][i]),
            addressclients_name_8=str(info2r["addressclients name_8"][i]),
        )

        # create a new doc_byte write the report inside and add it to the doc4df dataframe
        doc_byte = BytesIO()
        document.write(doc_byte)
        doc4df = pd.DataFrame(
            {
                "file_name": [
                    str(info2r["report_name"][i])
                    + "-"
                    + str(info2r["report_id"][i])
                    + ".docx"
                ],
                "bytes": [doc_byte],
            }
        )
        doc_df = doc_df.append(doc4df)

    info2r = info2r.set_index("cumsum")
    info2r_no_m = info2r[
        (info2r["real final type"].str.contains("p", regex=False))
        | (info2r["real final type"].str.contains("d", regex=False))
    ]

    info2r_nona = info2r_no_m[info2r_no_m.index.notnull()]

    final_reports4word = mont2r_check_w_cumsum.merge(
        info2r_nona[["clients name", "report_name", "report_id"]],
        on="clients name",
        how="left",
    )
    final_reports4word = final_reports4word.set_index("cumsum")
    doc_df = doc_df.set_index("file_name")

    for title in ls_cumsum:
        data = check[check["cumsum"] == title].round()
        table = [["סהכ", data["nis sum"].sum().round(0)]]
        last_row = pd.DataFrame(table, columns=["clients name", "nis sum"])
        orderd = pd.concat([data, last_row])
        orderd = orderd[
            [
                "transaction id",
                "transaction type",
                "date",
                "clients name",
                "check number",
                "bank number",
                "account number",
                "branch number",
                "in/out",
                "nis sum",
            ]
        ].fillna(" ")

        orderd = change_names_and_order(cox_dirrectory + r"/change_names.xlsx", orderd)

        orderd["תאריך פעולה"] = orderd["תאריך פעולה"].astype(str)

        file_name = (
            str(final_reports4word["report_name"].loc[title])
            + "-"
            + str(final_reports4word["report_id"].loc[title])
            + ".docx"
        )
        doc = docx.Document(doc_df.loc[file_name].values[0])

        doc.add_page_break()
        t = doc.add_table(orderd.shape[0] + 1, orderd.shape[1])

        for j in range(orderd.shape[-1]):
            t.cell(0, j).text = orderd.columns[j]

            # add the rest of the data frame
        for i in range(orderd.shape[0]):
            for j in range(orderd.shape[-1]):
                t.cell(i + 1, j).text = str(orderd.values[i, j])

        t.style = "Grid Table 4 Accent 5"

        doc.save(doc_df.loc[file_name].values[0])

    info2r = info2r.set_index("cumsum_xc")
    info2r_no_m = info2r[
        (info2r["real final type"].str.contains("p", regex=False))
        | (info2r["real final type"].str.contains("d", regex=False))
    ]

    info2r_nona = info2r_no_m[info2r_no_m.index.notnull()]

    final_reports4word = mont2r_xc_w_cumsum.merge(
        info2r_nona[["clients name", "report_name", "report_id"]],
        on="clients name",
        how="left",
    )
    final_reports4word = final_reports4word.set_index("cumsum_xc")

    for title in ls_cumsum_xc:
        df = report_tr_min_xc[report_tr_min_xc["cumsum_xc"] == title]
        data = [
            [
                "סהכ",
                df["nis sum"][df["curr"] == "ILS"].sum(),
                df["amount"][df["curr"] != "ILS"].sum(),
            ]
        ]
        last_row = pd.DataFrame(data, columns=["clients name", "nis sum", "amount"])
        for j in df.index:
            if df["curr"][j] == "ILS":
                df["amount"][j] = ""
            else:
                df["nis sum"][j] = ""
        orderd = pd.concat([df, last_row])
        orderd = orderd[
            [
                "transaction id",
                "date",
                "transaction type",
                "in/out",
                "clients name",
                "curr",
                "amount",
                "nis sum",
            ]
        ].fillna("")
        orderd = change_names_and_order(
            cox_dirrectory + r"/change_names_xc.xlsx", orderd
        )

        file_name = (
            str(final_reports4word["report_name"].loc[title])
            + "-"
            + str(final_reports4word["report_id"].loc[title])
            + ".docx"
        )
        doc = docx.Document(doc_df.loc[file_name].values[0])

        doc.add_page_break()
        t = doc.add_table(orderd.shape[0] + 1, orderd.shape[1])

        for j in range(orderd.shape[-1]):
            t.cell(0, j).text = orderd.columns[j]

            # add the rest of the data frame
        for i in range(orderd.shape[0]):
            for j in range(orderd.shape[-1]):
                t.cell(i + 1, j).text = str(orderd.values[i, j])

        t.style = "Grid Table 4 Accent 5"

        doc.save(doc_df.loc[file_name].values[0])

    info2r = info2r.set_index("cumsum_ow")
    info2r_only_m = info2r[(info2r["real final type"].str.contains("m", regex=False))]
    if info2r_only_m.empty:
        info2r_only_m = info2r_only_m
    else:
        info2r_ow = info2r_only_m[
            ["report_name", "report_id", "clients name", "account number"]
        ].merge(
            mont2r_check_account_w_cumsum[["clients name", "cumsum_ow"]],
            on="clients name",
            how="left",
        )
        info2r_ow = info2r_ow.drop_duplicates("cumsum_ow")
        final_reports4word = info2r_ow.set_index("cumsum_ow")

        for title in ls_cumsum_ow:
            df = report_tr_min_ow[report_tr_min_ow["cumsum_ow"] == name]
            data = [["סהכ", df["nis sum"].sum()]]
            last_row = pd.DataFrame(data, columns=["clients name", "nis sum"])
            orderd = pd.concat([df, last_row])
            orderd = orderd[
                [
                    "transaction id",
                    "date",
                    "transaction type",
                    "in/out",
                    "clients name",
                    "check number",
                    "bank number",
                    "account number",
                    "branch number",
                    "nis sum",
                ]
            ].fillna(" ")

            orderd = change_names_and_order(
                cox_dirrectory + r"/change_names.xlsx", orderd
            )

            file_name = (
                str(final_reports4word["report_name"].loc[title])
                + "-"
                + str(final_reports4word["report_id"].loc[title])
                + ".docx"
            )
            doc = docx.Document(doc_df.loc[file_name].values[0])

            doc.add_page_break()
            t = doc.add_table(orderd.shape[0] + 1, orderd.shape[1])

            for j in range(orderd.shape[-1]):
                t.cell(0, j).text = orderd.columns[j]

                # add the rest of the data frame
            for i in range(orderd.shape[0]):
                for j in range(orderd.shape[-1]):
                    t.cell(i + 1, j).text = str(orderd.values[i, j])

            t.style = "Grid Table 4 Accent 5"

            doc.save(doc_df.loc[file_name].values[0])

    # Create an in-memory zip file
    doc_df = doc_df.append(
        pd.DataFrame(
            {"file_name": ["monitoring.xlsx"], "bytes": [xlsx_data]}
        ).set_index("file_name")
    )
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
        for file_name, file_content in doc_df.iterrows():
            # Add each file to the zip file
            file_content = file_content[0]  # Access the BytesIO object
            file_content.seek(0)  # Reset the file pointer to the beginning
            zf.writestr(file_name, file_content.read())

    # Reset the buffer's file pointer to the beginning
    zip_buffer.seek(0)

    download = st.download_button(
        label="📥 Download Reports Zip", data=zip_buffer, file_name="Reports.zip"
    )

    st.error("הניטור הושלם בהצלחה")
